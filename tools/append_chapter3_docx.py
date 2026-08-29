import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


INPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2.docx")
OUTPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2_3.docx")


CHAPTER3 = r"""
# CHƯƠNG 3. KẾT QUẢ THỰC HIỆN

## 3.1. Chuẩn bị phát hành

Sau quá trình lập kế hoạch, phân tích yêu cầu và triển khai theo các Sprint đã trình bày ở Chương 2, hệ thống SocialContent Hub đã hoàn thiện các nhóm chức năng cốt lõi phục vụ quy trình tự động hóa sản xuất nội dung video ngắn bằng AI. Các chức năng đã được phát triển theo từng mốc phát hành gồm Release V1.0, Release V2.0 và giai đoạn Production.

Ở Release V1.0, hệ thống tập trung hoàn thiện nền tảng xác thực, phân quyền người dùng, thu thập dữ liệu bài viết từ VNExpress, chuẩn hóa dữ liệu thô, gom nhóm nội dung theo truyện/tập, đánh giá chất lượng dữ liệu và lập kế hoạch nội dung bằng AI. Đây là nền tảng đầu vào quan trọng để các phân hệ tạo video và đăng tải nội dung có dữ liệu sạch, có cấu trúc và có thể kiểm soát chất lượng.

Ở Release V2.0, hệ thống hoàn thiện pipeline sản xuất video gồm tạo kịch bản, chia cảnh, tạo giọng đọc, sinh phụ đề, đồng bộ timeline 1:1 giữa hình ảnh - âm thanh - phụ đề, render video MP4, xem trước và duyệt video. Đồng thời, hệ thống bổ sung chức năng kết nối tài khoản TikTok, cấu hình chiến lược kênh, lên lịch đăng bài, đăng bài thủ công và tự động đăng bài theo lịch.

Ở giai đoạn Production, hệ thống được chuẩn bị để vận hành ổn định trên môi trường triển khai chính thức với các thành phần microservices, cơ sở dữ liệu PostgreSQL, MongoDB, Kafka, background scheduler, trang quản trị hệ thống, audit logs và cấu hình vận hành.

**Bảng 3.1. Tổng hợp thành phần chuẩn bị phát hành**

| Thành phần | Mục tiêu | Kết quả thực hiện |
|---|---|---|
| Frontend Dashboard | Cung cấp giao diện cho người dùng sáng tạo nội dung và quản trị viên | Hoàn thành các màn hình Login, Dashboard, Crawl, Content, Planning, Generate Video, Approvals, Publishing Queue, Settings và Users |
| API Service | Cung cấp REST API, xác thực, phân quyền, quản lý người dùng và điều phối nghiệp vụ | Hoàn thành API xác thực, users, crawl jobs, contents, planning runs, media workflows, social profiles, publishing queue và admin |
| Data Ingestion Engine | Thu thập và xử lý dữ liệu đầu vào | Hoàn thành luồng tạo crawl job, xử lý job, lấy dữ liệu VNExpress, chuẩn hóa, ghi log và phát sự kiện qua Kafka |
| AI Media Engine | Lập kế hoạch nội dung và sản xuất video | Hoàn thành AI planning, tạo candidate, chia nội dung nhiều phần, tạo voice, fit timeline và điều phối tạo video |
| Remotion Worker | Render video MP4 từ dữ liệu timeline | Hoàn thành worker Node.js phục vụ ghép media, subtitle và xuất video |
| PostgreSQL | Lưu dữ liệu nghiệp vụ chuẩn hóa | Hoàn thành các bảng users, roles, crawl_jobs, content_items, stories, episodes, media_workflow, planning_runs, publishing_queue_items, audit_logs |
| MongoDB | Lưu dữ liệu trung gian, dữ liệu xử lý và planning input/output | Hoàn thành lưu processed_documents, planning_inputs, planning_outputs và dữ liệu phục vụ xử lý AI |
| Kafka | Kết nối bất đồng bộ giữa các service | Hoàn thành các topic cho crawl, normalization, planning, video generation, publishing và dead-letter |
| Scheduler & Admin | Tự động hóa vận hành và hỗ trợ quản trị | Hoàn thành cấu hình interval, theo dõi trạng thái scheduler, audit logs và quản trị người dùng |

Trước khi phát hành, các nhóm chức năng chính được rà soát theo tiêu chí: có giao diện hoặc API sử dụng được, có dữ liệu lưu trữ ổn định, có phân quyền theo vai trò, có trạng thái xử lý rõ ràng, có khả năng ghi log khi lỗi và có thể kiểm thử theo kịch bản người dùng.

## 3.2. Kiểm thử

Quá trình kiểm thử được thực hiện nhằm xác nhận các chức năng đã triển khai đáp ứng đúng yêu cầu nghiệp vụ, đồng thời kiểm tra khả năng phối hợp giữa frontend, backend, cơ sở dữ liệu và các worker xử lý nền. Do hệ thống có nhiều thành phần bất đồng bộ, việc kiểm thử không chỉ tập trung vào từng API riêng lẻ mà còn kiểm tra toàn bộ luồng dữ liệu từ lúc người dùng tạo yêu cầu đến khi hệ thống trả về kết quả cuối cùng.

### 3.2.1. Kiểm thử tích hợp

Kiểm thử tích hợp tập trung vào khả năng giao tiếp giữa các thành phần: React frontend, FastAPI API Service, PostgreSQL, MongoDB, Kafka, Data Ingestion Engine, AI Media Engine và Remotion Worker. Các kịch bản kiểm thử được thiết kế theo hướng bám sát luồng vận hành thực tế của hệ thống.

**Bảng 3.2. Kết quả kiểm thử tích hợp**

| Nhóm kiểm thử | Kịch bản kiểm thử | Kết quả mong đợi | Trạng thái |
|---|---|---|---|
| Xác thực và phân quyền | Người dùng đăng nhập bằng tài khoản hợp lệ | API trả về JWT token, frontend lưu phiên làm việc và điều hướng vào Dashboard | Đạt |
| Xác thực và phân quyền | Người dùng không đủ quyền truy cập tài nguyên của người khác | API phản hồi 403 Forbidden, dữ liệu không bị rò rỉ giữa các tài khoản | Đạt |
| Crawl VNExpress | Người dùng tạo crawl job từ nguồn VNExpress | Crawl job được lưu vào PostgreSQL, phát event crawl.job.created qua Kafka | Đạt |
| Điều phối crawl | Data Ingestion Engine nhận job và xử lý nguồn dữ liệu | Trạng thái job chuyển qua PENDING/RUNNING/COMPLETED hoặc FAILED, log được ghi nhận | Đạt |
| Chuẩn hóa dữ liệu | Bài viết thô được xử lý sau khi crawl | Nội dung được làm sạch, metadata chuẩn hóa và bản ghi được lưu phục vụ AI planning | Đạt |
| Deduplication | Hệ thống gặp nội dung hoặc ảnh trùng lặp | Fingerprint phát hiện trùng lặp, cập nhật duplicate_count và quality_score | Đạt |
| AI Planning | Người dùng tạo content project và yêu cầu AI đề xuất chủ đề | AI trả về chủ đề, góc triển khai, kiểu video và kế hoạch nhiều phần | Đạt |
| Video Pipeline | Người dùng tạo video từ kế hoạch đã duyệt | Hệ thống tạo story, scene, voice, subtitle, timeline và render job | Đạt |
| Render MP4 | Remotion Worker xử lý dữ liệu timeline | File MP4 được xuất ra đúng cấu trúc, có hình ảnh, audio và phụ đề | Đạt |
| Publishing Queue | Người dùng lên lịch hoặc đăng bài thủ công | Queue item được tạo, cập nhật trạng thái queued/publishing/published/failed | Đạt |
| Scheduler | Đến thời điểm đăng bài tự động | Background Scheduler quét hàng đợi và thực hiện đăng bài khi thỏa điều kiện | Đạt |
| Audit Logs | Người dùng tạo job, chỉnh cấu hình hoặc đăng bài | Hệ thống ghi lại actor, action, target, thời gian và metadata liên quan | Đạt |

Kết quả kiểm thử tích hợp cho thấy các thành phần chính đã kết nối được với nhau theo đúng kiến trúc microservices. Đặc biệt, cơ chế Kafka giúp tách biệt các tác vụ xử lý nền, tránh việc API phải chờ toàn bộ quá trình crawl, planning hoặc render video hoàn tất trong một request duy nhất.

### 3.2.2. Kiểm thử chấp nhận

Kiểm thử chấp nhận được thực hiện dựa trên các user story đã xác định trong Product Backlog. Mục tiêu là xác nhận hệ thống đáp ứng được nhu cầu sử dụng của hai nhóm vai trò chính: người dùng sáng tạo nội dung và quản trị viên/vận hành hệ thống.

**Bảng 3.3. Kết quả kiểm thử chấp nhận theo nghiệp vụ**

| Nghiệp vụ | Tiêu chí chấp nhận | Kết quả |
|---|---|---|
| Đăng nhập và sử dụng hệ thống | Người dùng đăng nhập được, chỉ sử dụng chức năng theo quyền hạn | Đạt |
| Thu thập dữ liệu | Người dùng tạo được private crawl job, theo dõi trạng thái và xem log xử lý | Đạt |
| Chuẩn hóa và đánh giá chất lượng | Dữ liệu crawl được làm sạch, đánh dấu READY/NEEDS_REVIEW và phát hiện trùng lặp | Đạt |
| Tạo dự án nội dung | Người dùng tạo được content project từ bài viết hoặc series đã đạt chất lượng | Đạt |
| AI đề xuất kế hoạch | AI đề xuất được chủ đề, góc khai thác, kiểu video và chia nội dung thành nhiều phần | Đạt |
| Duyệt kế hoạch | Người dùng chỉnh sửa, duyệt, từ chối hoặc yêu cầu tạo lại kế hoạch | Đạt |
| Tạo video | Hệ thống tạo được kịch bản, scene, voice, subtitle và timeline phục vụ render | Đạt |
| Xem trước và duyệt video | Người dùng xem preview video, kiểm tra lỗi và duyệt trước khi đăng | Đạt |
| Kết nối tài khoản TikTok | Người dùng tạo phiên QR, xác thực và lưu social profile | Đạt |
| Đăng bài thủ công | Người dùng chủ động đăng ngay video đã duyệt | Đạt |
| Đăng bài tự động | Scheduler tự động đăng video khi đến thời gian cấu hình và kênh bật auto publish | Đạt |
| Quản trị hệ thống | Quản trị viên xem người dùng, audit logs và cấu hình scheduler | Đạt |

Nhìn chung, hệ thống đáp ứng được phạm vi chức năng đã đặt ra trong đồ án. Các lỗi phát sinh chủ yếu thuộc nhóm cấu hình môi trường, kết nối API bên thứ ba hoặc dữ liệu đầu vào chưa đủ chất lượng; các trường hợp này đã được xử lý bằng cơ chế thông báo lỗi, trạng thái FAILED/NEEDS_REVIEW và log chi tiết.

### 3.2.3. Kiểm thử giao diện người dùng

Kiểm thử giao diện được thực hiện trên các màn hình chính của frontend React/Vite. Mục tiêu là bảo đảm người dùng có thể thao tác xuyên suốt quy trình tạo nội dung mà không phải gọi API thủ công.

**Bảng 3.4. Kết quả kiểm thử giao diện**

| Màn hình | Chức năng kiểm thử | Kết quả mong đợi | Trạng thái |
|---|---|---|---|
| Login | Đăng nhập, hiển thị lỗi sai tài khoản | Người dùng vào được hệ thống hoặc nhận thông báo lỗi rõ ràng | Đạt |
| Dashboard | Xem thống kê tổng quan và trạng thái vận hành | Dữ liệu tổng quan hiển thị đúng theo API | Đạt |
| Crawl | Tạo job, xem danh sách job và log | Người dùng theo dõi được tiến độ crawl | Đạt |
| Content | Xem nội dung đã chuẩn hóa và chi tiết dữ liệu nguồn | Nội dung, metadata và trạng thái hiển thị đầy đủ | Đạt |
| Planning | Tạo content project, yêu cầu AI planning, duyệt kế hoạch | Kế hoạch AI hiển thị và cập nhật trạng thái đúng | Đạt |
| Generate Video | Quản lý workflow tạo video | Người dùng theo dõi được các bước tạo scene, voice, subtitle và render | Đạt |
| Approvals | Xem trước và duyệt nội dung | Video đủ điều kiện mới được chuyển sang bước đăng bài | Đạt |
| Publishing Queue | Lên lịch, đăng ngay, xem trạng thái queue | Queue item cập nhật đúng trạng thái xử lý | Đạt |
| Settings | Cấu hình chiến lược kênh và scheduler | Tham số được lưu và áp dụng cho xử lý nền | Đạt |
| Users | Quản lý người dùng và vai trò | Admin thao tác được theo quyền SYSTEM_ADMIN | Đạt |

### 3.2.4. Đánh giá kết quả kiểm thử

Kết quả kiểm thử cho thấy hệ thống có thể vận hành theo luồng end-to-end: đăng nhập, tạo job crawl dữ liệu, chuẩn hóa dữ liệu, lập kế hoạch nội dung bằng AI, tạo video, duyệt nội dung, đưa vào hàng đợi và đăng bài. Các phân hệ sử dụng cơ chế trạng thái rõ ràng nên người dùng và quản trị viên có thể theo dõi tiến độ xử lý, nhận biết lỗi và thực hiện lại thao tác khi cần.

Một số rủi ro còn tồn tại gồm phụ thuộc vào chất lượng dữ liệu từ nguồn VNExpress, phụ thuộc vào độ ổn định của AI provider, thời gian render video có thể tăng khi dữ liệu đầu vào dài và giới hạn từ TikTok API khi đăng bài trực tiếp. Các rủi ro này đã được ghi nhận để tiếp tục tối ưu trong các phiên bản sau.

## 3.3. Hướng dẫn sử dụng phần mềm

### 3.3.1. Giới thiệu tổng quan

SocialContent Hub là hệ thống hỗ trợ tự động hóa quy trình sản xuất nội dung video ngắn bằng AI. Người dùng có thể thu thập dữ liệu bài viết từ nguồn tin tức, chuẩn hóa dữ liệu, tạo dự án nội dung, dùng AI đề xuất chủ đề và kế hoạch video, chuyển kế hoạch thành video, kiểm duyệt và đăng tải nội dung lên TikTok theo hình thức thủ công hoặc tự động theo lịch.

Hệ thống có hai nhóm người dùng chính:

- Người dùng sáng tạo nội dung (CREATOR): tạo job crawl dữ liệu, quản lý dữ liệu đã thu thập, lập kế hoạch nội dung bằng AI, tạo video, duyệt video, kết nối tài khoản TikTok và quản lý lịch đăng bài.
- Quản trị viên/vận hành hệ thống (SYSTEM_ADMIN): quản lý người dùng, giám sát toàn bộ dữ liệu, kiểm tra audit logs, cấu hình scheduler và theo dõi tình trạng hoạt động của hệ thống.

### 3.3.2. Yêu cầu hệ thống

Đối với người dùng cuối, hệ thống chỉ yêu cầu trình duyệt web hiện đại như Google Chrome, Microsoft Edge hoặc Safari. Người dùng cần có kết nối Internet ổn định khi sử dụng các chức năng liên quan đến AI, tải video hoặc đăng bài lên nền tảng TikTok.

Đối với môi trường triển khai, hệ thống yêu cầu máy chủ có Docker và Docker Compose để chạy các service, PostgreSQL có pgvector để lưu dữ liệu nghiệp vụ và embedding, MongoDB để lưu dữ liệu xử lý trung gian, Kafka để truyền sự kiện bất đồng bộ, Node.js/Remotion phục vụ render video và Python/FastAPI cho các service backend.

**Bảng 3.5. Yêu cầu môi trường triển khai**

| Thành phần | Phiên bản/công cụ đề xuất | Mục đích sử dụng |
|---|---|---|
| Docker, Docker Compose | Phiên bản mới ổn định | Đóng gói và điều phối các service |
| PostgreSQL + pgvector | PostgreSQL 16 | Lưu user, crawl job, content, workflow, planning, queue, audit logs và embedding |
| MongoDB | MongoDB 7 | Lưu processed documents, planning input/output và dữ liệu trung gian |
| Kafka, Zookeeper | Confluent Kafka 7.4 | Truyền sự kiện giữa API, ingestion, planning, video worker và scheduler |
| Python | 3.11 trở lên | Chạy FastAPI services, crawler, AI planning và video pipeline |
| Node.js | 18 trở lên | Chạy frontend React/Vite và Remotion Worker |
| Trình duyệt | Chrome/Edge/Safari mới | Sử dụng giao diện hệ thống |

### 3.3.3. Hướng dẫn cài đặt và khởi chạy

Quy trình cài đặt hệ thống được thực hiện theo các bước chính sau:

1. Tải mã nguồn dự án về máy chủ hoặc máy phát triển.
2. Cấu hình tệp môi trường `.env` cho backend, bao gồm `DATABASE_URL`, `MONGO_URI`, `KAFKA_BOOTSTRAP_SERVERS`, API key cho AI provider và thông tin TikTok OAuth nếu sử dụng chức năng đăng bài.
3. Khởi chạy hạ tầng bằng Docker Compose tại thư mục backend với lệnh `docker compose up --build`.
4. Kiểm tra API Service thông qua endpoint `/health`, bảo đảm service trả về trạng thái hoạt động.
5. Khởi chạy frontend bằng lệnh `npm install` và `npm run dev` trong thư mục frontend.
6. Truy cập giao diện web tại địa chỉ local của Vite, đăng nhập tài khoản và bắt đầu sử dụng các chức năng theo vai trò.

Trong môi trường phát triển, API Service mặc định phục vụ các endpoint dưới tiền tố `/api/v1`, frontend React/Vite chạy ở cổng 5173, PostgreSQL ở cổng 5432, MongoDB ở cổng 27017, Kafka ở cổng 9092 và pgAdmin ở cổng 5050.

### 3.3.4. Luồng sử dụng dành cho người sáng tạo nội dung

Người dùng sáng tạo nội dung thao tác theo luồng sau:

1. Đăng nhập vào hệ thống.
2. Tạo crawl job bằng cách chọn nguồn VNExpress, nhập URL/từ khóa/chuyên mục và cấu hình số lượng dữ liệu cần lấy.
3. Theo dõi tiến độ crawl, trạng thái job và nhật ký xử lý.
4. Kiểm tra dữ liệu đã chuẩn hóa, điểm chất lượng và trạng thái READY/NEEDS_REVIEW.
5. Tạo content project từ bài viết hoặc series đã đạt chất lượng.
6. Yêu cầu AI đề xuất chủ đề, góc triển khai và kiểu video.
7. Duyệt, chỉnh sửa hoặc yêu cầu tạo lại kế hoạch nội dung.
8. Chuyển kế hoạch sang bước tạo video, tạo scene, voice, subtitle và render MP4.
9. Xem trước video, kiểm tra chất lượng và duyệt nội dung.
10. Kết nối tài khoản TikTok, cấu hình chiến lược kênh và chọn đăng ngay hoặc lên lịch đăng tự động.

### 3.3.5. Luồng sử dụng dành cho quản trị viên

Quản trị viên sử dụng hệ thống để giám sát và vận hành:

1. Đăng nhập bằng tài khoản có vai trò SYSTEM_ADMIN.
2. Quản lý người dùng, tạo tài khoản, cập nhật trạng thái hoặc phân quyền.
3. Theo dõi toàn bộ crawl job, nội dung, workflow, queue và trạng thái xử lý.
4. Kiểm tra audit logs để truy vết thao tác tạo job, duyệt nội dung, đăng bài hoặc thay đổi cấu hình.
5. Cấu hình thời gian quét của Background Scheduler như `vnexpress_interval_minutes` và `publish_queue_interval_minutes`.
6. Theo dõi tình trạng scheduler, số lượng item đã publish, failed hoặc skipped để xử lý sự cố khi cần.

## 3.4. Phát hành và triển khai

Hệ thống được triển khai theo hướng containerization, trong đó mỗi service được đóng gói độc lập bằng Docker và giao tiếp với nhau thông qua mạng nội bộ cùng các biến môi trường cấu hình. Cách triển khai này phù hợp với kiến trúc microservices vì từng thành phần có thể được build, khởi động, giám sát và mở rộng tương đối độc lập.

### 3.4.1. Đóng gói ứng dụng bằng Docker

Backend được chia thành nhiều service độc lập: API Service, Data Ingestion Engine, AI Media Engine và Remotion Worker. Các service Python sử dụng FastAPI hoặc worker nền để xử lý nghiệp vụ, trong khi Remotion Worker sử dụng Node.js để render video MP4. PostgreSQL, MongoDB, Kafka và Zookeeper được chạy bằng image chính thức hoặc image phù hợp với yêu cầu của hệ thống.

Frontend được phát triển bằng React, TypeScript và Vite. Trong môi trường production, frontend có thể được build thành static files và phục vụ qua Nginx hoặc một web server tương đương. Việc tách frontend và backend giúp quá trình triển khai linh hoạt hơn, đồng thời dễ cấu hình CORS, reverse proxy và domain riêng cho giao diện người dùng.

**Bảng 3.6. Thành phần Docker trong hệ thống**

| Service | Công nghệ | Vai trò |
|---|---|---|
| postgres | PostgreSQL + pgvector | Lưu dữ liệu nghiệp vụ quan hệ và embedding |
| mongodb | MongoDB | Lưu dữ liệu xử lý trung gian, dữ liệu crawl/planning |
| kafka, zookeeper | Kafka event broker | Điều phối sự kiện bất đồng bộ |
| api-service | FastAPI | REST API, auth, users, social profiles, crawl jobs, planning, publishing |
| data-ingestion-worker | FastAPI/worker | Crawl, normalize, grouping, dedup và ghi log xử lý dữ liệu |
| ai-media-planning-worker | FastAPI/worker | AI planning, chọn chủ đề, chia nội dung nhiều phần |
| ai-media-worker | Python worker | Tạo kịch bản, voice, subtitle, timeline và điều phối video |
| remotion-worker | Node.js/Remotion | Render MP4 từ timeline và tài nguyên media |
| frontend | React/Vite | Giao diện người dùng và quản trị |

### 3.4.2. Điều phối bằng Docker Compose

Docker Compose được sử dụng để khai báo các service, volume, biến môi trường, cổng truy cập và phụ thuộc khởi động. Các container giao tiếp trong cùng mạng nội bộ thông qua tên service, ví dụ API Service kết nối PostgreSQL qua host `postgres`, MongoDB qua host `mongodb` và Kafka qua host `kafka:29092`.

Các volume được sử dụng để lưu dữ liệu bền vững cho PostgreSQL, MongoDB và thư mục video render. Nhờ đó, dữ liệu không bị mất khi container được khởi động lại. Đối với các service xử lý nền, cấu hình `depends_on`, biến môi trường và log container hỗ trợ việc theo dõi trạng thái trong quá trình vận hành.

### 3.4.3. Cấu hình môi trường và bảo mật

Các thông tin nhạy cảm như database URL, Mongo URI, Kafka bootstrap servers, AI provider API key, TikTok client key, TikTok client secret và redirect URI được đặt trong tệp `.env`. Việc tách cấu hình khỏi mã nguồn giúp dễ thay đổi giữa môi trường development, staging và production.

Về bảo mật, hệ thống áp dụng xác thực JWT cho người dùng, phân quyền theo vai trò, kiểm soát phạm vi dữ liệu theo owner_id, chặn thao tác vượt quyền bằng HTTP 403, không cho user thường xem dữ liệu của tài khoản khác và ghi audit logs cho các thao tác quan trọng. Khi triển khai production, hệ thống nên đặt sau reverse proxy như Nginx, bật HTTPS/SSL và giới hạn truy cập trực tiếp vào database hoặc message broker từ bên ngoài.

### 3.4.4. Quy trình triển khai đề xuất

Quy trình triển khai đề xuất gồm các bước:

1. Chuẩn bị máy chủ có Docker, Docker Compose, Git và domain nếu triển khai public.
2. Tải mã nguồn lên server, tạo tệp `.env` theo cấu hình production.
3. Build và khởi chạy các container bằng `docker compose up --build -d`.
4. Kiểm tra log của từng service, đặc biệt API Service, Data Ingestion Engine, AI Media Engine, Remotion Worker và Kafka.
5. Kiểm tra endpoint `/health` của API Service.
6. Chạy thử luồng end-to-end: đăng nhập, tạo crawl job, chuẩn hóa dữ liệu, AI planning, tạo video, duyệt và đưa vào hàng đợi đăng bài.
7. Cấu hình reverse proxy, HTTPS và chính sách restart cho các service production.
8. Theo dõi audit logs, scheduler snapshot và trạng thái publishing queue sau khi đưa vào sử dụng.

## 3.5. Hỗ trợ và bảo trì

Sau khi triển khai, hệ thống cần được theo dõi thường xuyên để đảm bảo các luồng xử lý nền hoạt động ổn định. Các nhóm hỗ trợ chính bao gồm giám sát log, xử lý lỗi dữ liệu đầu vào, kiểm tra trạng thái queue, theo dõi lỗi AI provider, xử lý lỗi TikTok API và tối ưu thời gian render video.

**Bảng 3.7. Kế hoạch hỗ trợ và bảo trì**

| Nhóm hỗ trợ | Nội dung thực hiện | Mục tiêu |
|---|---|---|
| Theo dõi log | Kiểm tra log container, crawl_logs và audit_logs | Phát hiện lỗi sớm trong quá trình crawl, planning, render và publishing |
| Xử lý dữ liệu lỗi | Kiểm tra các bản ghi NEEDS_REVIEW, FAILED hoặc duplicate | Đảm bảo dữ liệu đầu vào cho AI planning có chất lượng |
| Bảo trì scheduler | Theo dõi interval, last_run, số item published/failed/skipped | Đảm bảo đăng bài tự động hoạt động đúng lịch |
| Kiểm tra AI provider | Theo dõi lỗi timeout, quota hoặc response không hợp lệ | Giảm lỗi khi sinh kế hoạch hoặc tạo nội dung |
| Kiểm tra TikTok API | Theo dõi token, scope, trạng thái publish và lỗi third-party | Đảm bảo luồng đăng bài thủ công/tự động ổn định |
| Sao lưu dữ liệu | Sao lưu PostgreSQL, MongoDB và thư mục video output | Tránh mất dữ liệu khi hệ thống gặp sự cố |
| Nâng cấp phiên bản | Cập nhật dependency, cải tiến giao diện, tối ưu worker | Tăng độ ổn định và khả năng mở rộng của hệ thống |

Trong các phiên bản tiếp theo, hệ thống có thể được cải tiến bằng cách bổ sung retry tự động cho các job thất bại, dashboard giám sát real-time cho admin, cảnh báo qua email/notification khi scheduler lỗi, tối ưu hàng đợi render video và tích hợp thêm nguồn dữ liệu đầu vào ngoài VNExpress.

## 3.6. Tổng kết chương 3

Chương 3 đã trình bày kết quả thực hiện của hệ thống SocialContent Hub sau quá trình phát triển theo Agile Scrum. Các nội dung chính bao gồm chuẩn bị phát hành, kiểm thử tích hợp, kiểm thử chấp nhận, hướng dẫn sử dụng, triển khai bằng Docker/Docker Compose và kế hoạch hỗ trợ sau triển khai.

Kết quả đạt được cho thấy hệ thống đã hoàn thiện luồng nghiệp vụ cốt lõi từ thu thập dữ liệu, chuẩn hóa, lập kế hoạch nội dung bằng AI, tạo video, kiểm duyệt, đăng bài thủ công/tự động đến quản trị hệ thống. Kiến trúc microservices kết hợp Kafka, PostgreSQL, MongoDB, AI Media Engine và Remotion Worker giúp hệ thống có khả năng mở rộng, dễ theo dõi và phù hợp với bài toán tự động hóa sản xuất nội dung video ngắn.

Mặc dù vẫn còn một số rủi ro liên quan đến dữ liệu đầu vào, độ ổn định của AI provider, giới hạn TikTok API và thời gian render video, hệ thống đã đáp ứng được mục tiêu chính của đồ án. Đây là cơ sở để tiếp tục hoàn thiện các chương tiếp theo, trong đó có thể phân tích sâu hơn về ưu điểm, hạn chế và định hướng phát triển trong tương lai.
"""


def iter_body_blocks(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def remove_existing_chapter3(document):
    deleting = False
    for block in list(iter_body_blocks(document)):
        if isinstance(block, Paragraph) and block.text.strip().startswith("CHƯƠNG 3"):
            deleting = True
        if deleting:
            element = block._element
            element.getparent().remove(element)


def clean_inline(text):
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def add_run_with_bold(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(clean_inline(part))


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(document, rows):
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.text = clean_inline(row[ci]) if ci < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8)
                    if ri == 0:
                        run.bold = True
            if ri == 0:
                shade_cell(cell)
    return table


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)
        if not is_separator:
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown(document, markdown):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2).replace("**", ""))
            p = document.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            if level == 1:
                run.font.size = Pt(16)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
            elif level == 2:
                run.font.size = Pt(14)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(6)
            else:
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith("- "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.add_run("- ")
            add_run_with_bold(p, stripped[2:])
            i += 1
            continue
        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.add_run(ordered.group(1) + ". ")
            add_run_with_bold(p, ordered.group(2))
            i += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            p = document.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        add_run_with_bold(p, stripped)
        i += 1


def main():
    doc = Document(INPUT_DOCX)
    remove_existing_chapter3(doc)
    doc.add_page_break()
    add_markdown(doc, CHAPTER3)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
