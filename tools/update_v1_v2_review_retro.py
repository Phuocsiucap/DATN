from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


INPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2_3_chi_tiet_2_5_2_2.docx")
OUTPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2_3_chi_tiet_v1_v2.docx")


def iter_body_blocks(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def find_paragraph(doc, predicate):
    for block in iter_body_blocks(doc):
        if isinstance(block, Paragraph) and predicate(block.text.strip()):
            return block
    raise RuntimeError("Cannot find paragraph")


def delete_range(doc, start_predicate, end_predicate):
    blocks = list(iter_body_blocks(doc))
    start_idx = None
    end_idx = None
    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and start_predicate(block.text.strip()):
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError("Cannot find replacement start")
    for i in range(start_idx + 1, len(blocks)):
        block = blocks[i]
        if isinstance(block, Paragraph) and end_predicate(block.text.strip()):
            end_idx = i
            break
    if end_idx is None:
        raise RuntimeError("Cannot find replacement end")
    for block in blocks[start_idx:end_idx]:
        block._element.getparent().remove(block._element)


def format_run(run, bold=False, size=11):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_spacing(paragraph, before=3, after=3, line=1.05):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_paragraph_before(anchor, text="", bold=False, size=11):
    p = anchor.insert_paragraph_before()
    r = p.add_run(text)
    format_run(r, bold=bold, size=size)
    set_paragraph_spacing(p)
    return p


def add_heading(anchor, text):
    p = add_paragraph_before(anchor, text, bold=True, size=12)
    return p


def add_label_paragraph(anchor, label, text):
    p = anchor.insert_paragraph_before()
    r1 = p.add_run(label)
    format_run(r1, bold=True, size=11)
    r2 = p.add_run(text)
    format_run(r2, size=11)
    set_paragraph_spacing(p)
    return p


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        r = p.add_run(part)
        format_run(r, bold=bold, size=size)
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def insert_table_before(doc, anchor, rows, widths=None, font_size=9, header_rows=1):
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = False
    set_table_fixed(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_text(cell, text, bold=r_idx < header_rows, size=font_size)
            if widths and c_idx < len(widths):
                set_cell_width(cell, widths[c_idx])
    anchor._p.addprevious(table._tbl)
    return table


def add_review_metadata(doc, anchor, version, sprint_name, goal, meeting_time):
    rows = [
        ["SPRINT REVIEW", ""],
        ["Tên Sprint", sprint_name],
        ["Mục tiêu", goal],
        ["Thời gian tổ chức", meeting_time],
        ["Địa điểm", "Google Meet"],
        ["Thời lượng cuộc họp dự kiến", "60 phút"],
        ["Thành phần tham gia", "Nhóm phát triển, Scrum Master, Product Owner và các bên liên quan"],
    ]
    table = insert_table_before(doc, anchor, rows, widths=[1.8, 5.3], font_size=10, header_rows=1)
    table.cell(0, 0).merge(table.cell(0, 1))
    set_cell_text(table.cell(0, 0), f"SPRINT REVIEW - {version}", bold=True, size=12)


def add_retro_metadata(doc, anchor, version, sprint_name, goal, meeting_time):
    rows = [
        ["SPRINT RETROSPECTIVE", ""],
        ["Tên Sprint", sprint_name],
        ["Mục tiêu", goal],
        ["Thời gian tổ chức", meeting_time],
        ["Địa điểm", "Google Meet"],
        ["Thời lượng cuộc họp dự kiến", "45 phút"],
        ["Thành phần tham gia", "Nhóm phát triển, Scrum Master, Product Owner"],
    ]
    table = insert_table_before(doc, anchor, rows, widths=[1.8, 5.3], font_size=10, header_rows=1)
    table.cell(0, 0).merge(table.cell(0, 1))
    set_cell_text(table.cell(0, 0), f"SPRINT RETROSPECTIVE - {version}", bold=True, size=12)


def insert_v1_sections(doc, anchor):
    add_heading(anchor, "2.5.1.3. Phát triển tính năng trên câu chuyện")
    add_paragraph_before(anchor, "Ở phiên bản V1.0, nhóm phát triển tập trung hoàn thiện nền tảng nghiệp vụ đầu tiên của hệ thống SocialContent Hub. Phạm vi triển khai bao gồm xác thực, phân quyền, đăng xuất, tạo và theo dõi crawl job, thu thập dữ liệu VNExpress, chuẩn hóa dữ liệu, nhóm nội dung theo truyện/tập, đánh giá chất lượng, phát hiện trùng lặp, tạo content project và lập kế hoạch nội dung bằng AI.")

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-01, US-02 và US-03: ", "nhóm chức năng xác thực, phân quyền và quản lý phiên đăng nhập.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "xây dựng lớp bảo vệ truy cập để người dùng chỉ sử dụng được hệ thống sau khi đăng nhập và chỉ thao tác trong phạm vi quyền được cấp.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "giao diện đăng nhập được thiết kế cho người sáng tạo nội dung và quản trị viên; backend kiểm tra thông tin đăng nhập, mã hóa mật khẩu, cấp JWT token, xác định vai trò CREATOR hoặc SYSTEM_ADMIN và gắn thông tin người dùng hiện tại vào request. Các API chính như crawl job, content project, social profile, media workflow và quản trị người dùng đều kiểm tra quyền trước khi trả dữ liệu. Chức năng đăng xuất xóa phiên ở phía client và đưa người dùng về trạng thái chưa xác thực.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Đăng nhập bằng tài khoản hợp lệ", "Người dùng nhận được phiên đăng nhập và truy cập được màn hình chính theo đúng vai trò.", "Đạt"],
            ["Đăng nhập bằng thông tin sai", "Hệ thống từ chối đăng nhập, không cấp token và hiển thị thông báo lỗi rõ ràng.", "Đạt"],
            ["Đăng xuất khỏi hệ thống", "Phiên làm việc phía client được xóa, người dùng không thể gọi API yêu cầu xác thực bằng phiên cũ.", "Đạt"],
            ["Truy cập API không đủ quyền", "API trả lỗi phân quyền, dữ liệu của người dùng khác không bị lộ.", "Đạt"],
            ["Quản trị cập nhật vai trò", "Vai trò mới được lưu và áp dụng ở các màn hình, API liên quan.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-06, US-07 và US-08: ", "nhóm chức năng tạo job thu thập dữ liệu, theo dõi tiến độ crawl và lấy dữ liệu VNExpress.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "cho phép người dùng tạo private crawl job theo nguồn và từ khóa, theo dõi tiến độ xử lý và thu thập bài viết VNExpress làm dữ liệu đầu vào cho AI planning.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "người dùng nhập nguồn VNExpress, từ khóa hoặc URL, cấu hình số lượng bài cần lấy và tạo crawl job. API Service lưu job thuộc sở hữu người dùng, phát sự kiện sang Kafka để crawl-orchestrator và crawler-service xử lý. Trong quá trình chạy, hệ thống cập nhật trạng thái PENDING, RUNNING, SUCCEEDED, FAILED hoặc PARTIAL_SUCCESS, ghi crawl_logs, số lượng URL phát hiện, số bài crawl thành công, lỗi nguồn và tiến độ phần trăm.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Tạo crawl job VNExpress hợp lệ", "Job được tạo, gắn user_id, chuyển trạng thái xử lý và sinh task crawl tương ứng.", "Đạt"],
            ["Theo dõi tiến độ crawl", "Giao diện hiển thị trạng thái, phần trăm tiến độ, số bài đã xử lý và log từng bước.", "Đạt"],
            ["Nguồn hoặc từ khóa không hợp lệ", "Hệ thống từ chối tạo job hoặc chuyển job sang FAILED kèm lý do cụ thể.", "Đạt"],
            ["Crawler lấy bài viết thành công", "Bài viết có tiêu đề, mô tả, nội dung, URL nguồn, metadata và hình ảnh được lưu vào kho dữ liệu thô.", "Đạt"],
            ["Người dùng xem job không thuộc sở hữu", "Hệ thống không trả dữ liệu job của người dùng khác.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-09, US-10 và US-11: ", "nhóm chức năng chuẩn hóa dữ liệu, nhóm nội dung và kiểm soát chất lượng.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "chuyển dữ liệu crawl thô thành dữ liệu chuẩn có thể dùng cho AI, đồng thời phát hiện nội dung hoặc hình ảnh trùng lặp để giảm nhiễu khi lập kế hoạch.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "story-processing-service làm sạch HTML, chuẩn hóa tiêu đề, mô tả, nội dung chính, ảnh đại diện và metadata nguồn. Hệ thống nhóm các bài liên quan theo truyện, chủ đề hoặc tập nội dung, sắp xếp thứ tự để chuẩn bị cho AI planning. Quality Score được tính dựa trên độ đầy đủ của nội dung, độ dài văn bản, ảnh, nguồn tham chiếu và trạng thái xử lý. Cơ chế deduplication sử dụng fingerprint để phát hiện bài hoặc ảnh trùng, từ đó đánh dấu dữ liệu canonical.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Chuẩn hóa bài crawl thô", "Dữ liệu sau xử lý có cấu trúc thống nhất, loại bỏ HTML rác và trường metadata thừa.", "Đạt"],
            ["Kiểm tra trường bắt buộc", "Bản ghi thiếu tiêu đề, nội dung hoặc URL nguồn được đánh dấu lỗi chất lượng.", "Đạt"],
            ["Nhóm nội dung theo truyện/tập", "Các bài cùng mạch nội dung được gom nhóm và sắp xếp thứ tự hợp lý.", "Đạt"],
            ["Tính Quality Score", "Hệ thống có điểm đánh giá chất lượng để ưu tiên dữ liệu tốt cho AI.", "Đạt"],
            ["Deduplication", "Nội dung hoặc hình ảnh trùng được phát hiện bằng fingerprint và không tạo nhiều bản canonical.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-12, US-13, US-14 và US-15: ", "nhóm chức năng tạo content project, AI đề xuất chủ đề, chia nội dung nhiều phần và duyệt kế hoạch.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "chuyển dữ liệu đã crawl và chuẩn hóa thành dự án nội dung có kế hoạch triển khai video rõ ràng, đủ điều kiện chuyển sang pipeline sản xuất video ở V2.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "người dùng chọn dữ liệu chuẩn để tạo content project. Hệ thống kiểm tra quyền sở hữu dữ liệu, liên kết project với nguồn crawl, nhóm nội dung và metadata. AI Planning Worker nhận nội dung đầu vào, mục tiêu kênh, kiểu video mong muốn và yêu cầu thời lượng để đề xuất chủ đề, góc triển khai, định dạng video. Với nội dung dài, AI chia thành nhiều phần, giữ mạch truyện và đưa ra kế hoạch theo từng tập. Người dùng có thể duyệt, từ chối, chỉnh sửa hoặc yêu cầu tạo lại kế hoạch trước khi chuyển sang tạo kịch bản video.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Tạo content project từ dữ liệu đã chuẩn hóa", "Project được tạo, gắn đúng dữ liệu nguồn, người sở hữu và trạng thái ban đầu.", "Đạt"],
            ["AI đề xuất chủ đề và góc triển khai", "Kết quả có chủ đề, insight, kiểu video, hook và lý do lựa chọn.", "Đạt"],
            ["AI chia nội dung thành nhiều phần", "Kế hoạch nhiều tập có thứ tự, mục tiêu từng tập và mạch nội dung liên tục.", "Đạt"],
            ["Duyệt kế hoạch", "Kế hoạch chuyển sang APPROVED và trở thành đầu vào hợp lệ cho V2.", "Đạt"],
            ["AI xử lý lỗi hoặc kết quả không đạt", "Hệ thống hiển thị lỗi, giữ trạng thái an toàn và cho phép tạo lại.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_heading(anchor, "2.5.1.4. Phát hành phiên bản sản phẩm phần mềm")
    add_paragraph_before(anchor, "Sau khi hoàn thành Sprint 1 và Sprint 2, phiên bản V1.0 được đóng gói để demo luồng nền tảng từ đăng nhập, thu thập dữ liệu, chuẩn hóa dữ liệu đến lập kế hoạch nội dung bằng AI. Đây là bản phát hành nền tảng, chưa tập trung vào tạo video hoàn chỉnh nhưng phải bảo đảm dữ liệu đầu vào và kế hoạch AI đủ tin cậy cho V2.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Hạng mục", "Nội dung"],
            ["Phiên bản", "V1.0 - Release nền tảng dữ liệu và AI Planning"],
            ["Ngày phát hành dự kiến", "23/08/2026"],
            ["Sprint bao phủ", "Sprint 1 và Sprint 2"],
            ["Phạm vi user story", "US-01, US-02, US-03, US-06, US-07, US-08, US-09, US-10, US-11, US-12, US-13, US-14, US-15"],
            ["Mục tiêu phát hành", "Demo được quy trình lấy dữ liệu VNExpress, chuẩn hóa, deduplicate, tạo content project và tạo kế hoạch nội dung bằng AI."],
            ["Tiêu chí chấp nhận", "Người dùng đăng nhập được, dữ liệu được lọc theo quyền, crawl job chạy có log, dữ liệu chuẩn đủ chất lượng, AI tạo kế hoạch có thể duyệt để chuyển sang V2."],
        ],
        widths=[2.0, 5.1],
        font_size=9,
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Tính năng hoàn thành", "Ý nghĩa nghiệp vụ", "Trạng thái"],
            ["Xác thực, phân quyền, đăng xuất", "Bảo vệ dữ liệu và phân tách quyền giữa Creator, Admin và người vận hành.", "Hoàn thành"],
            ["Crawl job và theo dõi tiến độ", "Cho phép người dùng chủ động tạo job thu thập dữ liệu và theo dõi quá trình xử lý.", "Hoàn thành"],
            ["Crawler VNExpress", "Cung cấp nguồn nội dung tiếng Việt có metadata và hình ảnh cho hệ thống.", "Hoàn thành"],
            ["Chuẩn hóa, grouping, Quality Score và deduplication", "Làm sạch dữ liệu trước khi đưa vào AI, giảm trùng lặp và lỗi nguồn.", "Hoàn thành"],
            ["Content project và AI Planning", "Biến dữ liệu đã crawl thành kế hoạch nội dung có chủ đề, góc triển khai và nhiều phần.", "Hoàn thành"],
            ["Duyệt kế hoạch", "Tạo điểm kiểm soát trước khi chuyển sang pipeline sản xuất video.", "Hoàn thành"],
        ],
        widths=[2.1, 4.1, 0.9],
        font_size=9,
    )

    add_heading(anchor, "2.5.1.5. Đánh giá phiên bản - Sprint Review")
    add_paragraph_before(anchor, "Buổi Sprint Review của V1.0 được tổ chức sau khi hoàn thành hai Sprint đầu tiên nhằm trình bày các chức năng nền tảng cho Product Owner và các bên liên quan. Nội dung đánh giá tập trung vào việc hệ thống đã tạo được đường ống dữ liệu đủ sạch và kế hoạch AI đủ rõ ràng để phục vụ sản xuất video ở V2.")
    add_paragraph_before(anchor, "Bảng 2.9. Sprint Review cho phiên bản V1.0")
    add_review_metadata(
        doc,
        anchor,
        "V1.0",
        "Sprint 1 và Sprint 2",
        "Hoàn thành nền tảng xác thực, phân quyền, thu thập dữ liệu, chuẩn hóa dữ liệu, nhóm nội dung và lập kế hoạch nội dung bằng AI cho Release V1.",
        "Sau mốc Release V1 ngày 23/08/2026",
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Team", "Role", "Sprint Status", "Task Demo", "Phản hồi của khách hàng/PO", "Điều chỉnh", "Trạng thái"],
            ["DEV Team", "Frontend Dev", "Hoàn thành US-01, US-02, US-03, US-06, US-07, US-12, US-15", "Demo màn hình đăng nhập, điều hướng theo quyền, danh sách crawl job, chi tiết log job, tạo content project và duyệt kế hoạch.", "Giao diện đủ luồng nghiệp vụ chính, cần làm rõ trạng thái job và thông báo lỗi để người dùng không nhầm giữa đang chạy và thất bại.", "Bổ sung nhãn trạng thái, thông báo lỗi dễ hiểu, bộ lọc job theo trạng thái và hiển thị owner của dữ liệu.", "Đạt"],
            ["DEV Team", "Backend Dev", "Hoàn thành US-01 đến US-03, US-06 đến US-11", "Demo API đăng nhập, phân quyền middleware, tạo crawl job, phát Kafka event, crawl VNExpress, chuẩn hóa dữ liệu, lưu PostgreSQL/MongoDB và deduplication.", "Backend xử lý đúng luồng nhưng cần log chi tiết hơn ở các job lỗi để hỗ trợ kiểm thử và vận hành.", "Bổ sung crawl_logs, failure_reason, tiến độ theo từng bước và kiểm tra dữ liệu theo user_id.", "Đạt"],
            ["DEV Team", "AI/Data Dev", "Hoàn thành US-09 đến US-15", "Demo dữ liệu canonical, Quality Score, nhóm nội dung, tạo project, AI đề xuất chủ đề, chia nội dung nhiều phần và duyệt kế hoạch.", "Kế hoạch AI có cấu trúc tốt nhưng một số prompt cần ổn định hơn khi dữ liệu nguồn dài hoặc thiếu metadata.", "Tinh chỉnh prompt, bổ sung validation kết quả AI và cơ chế tạo lại kế hoạch khi kết quả chưa đạt.", "Đạt"],
            ["DEV Team", "Tester", "Hoàn thành kiểm thử chức năng V1.0", "Demo test case đăng nhập, phân quyền, tạo job, theo dõi crawl, chuẩn hóa dữ liệu, deduplication, tạo content project và duyệt kế hoạch.", "Các chức năng chính đạt yêu cầu nghiệm thu, cần tăng kiểm thử dữ liệu lỗi và quyền truy cập chéo giữa người dùng.", "Bổ sung test case dữ liệu thiếu, job thất bại, truy cập trái quyền và AI trả về kết quả rỗng.", "Đạt"],
            ["Product Owner", "PO/Scrum Master", "Hoàn thành nghiệm thu Release V1", "Tổng hợp kết quả demo, xác nhận phạm vi Release V1 và thống nhất đầu vào cho V2.", "Release V1 đạt mục tiêu nền tảng, có thể chuyển sang phát triển pipeline tạo video.", "Ưu tiên Sprint 3 cho kịch bản, scene, TTS, subtitle sync và render MP4.", "Đạt"],
        ],
        widths=[0.8, 0.8, 1.0, 1.7, 1.6, 1.5, 0.55],
        font_size=7.8,
    )

    add_heading(anchor, "2.5.1.6. Hồi tưởng phiên bản - Sprint Retrospective")
    add_paragraph_before(anchor, "Ngay sau Sprint Review, nhóm phát triển thực hiện Sprint Retrospective để đánh giá cách phối hợp trong hai Sprint đầu. Trọng tâm hồi tưởng là chất lượng dữ liệu, tính rõ ràng của trạng thái job, sự phối hợp giữa backend, crawler, AI worker và khả năng kiểm thử độc lập từng nghiệp vụ.")
    add_paragraph_before(anchor, "Bảng 2.10. Sprint Retrospective cho phiên bản V1.0")
    add_retro_metadata(
        doc,
        anchor,
        "V1.0",
        "Sprint 1 và Sprint 2",
        "Hoàn thành nền tảng dữ liệu và chức năng lập kế hoạch nội dung bằng AI cho Release V1.",
        "Sau mốc Release V1 ngày 23/08/2026",
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Nhóm nội dung", "Đánh giá"],
            ["Nội dung đã làm tốt", "Nhóm đã hoàn thành đúng trọng tâm V1: xác thực, phân quyền, crawl job, VNExpress crawler, chuẩn hóa, grouping, Quality Score, deduplication và AI Planning. Các service backend được tách theo nhiệm vụ, dữ liệu người dùng được gắn owner rõ ràng, luồng từ crawl đến content project đủ điều kiện demo. Việc dùng Kafka event giúp các bước crawl và xử lý dữ liệu có thể mở rộng sang các worker khác."],
            ["Nội dung cần điều chỉnh", "Một số màn hình cần hiển thị trạng thái rõ hơn để phân biệt job đang chờ, đang chạy, thành công một phần và thất bại. Dữ liệu crawl từ nguồn thực tế có thể thiếu ảnh hoặc thiếu nội dung, do đó phần validation cần mạnh hơn. Kết quả AI Planning phụ thuộc nhiều vào chất lượng prompt và dữ liệu đầu vào, cần có cơ chế retry và đánh giá kết quả trước khi cho duyệt."],
            ["Hành động cải tiến", "Trong V2, nhóm thống nhất bổ sung trạng thái workflow chi tiết hơn cho media pipeline, tăng logging ở các job bất đồng bộ, chuẩn hóa failure_reason ở API, kiểm thử quyền truy cập chéo, và xây dựng validation cho từng bước tạo video. Các test case cần chia theo luồng thành công, luồng dữ liệu thiếu, luồng lỗi worker và luồng người dùng không đủ quyền."],
            ["Bài học rút ra", "Cần thiết kế dữ liệu trung gian đủ chặt ngay từ đầu vì AI Planning và sản xuất video phụ thuộc trực tiếp vào dữ liệu sau crawl. Các bước bất đồng bộ phải có log và trạng thái rõ ràng để người dùng không bị mất niềm tin khi hệ thống xử lý nền. Mỗi user story nên có tiêu chí nghiệm thu gắn với dữ liệu thật thay vì chỉ kiểm thử màn hình."],
            ["Rủi ro chuyển sang V2", "Nếu timeline, voice, subtitle và render không dùng cùng một version dữ liệu, video có thể bị lệch tiếng hoặc sai nội dung. Vì vậy V2 cần kiểm soát revision của script, audio, subtitle và render output, đồng thời khóa thao tác khi job đang chạy."],
        ],
        widths=[1.7, 5.4],
        font_size=9,
    )


def insert_v2_sections(doc, anchor):
    add_heading(anchor, "2.5.2.3. Phát triển tính năng trên câu chuyện")
    add_paragraph_before(anchor, "Ở phiên bản V2.0, phạm vi phát triển chuyển từ nền tảng dữ liệu sang pipeline sản xuất và xuất bản video hoàn chỉnh. Các chức năng được triển khai phải bảo đảm kế hoạch nội dung đã duyệt ở V1 có thể biến thành kịch bản, scene, voice, phụ đề, file MP4, sau đó được QA, duyệt, lên lịch và đăng lên TikTok theo cách thủ công hoặc tự động.")

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-16, US-17 và US-18: ", "nhóm chức năng biên soạn kịch bản, chỉnh sửa AI và tạo giọng đọc.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "chuyển kế hoạch nội dung đã duyệt thành kịch bản sản xuất video, chia scene, cho phép chỉnh bằng AI và tạo file audio giọng đọc làm mốc đồng bộ.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "AI Media Engine nhận project plan đã APPROVED, sinh kịch bản video, hook, body, CTA và danh sách scene. Mỗi scene có lời thoại, visual prompt, mô tả cảm xúc và thời lượng dự kiến. Người dùng có thể chỉnh một scene hoặc toàn bộ kịch bản bằng prompt AI, hệ thống lưu revision để truy vết. Khi kịch bản ổn định, TTS Engine tạo file MP3, lưu duration chính xác và gắn audio với script revision hiện hành.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Tạo kịch bản từ kế hoạch đã duyệt", "Kịch bản có hook, nội dung chính, CTA và liên kết đúng content project.", "Đạt"],
            ["Chia kịch bản thành scene", "Mỗi scene có thứ tự, lời thoại, visual prompt, thời lượng dự kiến và trạng thái hợp lệ.", "Đạt"],
            ["Chỉnh scene bằng AI", "AI cập nhật riêng scene được chọn, giữ ngữ cảnh trước/sau và lưu revision.", "Đạt"],
            ["Tạo voice TTS", "File MP3 được tạo, lưu đường dẫn, duration và gắn với script revision hiện tại.", "Đạt"],
            ["Script thay đổi sau khi tạo voice", "Hệ thống cảnh báo audio cũ không khớp và yêu cầu tạo lại trước khi render.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-19, US-20 và US-21: ", "nhóm chức năng đồng bộ phụ đề, render MP4 và duy trì series context.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "đồng bộ 1:1 giữa âm thanh, phụ đề và hình ảnh, xuất file MP4 hoàn chỉnh và giữ mạch nội dung xuyên suốt giữa các tập video.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "sau khi audio hoàn thành, hệ thống tạo subtitle segment, chạy thuật toán fit_video_clips_to_text để căn timestamp theo duration audio và phân bổ thời lượng cho từng visual clip. Render Worker sử dụng timeline cuối để ghép voice, subtitle, hình ảnh và hiệu ứng thành MP4. Với video series, MongoDB lưu series context gồm nhân vật, bối cảnh, sự kiện chính, tone giọng và cliffhanger để tập sau có thể kế thừa.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Tạo phụ đề từ lời thoại", "Subtitle được tách thành các segment ngắn, có start_time và end_time.", "Đạt"],
            ["Đồng bộ 1:1 bằng fit_video_clips_to_text", "Thời lượng hình ảnh, phụ đề và audio khớp nhau, không bị trôi tiếng.", "Đạt"],
            ["Render MP4", "Video dọc hoàn chỉnh có audio, subtitle, hình ảnh, thumbnail và duration.", "Đạt"],
            ["Render thất bại", "Job chuyển FAILED, có failure_reason và có thể retry từ timeline hiện tại.", "Đạt"],
            ["Series context", "Tập tiếp theo kế thừa đúng nhân vật, bối cảnh, tone và mạch truyện từ tập trước.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-22 và US-23: ", "nhóm chức năng xem trước, kiểm tra chất lượng và duyệt nội dung trước khi đăng.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "thiết lập cổng QA để video chỉ được đưa vào hàng đợi xuất bản khi đã xem trước và được duyệt.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "màn hình QA Preview tải file MP4 mới nhất qua media proxy, cho phép phát, tua, kiểm tra âm thanh, phụ đề, hình ảnh và version render. Người dùng có thể ghi nhận lỗi theo timestamp, yêu cầu render lại hoặc duyệt video. Khi duyệt, hệ thống lưu người duyệt, thời điểm duyệt, version được duyệt và mở khóa chức năng lên lịch hoặc đăng ngay.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Xem preview MP4", "Trình phát hiển thị đúng video mới nhất, có duration và thumbnail.", "Đạt"],
            ["Ghi nhận lỗi QA", "Issue được lưu theo video, scene, timestamp và mức độ nghiêm trọng.", "Đạt"],
            ["Duyệt video đạt yêu cầu", "Video chuyển APPROVED và được phép lên lịch hoặc đăng ngay.", "Đạt"],
            ["Từ chối video", "Video chuyển NEED_FIX/REJECTED, chức năng đăng bị khóa cho đến khi có bản sửa.", "Đạt"],
            ["Audit phê duyệt", "Hệ thống ghi người thao tác, trạng thái cũ/mới và version video.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-04 và US-05: ", "nhóm chức năng kết nối TikTok bằng QR và cấu hình chiến lược kênh.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "liên kết tài khoản TikTok của người dùng với hệ thống và lưu chiến lược đăng nội dung để scheduler có căn cứ vận hành.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "người dùng tạo phiên QR TikTok, hệ thống polling trạng thái xác thực, nhận token, scope và thông tin kênh sau khi quét thành công. Social profile được lưu theo user_id để bảo đảm phân quyền. Người dùng cấu hình SocialProfileStrategy gồm chủ đề ưu tiên, chủ đề cần tránh, tone nội dung, tần suất đăng, khung giờ và auto_publish_enabled. Strategy này được dùng khi gợi ý lịch đăng và khi auto scheduler quyết định có được tự động đăng hay không.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Tạo phiên QR TikTok", "Hệ thống trả session_id, QR, expires_at và trạng thái pending.", "Đạt"],
            ["Quét QR thành công", "Social profile được lưu ACTIVE cùng token, scope và thông tin kênh.", "Đạt"],
            ["QR hết hạn hoặc token thiếu quyền", "Hệ thống thông báo cần tạo lại QR hoặc cấp quyền đăng video.", "Đạt"],
            ["Tạo chiến lược kênh", "Strategy được lưu theo profile, gồm chủ đề, tone, tần suất, khung giờ và auto_publish_enabled.", "Đạt"],
            ["Truy cập strategy không thuộc sở hữu", "API từ chối, không lộ dữ liệu profile của người dùng khác.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_label_paragraph(anchor, "Phát triển tính năng dựa trên US-24, US-25, US-26, US-27 và US-28: ", "nhóm chức năng lên lịch, theo dõi hàng đợi, đăng thủ công, tự động đăng và thống kê hiệu suất.")
    add_label_paragraph(anchor, "Nhiệm vụ: ", "hoàn thiện vòng đời xuất bản video từ APPROVED đến QUEUED, PUBLISHING, PUBLISHED hoặc FAILED, đồng thời thu thập metrics sau khi đăng.")
    add_label_paragraph(anchor, "Mô tả triển khai: ", "video đã APPROVED được đưa vào publishing queue theo profile và scheduled_at. Người dùng có thể xem queue, lọc theo trạng thái, cập nhật lịch hoặc bấm Đăng ngay để gọi TikTok API tức thời. Background Scheduler chạy định kỳ, kiểm tra ENABLE_SCHEDULER, auto_publish_enabled, token, video version và thời điểm scheduled_at để tự động đăng bài đến hạn. Sau khi đăng thành công, hệ thống lưu SocialPost, platform_post_id, link bài đăng và thu thập views, likes, comments, shares để hiển thị dashboard metrics.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Nội dung kiểm thử", "Kết quả mong đợi", "Trạng thái"],
            ["Lên lịch video đã duyệt", "Queue item được tạo với trạng thái QUEUED, profile_id, scheduled_at và caption.", "Đạt"],
            ["Theo dõi hàng đợi", "Giao diện hiển thị QUEUED, PUBLISHING, PUBLISHED, FAILED, failure_reason và retry_count.", "Đạt"],
            ["Đăng ngay thủ công", "Video được gọi TikTok API tức thời, item chuyển PUBLISHED hoặc FAILED rõ lý do.", "Đạt"],
            ["Auto Scheduler đăng đúng giờ", "Item đến hạn được khóa, đăng tự động và tránh duplicate publish.", "Đạt"],
            ["Thu thập post metrics", "Dashboard hiển thị views, likes, comments, shares và thời điểm cập nhật.", "Đạt"],
        ],
        widths=[2.3, 3.9, 0.9],
        font_size=9,
    )

    add_heading(anchor, "2.5.2.4. Phát hành phiên bản sản phẩm phần mềm")
    add_paragraph_before(anchor, "Phiên bản V2.0 được phát hành sau Sprint 3 và Sprint 4 với mục tiêu chứng minh hệ thống không chỉ lập kế hoạch nội dung mà còn sản xuất được video MP4 và đưa video vào quy trình xuất bản thực tế trên TikTok. Đây là bản phát hành quan trọng nhất về mặt nghiệp vụ vì nối liền AI Planning, Media Workflow, QA Approval và Publishing Queue.")
    insert_table_before(
        doc,
        anchor,
        [
            ["Hạng mục", "Nội dung"],
            ["Phiên bản", "V2.0 - Release sản xuất video và xuất bản TikTok"],
            ["Ngày phát hành dự kiến", "13/09/2026"],
            ["Sprint bao phủ", "Sprint 3 và Sprint 4"],
            ["Phạm vi user story", "US-16, US-17, US-18, US-19, US-20, US-21, US-22, US-23, US-04, US-05, US-24, US-25, US-26, US-27, US-28"],
            ["Mục tiêu phát hành", "Demo được quy trình tạo kịch bản, scene, voice, subtitle, render MP4, QA, duyệt, kết nối TikTok, lên lịch, đăng ngay, tự động đăng và xem metrics."],
            ["Tiêu chí chấp nhận", "Video MP4 đồng bộ audio/phụ đề/hình ảnh, TikTok profile kết nối được, queue hoạt động minh bạch, đăng thủ công và tự động có trạng thái rõ ràng, metrics hiển thị được sau khi đăng."],
        ],
        widths=[2.0, 5.1],
        font_size=9,
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Tính năng hoàn thành", "Ý nghĩa nghiệp vụ", "Trạng thái"],
            ["Script, scene, AI edit và TTS", "Biến kế hoạch AI thành dữ liệu sản xuất video có lời thoại và audio.", "Hoàn thành"],
            ["Subtitle sync và render MP4", "Tạo video hoàn chỉnh không bị lệch tiếng, sẵn sàng preview.", "Hoàn thành"],
            ["Series context", "Duy trì mạch truyện và nhận diện nội dung giữa nhiều tập.", "Hoàn thành"],
            ["QA Preview và Approval", "Kiểm soát chất lượng trước khi đưa video vào hàng đợi đăng.", "Hoàn thành"],
            ["TikTok QR và Channel Strategy", "Liên kết kênh thật và thiết lập chính sách đăng nội dung.", "Hoàn thành"],
            ["Publishing Queue, Manual Publish, Auto Publish và Metrics", "Hoàn thiện vòng đời xuất bản và đánh giá hiệu quả bài đăng.", "Hoàn thành"],
        ],
        widths=[2.1, 4.1, 0.9],
        font_size=9,
    )

    add_heading(anchor, "2.5.2.5. Đánh giá phiên bản - Sprint Review")
    add_paragraph_before(anchor, "Sprint Review của V2.0 tập trung demo luồng nghiệp vụ đầy đủ từ kế hoạch nội dung đã duyệt đến video MP4 và xuất bản TikTok. Các bên liên quan đánh giá không chỉ giao diện mà cả tính đúng của trạng thái workflow, độ đồng bộ audio/phụ đề, khả năng retry khi lỗi và sự phân tách giữa đăng thủ công với tự động đăng theo lịch.")
    add_paragraph_before(anchor, "Bảng 2.14. Sprint Review cho phiên bản V2.0")
    add_review_metadata(
        doc,
        anchor,
        "V2.0",
        "Sprint 3 và Sprint 4",
        "Hoàn thành pipeline tạo video MP4, QA, duyệt nội dung, kết nối TikTok, cấu hình chiến lược kênh, hàng đợi đăng bài, đăng thủ công, tự động đăng theo lịch và theo dõi metrics.",
        "Sau mốc Release V2 ngày 13/09/2026",
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Team", "Role", "Sprint Status", "Task Demo", "Phản hồi của khách hàng/PO", "Điều chỉnh", "Trạng thái"],
            ["DEV Team", "Frontend Dev", "Hoàn thành US-16, US-17, US-22, US-23, US-24, US-25, US-26, US-27", "Demo màn hình media workflow, danh sách scene, chỉnh AI, preview MP4, duyệt video, queue, đăng ngay và dashboard metrics.", "Luồng người dùng đã đủ nhưng cần nhấn mạnh version mới nhất của script/audio/render để tránh duyệt nhầm bản cũ.", "Bổ sung nhãn latest version, trạng thái NEED_FIX, APPROVED, QUEUED và cảnh báo khi dữ liệu đã thay đổi.", "Đạt"],
            ["DEV Team", "Backend Dev", "Hoàn thành US-04, US-05, US-16 đến US-28", "Demo API social profile TikTok QR, strategy, media workflow state, publishing queue, manual publish, auto scheduler, retry và metrics.", "API đáp ứng nghiệp vụ chính, cần thống nhất failure_reason để người dùng biết lỗi do token, file, TikTok hay scheduler.", "Chuẩn hóa mã lỗi, bổ sung audit log cho phê duyệt và publish, khóa item khi PUBLISHING để tránh đăng trùng.", "Đạt"],
            ["DEV Team", "AI Media Dev", "Hoàn thành US-16 đến US-21", "Demo tạo script, chia scene, tạo voice TTS, sinh subtitle, chạy fit_video_clips_to_text, render MP4 và lưu series context.", "Chất lượng video đạt yêu cầu demo, điểm quan trọng nhất là phụ đề và giọng đọc không bị lệch.", "Giữ duration audio làm nguồn chuẩn, bắt buộc tạo lại subtitle/render khi script hoặc audio đổi revision.", "Đạt"],
            ["DEV Team", "Integration/DevOps", "Hoàn thành US-20, US-24, US-27, US-28", "Demo worker render, storage video, scheduler quét item đến hạn, cấu hình ENABLE_SCHEDULER và trạng thái auto publish.", "Cần có khả năng quan sát scheduler vì lỗi tự động đăng khó phát hiện nếu chỉ nhìn giao diện người dùng.", "Bổ sung log chu kỳ scheduler, số item quét, số item thành công/thất bại và trạng thái paused khi tắt scheduler.", "Đạt"],
            ["DEV Team", "Tester", "Hoàn thành kiểm thử chức năng V2.0", "Demo test case timeline, TTS, subtitle sync, render, QA approval, TikTok QR, queue, manual publish, auto publish và metrics.", "Cần kiểm thử sâu hơn ở luồng lỗi: token hết hạn, TikTok trả lỗi, video chưa duyệt, item bị scheduler xử lý đồng thời.", "Tăng test case negative, kiểm tra chống duplicate publish và retry sau FAILED.", "Đạt"],
            ["Product Owner", "PO/Scrum Master", "Hoàn thành nghiệm thu Release V2", "Tổng hợp kết quả demo, xác nhận pipeline video và publishing đủ điều kiện chuyển sang Production hardening.", "Release V2 đạt mục tiêu nghiệp vụ, có thể triển khai bước Production gồm admin, audit, system settings và môi trường Docker.", "Ưu tiên Sprint 5 cho cấu hình vận hành, phân quyền admin, audit logs và kiểm thử triển khai.", "Đạt"],
        ],
        widths=[0.8, 0.8, 1.0, 1.7, 1.6, 1.5, 0.55],
        font_size=7.8,
    )

    add_heading(anchor, "2.5.2.6. Hồi tưởng phiên bản - Sprint Retrospective")
    add_paragraph_before(anchor, "Sprint Retrospective của V2.0 được thực hiện sau khi nhóm hoàn thành pipeline video và publishing. Nội dung hồi tưởng tập trung vào các điểm dễ phát sinh lỗi trong hệ thống bất đồng bộ: version dữ liệu, đồng bộ audio/subtitle, render job, token TikTok, queue lock, retry và scheduler.")
    add_paragraph_before(anchor, "Bảng 2.15. Sprint Retrospective cho phiên bản V2.0")
    add_retro_metadata(
        doc,
        anchor,
        "V2.0",
        "Sprint 3 và Sprint 4",
        "Hoàn thành pipeline tạo video, kiểm duyệt chất lượng và xuất bản TikTok cho Release V2.",
        "Sau mốc Release V2 ngày 13/09/2026",
    )
    insert_table_before(
        doc,
        anchor,
        [
            ["Nhóm nội dung", "Đánh giá"],
            ["Nội dung đã làm tốt", "Nhóm đã hoàn thành được luồng end-to-end từ kế hoạch V1 sang script, scene, voice TTS, subtitle sync, render MP4, QA preview, duyệt nội dung, kết nối TikTok, publishing queue, manual publish, auto publish và post metrics. Việc tách rõ US-27 đăng thủ công và US-28 tự động đăng giúp kiểm thử độc lập hơn. Thuật toán fit_video_clips_to_text giải quyết đúng vấn đề trôi tiếng/lệch phụ đề."],
            ["Nội dung cần điều chỉnh", "Pipeline V2 có nhiều bước bất đồng bộ nên cần kiểm soát version dữ liệu chặt hơn. Khi script thay đổi, audio, subtitle và render cũ phải được đánh dấu không còn mới nhất. Queue publish cần cơ chế khóa khi PUBLISHING để tránh scheduler và thao tác đăng ngay cùng xử lý một item. Các lỗi TikTok cần hiển thị dễ hiểu hơn cho người dùng."],
            ["Hành động cải tiến", "Trong giai đoạn Production, nhóm cần hoàn thiện audit logs cho duyệt và publish, system settings cho scheduler, dashboard quản trị, cấu hình retry, kiểm tra token TikTok định kỳ và log vận hành. Cần bổ sung test case race condition, duplicate publish, token hết hạn, file MP4 lỗi, scheduler paused và metrics pending."],
            ["Bài học rút ra", "Đối với hệ thống tạo video, audio duration phải là nguồn chuẩn cho timeline; mọi thay đổi script đều kéo theo phụ thuộc voice, subtitle và render. Đối với publishing, trạng thái queue phải được thiết kế như một state machine rõ ràng, vì chỉ một trạng thái sai có thể dẫn đến đăng nhầm hoặc đăng trùng."],
            ["Rủi ro chuyển sang Production", "Khi triển khai thật, rủi ro chính nằm ở cấu hình môi trường, độ ổn định worker, token TikTok, storage video, Kafka, scheduler và quyền admin. Vì vậy Sprint 5 phải tập trung vào vận hành, quan sát log, quản trị cấu hình và kiểm thử triển khai Docker."],
        ],
        widths=[1.7, 5.4],
        font_size=9,
    )


def main():
    doc = Document(INPUT_DOCX)

    delete_range(
        doc,
        lambda text: text.startswith("2.5.1.3."),
        lambda text: text.startswith("2.5.2. Phiên bản phần mềm V2.0"),
    )
    v1_anchor = find_paragraph(doc, lambda text: text.startswith("2.5.2. Phiên bản phần mềm V2.0"))
    insert_v1_sections(doc, v1_anchor)

    delete_range(
        doc,
        lambda text: text.startswith("2.5.2.3."),
        lambda text: text.startswith("2.5.3. Phiên bản Production"),
    )
    v2_anchor = find_paragraph(doc, lambda text: text.startswith("2.5.3. Phiên bản Production"))
    insert_v2_sections(doc, v2_anchor)

    doc.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
