from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, Cm
from docx.table import Table
from docx.text.paragraph import Paragraph


INPUT = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc.docx")
OUTPUT = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_2_5_1_2.docx")


SPRINT_ROWS = [
    ["Sprint", "Mục tiêu", "Ngày bắt đầu", "", "Ngày kết thúc", "", "Tiến độ"],
    [
        "Sprint 1",
        "Hoàn thành nền tảng xác thực, phân quyền, đăng xuất, thu thập dữ liệu, chuẩn hóa dữ liệu và nhóm nội dung",
        "20/07/2026",
        "",
        "09/08/2026",
        "",
        "100%",
    ],
    [
        "Sprint 2",
        "Hoàn thành các chức năng AI chọn chủ đề, đề xuất kiểu video, chia nội dung nhiều tập và duyệt kế hoạch",
        "10/08/2026",
        "",
        "23/08/2026",
        "",
        "100%",
    ],
    ["", "", "", "", "", "", ""],
    [
        "Tiêu đề",
        "Mục tiêu",
        "Người thực hiện",
        "Estimate Effort (h)",
        "Ngày bắt đầu",
        "Ngày kết thúc",
        "Priority",
    ],
    [
        "US-01: Đăng nhập vào hệ thống",
        "Người dùng đăng nhập để sử dụng chức năng theo quyền hạn",
        "Nguyễn Văn Phước",
        "8",
        "20/07/2026",
        "26/07/2026",
        "High",
    ],
    [
        "US-02: Kiểm soát phạm vi truy cập dữ liệu",
        "Người dùng được gán vai trò phù hợp và chỉ sử dụng chức năng theo quyền hạn",
        "Nguyễn Văn Phước",
        "8",
        "20/07/2026",
        "26/07/2026",
        "Medium",
    ],
    [
        "US-03: Đăng xuất khỏi hệ thống",
        "Người dùng hủy phiên đăng nhập an toàn khỏi hệ thống",
        "Nguyễn Văn Phước",
        "4",
        "20/07/2026",
        "26/07/2026",
        "High",
    ],
    [
        "US-06: Tạo job lấy dữ liệu",
        "Người dùng tạo được private crawl job theo nguồn và từ khóa",
        "Nguyễn Văn Phước",
        "12",
        "27/07/2026",
        "02/08/2026",
        "High",
    ],
    [
        "US-07: Theo dõi tiến độ crawl",
        "Người dùng xem được trạng thái, tiến độ và log crawl của job thuộc sở hữu",
        "Nguyễn Văn Phước",
        "12",
        "27/07/2026",
        "02/08/2026",
        "High",
    ],
    [
        "US-08: Lấy dữ liệu VNExpress",
        "Hệ thống lấy được bài viết, metadata và hình ảnh từ VNExpress",
        "Nguyễn Văn Phước",
        "12",
        "27/07/2026",
        "02/08/2026",
        "High",
    ],
    [
        "US-09: Chuẩn hóa dữ liệu thô",
        "Dữ liệu crawl được làm sạch và chuyển thành dữ liệu chuẩn",
        "Nguyễn Văn Phước",
        "20",
        "03/08/2026",
        "09/08/2026",
        "High",
    ],
    [
        "US-10: Nhóm nội dung theo truyện/tập",
        "Nội dung được sắp xếp đúng thứ tự để chuẩn bị cho AI planning",
        "Nguyễn Văn Phước",
        "20",
        "03/08/2026",
        "09/08/2026",
        "Medium",
    ],
    [
        "US-11: Quality Score & Deduplication",
        "Hệ thống phát hiện trùng lặp hình ảnh và nội dung bằng fingerprint",
        "Nguyễn Văn Phước",
        "20",
        "03/08/2026",
        "09/08/2026",
        "High",
    ],
    [
        "US-12: Tạo content project",
        "Người dùng tạo được dự án nội dung từ dữ liệu đã crawl",
        "Nguyễn Văn Phước",
        "12",
        "10/08/2026",
        "16/08/2026",
        "High",
    ],
    [
        "US-13: AI đề xuất chủ đề",
        "AI đề xuất chủ đề, góc triển khai và kiểu video phù hợp",
        "Nguyễn Văn Phước",
        "20",
        "10/08/2026",
        "16/08/2026",
        "High",
    ],
    [
        "US-14: AI chia nội dung thành nhiều phần",
        "Kế hoạch chuỗi video có mạch truyện rõ ràng giữa các phần",
        "Nguyễn Văn Phước",
        "32",
        "17/08/2026",
        "23/08/2026",
        "High",
    ],
    [
        "US-15: Duyệt hoặc chỉnh sửa kế hoạch",
        "Người dùng duyệt, từ chối hoặc yêu cầu tạo lại kế hoạch",
        "Nguyễn Văn Phước",
        "12",
        "17/08/2026",
        "23/08/2026",
        "Medium",
    ],
]


STORIES = [
    (
        "US-01: Đăng nhập vào hệ thống",
        [
            "Feature: Xác thực người dùng",
            "Scenario 1: Đăng nhập thành công bằng tài khoản hợp lệ",
            "Given người dùng đã có tài khoản hợp lệ và đang ở màn hình đăng nhập",
            "When người dùng nhập đúng email, mật khẩu và chọn nút Đăng nhập",
            "Then hệ thống xác thực thông tin đăng nhập",
            "And cấp JWT access token cho phiên làm việc",
            "And điều hướng người dùng vào dashboard theo đúng vai trò",
            "Scenario 2: Đăng nhập thất bại do sai thông tin",
            "Given người dùng đang ở màn hình đăng nhập",
            "When người dùng nhập sai email hoặc mật khẩu",
            "Then hệ thống không cấp phiên đăng nhập",
            "And hiển thị thông báo tài khoản hoặc mật khẩu không chính xác",
            "Scenario 3: Truy cập chức năng khi chưa đăng nhập",
            "Given người dùng chưa có access token hợp lệ",
            "When người dùng truy cập màn hình yêu cầu xác thực",
            "Then hệ thống chặn truy cập",
            "And chuyển người dùng về màn hình đăng nhập",
        ],
    ),
    (
        "US-02: Kiểm soát phạm vi truy cập dữ liệu",
        [
            "Feature: Phân quyền và kiểm soát phạm vi dữ liệu",
            "Scenario 1: Người dùng CREATOR chỉ xem dữ liệu cá nhân",
            "Given người dùng đăng nhập với vai trò CREATOR",
            "When người dùng mở danh sách crawl job, content project hoặc kế hoạch AI",
            "Then hệ thống chỉ hiển thị các dữ liệu do chính người dùng đó tạo",
            "And không trả về dữ liệu thuộc tài khoản khác",
            "Scenario 2: Quản trị viên SYSTEM_ADMIN xem dữ liệu toàn hệ thống",
            "Given quản trị viên đăng nhập với vai trò SYSTEM_ADMIN",
            "When quản trị viên truy cập trang quản trị hoặc báo cáo hệ thống",
            "Then hệ thống cho phép xem toàn bộ người dùng, job, dữ liệu và nhật ký xử lý",
            "And các thao tác quản trị được ghi nhận vào audit log",
            "Scenario 3: Từ chối truy cập vượt quyền",
            "Given người dùng không có quyền với tài nguyên được yêu cầu",
            "When người dùng gọi API hoặc mở màn hình ngoài phạm vi quyền hạn",
            "Then hệ thống phản hồi 403 Forbidden",
            "And hiển thị thông báo không có quyền truy cập",
        ],
    ),
    (
        "US-03: Đăng xuất khỏi hệ thống",
        [
            "Feature: Đăng xuất và kết thúc phiên làm việc",
            "Scenario 1: Đăng xuất thành công",
            "Given người dùng đã đăng nhập vào hệ thống",
            "When người dùng chọn chức năng Đăng xuất",
            "Then hệ thống xóa access token khỏi phía client",
            "And kết thúc phiên làm việc hiện tại",
            "And chuyển người dùng về màn hình đăng nhập",
            "Scenario 2: Gửi request sau khi đã đăng xuất",
            "Given người dùng đã thực hiện đăng xuất",
            "When người dùng gửi request tới API yêu cầu xác thực",
            "Then hệ thống phản hồi 401 Unauthorized",
            "And yêu cầu người dùng đăng nhập lại",
        ],
    ),
    (
        "US-06: Tạo job lấy dữ liệu",
        [
            "Feature: Tạo private crawl job",
            "Scenario 1: Hiển thị form tạo job crawl",
            "Given người dùng đã đăng nhập",
            "When người dùng truy cập màn hình thu thập dữ liệu",
            "Then hệ thống hiển thị form gồm nguồn dữ liệu, từ khóa hoặc URL, giới hạn crawl và tùy chọn xử lý",
            "Scenario 2: Tạo job crawl VNExpress thành công",
            "Given người dùng nhập nguồn VNExpress và tham số hợp lệ",
            "When người dùng chọn Tạo job",
            "Then hệ thống tạo crawl job ở trạng thái PENDING",
            "And gắn owner_id của người dùng vào job",
            "And phát sự kiện job_created qua Kafka để crawler xử lý",
            "Scenario 3: Không tạo job khi thiếu dữ liệu bắt buộc",
            "Given người dùng đang ở form tạo job",
            "When người dùng không nhập nguồn, URL hoặc từ khóa bắt buộc",
            "Then hệ thống không tạo job",
            "And hiển thị thông báo yêu cầu bổ sung dữ liệu đầu vào",
        ],
    ),
    (
        "US-07: Theo dõi tiến độ crawl",
        [
            "Feature: Giám sát trạng thái và log crawl job",
            "Scenario 1: Xem danh sách job thuộc sở hữu",
            "Given người dùng đã tạo một hoặc nhiều crawl job",
            "When người dùng mở màn hình Crawl Jobs",
            "Then hệ thống hiển thị danh sách job thuộc sở hữu của người dùng",
            "And mỗi job có trạng thái PENDING, RUNNING, COMPLETED hoặc FAILED",
            "Scenario 2: Theo dõi tiến độ xử lý real-time",
            "Given crawler đang xử lý một job",
            "When hệ thống cập nhật số lượng bài viết đã phát hiện, đã tải và đã chuẩn hóa",
            "Then giao diện hiển thị phần trăm tiến độ mới nhất",
            "And cập nhật trạng thái job mà không cần tải lại trang",
            "Scenario 3: Xem nhật ký lỗi khi crawl thất bại",
            "Given crawl job chuyển sang trạng thái FAILED",
            "When người dùng mở chi tiết job",
            "Then hệ thống hiển thị log lỗi, bước lỗi và thời điểm xảy ra lỗi",
            "And giữ nguyên dữ liệu đã xử lý được để phục vụ chạy lại",
        ],
    ),
    (
        "US-08: Lấy dữ liệu VNExpress",
        [
            "Feature: Thu thập bài viết, metadata và hình ảnh từ VNExpress",
            "Scenario 1: Crawl bài viết VNExpress bằng URL hợp lệ",
            "Given crawler nhận được job VNExpress từ Kafka",
            "When crawler tải nội dung bài viết từ URL hợp lệ",
            "Then hệ thống trích xuất tiêu đề, mô tả, nội dung chính, chuyên mục, thời gian đăng và danh sách hình ảnh",
            "And lưu dữ liệu thô vào MongoDB để phục vụ chuẩn hóa",
            "Scenario 2: Crawl theo từ khóa hoặc chuyên mục",
            "Given job crawl có cấu hình từ khóa hoặc chuyên mục VNExpress",
            "When crawler tìm thấy danh sách bài viết phù hợp",
            "Then hệ thống tạo danh sách bản ghi thô theo từng bài viết",
            "And gắn metadata nguồn, URL gốc và thời điểm thu thập",
            "Scenario 3: Xử lý bài viết thiếu ảnh hoặc thiếu metadata",
            "Given một bài viết VNExpress thiếu ảnh đại diện hoặc thiếu mô tả",
            "When crawler hoàn tất trích xuất dữ liệu",
            "Then hệ thống vẫn lưu nội dung hợp lệ",
            "And đánh dấu các trường thiếu để bước chuẩn hóa chấm điểm chất lượng",
        ],
    ),
    (
        "US-09: Chuẩn hóa dữ liệu thô",
        [
            "Feature: Làm sạch và chuẩn hóa dữ liệu crawl",
            "Scenario 1: Làm sạch nội dung HTML và ký tự nhiễu",
            "Given dữ liệu thô đã được lưu trong MongoDB",
            "When Normalization Service xử lý bản ghi",
            "Then hệ thống loại bỏ thẻ HTML, ký tự rác, khoảng trắng dư và nội dung quảng cáo",
            "And tạo nội dung văn bản sạch để lưu vào PostgreSQL",
            "Scenario 2: Chuẩn hóa metadata bài viết",
            "Given bản ghi thô có tiêu đề, URL, thời gian đăng và danh sách ảnh",
            "When hệ thống chuyển đổi sang mô hình dữ liệu chuẩn",
            "Then hệ thống chuẩn hóa tiêu đề, slug, source_url, published_at và image_urls",
            "And liên kết bản ghi chuẩn với crawl job gốc",
            "Scenario 3: Đánh dấu dữ liệu cần kiểm tra",
            "Given dữ liệu thô thiếu tiêu đề, thiếu nội dung hoặc nội dung quá ngắn",
            "When hệ thống không đạt ngưỡng dữ liệu tối thiểu",
            "Then hệ thống đặt trạng thái NEEDS_REVIEW",
            "And ghi lý do để người dùng hoặc admin kiểm tra lại",
        ],
    ),
    (
        "US-10: Nhóm nội dung theo truyện/tập",
        [
            "Feature: Gom nhóm nội dung theo series và thứ tự tập",
            "Scenario 1: Nhận diện các bài viết thuộc cùng một truyện",
            "Given hệ thống có nhiều bài viết đã chuẩn hóa",
            "When Story Processing Service phân tích tiêu đề, URL và nội dung",
            "Then hệ thống gom các bài viết liên quan vào cùng một content series",
            "And lưu quan hệ giữa series và từng tập",
            "Scenario 2: Sắp xếp tập theo đúng thứ tự",
            "Given một series có nhiều tập được phát hiện",
            "When hệ thống trích xuất số tập hoặc thứ tự xuất hiện từ metadata",
            "Then hệ thống sắp xếp danh sách tập theo trình tự nội dung",
            "And đánh dấu tập đầu, tập tiếp theo và tập cuối nếu xác định được",
            "Scenario 3: Xử lý trường hợp thiếu số tập",
            "Given một số bài viết không có số tập rõ ràng",
            "When hệ thống không chắc chắn về thứ tự",
            "Then hệ thống giữ dữ liệu trong series",
            "And gắn cờ NEEDS_REVIEW_ORDER để tránh tạo kế hoạch sai mạch truyện",
        ],
    ),
    (
        "US-11: Quality Score & Deduplication",
        [
            "Feature: Chấm điểm chất lượng và phát hiện trùng lặp dữ liệu",
            "Scenario 1: Tính điểm chất lượng nội dung",
            "Given bài viết đã được chuẩn hóa",
            "When hệ thống kiểm tra độ dài nội dung, số lượng ảnh, metadata bắt buộc và trạng thái xử lý",
            "Then hệ thống tính quality_score cho từng bản ghi",
            "And phân loại dữ liệu thành READY hoặc NEEDS_REVIEW",
            "Scenario 2: Phát hiện trùng lặp nội dung văn bản",
            "Given hai bài viết có nội dung tương đồng nhưng khác URL",
            "When hệ thống tạo text fingerprint từ tiêu đề chuẩn hóa và phần nội dung chính",
            "Then hệ thống đánh dấu các bản ghi có fingerprint trùng hoặc gần trùng",
            "And tăng duplicate_count cho nhóm dữ liệu liên quan",
            "Scenario 3: Phát hiện trùng lặp hình ảnh",
            "Given hai bài viết sử dụng cùng ảnh nhưng khác subdomain CDN",
            "When hệ thống áp dụng image fingerprinting cho danh sách ảnh",
            "Then hệ thống nhận diện ảnh trùng lặp",
            "And tránh chọn nhiều dữ liệu giống nhau cho AI planning",
        ],
    ),
    (
        "US-12: Tạo content project",
        [
            "Feature: Khởi tạo dự án nội dung từ dữ liệu đã crawl",
            "Scenario 1: Tạo content project từ bài viết READY",
            "Given người dùng có dữ liệu đã chuẩn hóa ở trạng thái READY",
            "When người dùng chọn bài viết và tạo dự án nội dung",
            "Then hệ thống tạo content project mới ở trạng thái DRAFT",
            "And liên kết project với dữ liệu nguồn và người sở hữu",
            "Scenario 2: Tạo project từ series nhiều tập",
            "Given người dùng chọn một content series đã được gom nhóm",
            "When người dùng yêu cầu tạo project",
            "Then hệ thống khởi tạo project chứa danh sách tập theo đúng thứ tự",
            "And lưu context series để chuyển sang bước AI planning",
            "Scenario 3: Chặn tạo project từ dữ liệu chưa đạt chất lượng",
            "Given dữ liệu nguồn đang ở trạng thái NEEDS_REVIEW",
            "When người dùng tạo content project",
            "Then hệ thống cảnh báo dữ liệu chưa đạt điều kiện",
            "And không cho chuyển sang AI planning cho đến khi dữ liệu được xác nhận",
        ],
    ),
    (
        "US-13: AI đề xuất chủ đề",
        [
            "Feature: AI phân tích nguồn nội dung và đề xuất hướng triển khai video",
            "Scenario 1: Đề xuất chủ đề và góc khai thác",
            "Given content project đã được tạo từ dữ liệu hợp lệ",
            "When người dùng chọn tạo kế hoạch bằng AI",
            "Then hệ thống gửi nội dung, metadata và mục tiêu sản xuất tới AI Planning Service",
            "And AI trả về danh sách chủ đề, góc triển khai, đối tượng mục tiêu và giọng điệu đề xuất",
            "Scenario 2: Đề xuất kiểu video phù hợp",
            "Given AI đã phân tích nội dung nguồn",
            "When hệ thống tạo kết quả planning",
            "Then AI đề xuất kiểu video như kể chuyện, tóm tắt tin tức hoặc series nhiều phần",
            "And giải thích lý do chọn kiểu video theo đặc điểm nội dung",
            "Scenario 3: Xử lý lỗi AI provider",
            "Given người dùng yêu cầu AI đề xuất chủ đề",
            "When AI provider quá tải hoặc trả về lỗi",
            "Then hệ thống lưu trạng thái planning là FAILED",
            "And hiển thị thông báo lỗi cùng tùy chọn thử lại",
        ],
    ),
    (
        "US-14: AI chia nội dung thành nhiều phần",
        [
            "Feature: AI chia nội dung dài thành kế hoạch chuỗi video",
            "Scenario 1: Chia nội dung thành nhiều phần có cấu trúc",
            "Given content project chứa bài viết dài hoặc series nhiều tập",
            "When người dùng yêu cầu AI chia nội dung",
            "Then hệ thống tạo kế hoạch gồm nhiều phần video",
            "And mỗi phần có tiêu đề, mục tiêu nội dung, tóm tắt, hook mở đầu và thời lượng dự kiến",
            "Scenario 2: Giữ mạch truyện giữa các phần",
            "Given kế hoạch gồm nhiều phần liên tiếp",
            "When AI tạo nội dung cho từng phần",
            "Then hệ thống lưu context của phần trước",
            "And phần sau kế thừa nhân vật, sự kiện chính và điểm dừng của phần trước",
            "Scenario 3: Điều chỉnh số phần theo yêu cầu",
            "Given người dùng muốn thay đổi số tập hoặc độ dài từng phần",
            "When người dùng cập nhật yêu cầu và tạo lại kế hoạch",
            "Then AI tái phân bổ nội dung",
            "And vẫn bảo đảm trình tự câu chuyện rõ ràng, không lặp ý và không bỏ sót ý chính",
        ],
    ),
    (
        "US-15: Duyệt hoặc chỉnh sửa kế hoạch",
        [
            "Feature: Duyệt, chỉnh sửa hoặc tạo lại kế hoạch nội dung AI",
            "Scenario 1: Duyệt kế hoạch phù hợp",
            "Given AI đã tạo kế hoạch nội dung cho project",
            "When người dùng kiểm tra và chọn Chấp nhận kế hoạch",
            "Then hệ thống cập nhật trạng thái kế hoạch thành APPROVED",
            "And cho phép chuyển sang bước sản xuất video ở phiên bản tiếp theo",
            "Scenario 2: Chỉnh sửa kế hoạch trước khi duyệt",
            "Given kế hoạch AI có một số tiêu đề, hook hoặc thứ tự phần chưa phù hợp",
            "When người dùng chỉnh sửa nội dung kế hoạch và lưu lại",
            "Then hệ thống cập nhật phiên bản kế hoạch mới",
            "And giữ lịch sử thay đổi để đối chiếu khi cần",
            "Scenario 3: Từ chối và yêu cầu tạo lại kế hoạch",
            "Given người dùng chưa hài lòng với kế hoạch hiện tại",
            "When người dùng nhập lý do và chọn tạo lại kế hoạch",
            "Then hệ thống lưu phản hồi của người dùng",
            "And gửi lại context cùng yêu cầu điều chỉnh cho AI Planning Service",
        ],
    ),
]


def iter_body_blocks(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def set_paragraph_text(paragraph, text, bold=False):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.bold = bold


def remove_block(block):
    element = block._element
    element.getparent().remove(element)


def insert_table_after(paragraph, table):
    paragraph._p.addnext(table._tbl)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)


def add_formatted_line(cell, text, font_size=10):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0

    label = None
    rest = None
    if text.startswith("Feature:"):
        label, rest = "Feature", text[len("Feature") :]
    elif text.startswith("Scenario "):
        prefix, rest_part = text.split(":", 1)
        label, rest = prefix, ":" + rest_part
    else:
        for token in ["Given", "When", "Then", "And"]:
            if text.startswith(token + " "):
                label, rest = token, text[len(token) :]
                break

    if label is None:
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        return p

    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(font_size)
    r2 = p.add_run(rest)
    r2.font.size = Pt(font_size)
    return p


def replace_story_table(table, lines):
    cell = table.cell(0, 0)
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for line in lines:
        add_formatted_line(cell, line)


def apply_sprint_table_format(table):
    widths = [1450, 2250, 1050, 850, 1250, 1250, 1260]
    set_table_width(table, widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if ci in (2, 3, 4, 5, 6):
                tc_pr = cell._tc.get_or_add_tcPr()
                if tc_pr.find(qn("w:noWrap")) is None:
                    tc_pr.append(OxmlElement("w:noWrap"))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7)
            if ri in (0, 4):
                shade_cell(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            elif ri == 3:
                shade_cell(cell, "F2F2F2")


def build_sprint_planning_table(document):
    table = document.add_table(rows=len(SPRINT_ROWS), cols=7)
    table.style = "Table Grid"
    for ri, row in enumerate(SPRINT_ROWS):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = text
    for ri in range(3):
        table.cell(ri, 2).merge(table.cell(ri, 3))
        table.cell(ri, 4).merge(table.cell(ri, 5))
    apply_sprint_table_format(table)
    return table


def main():
    document = Document(INPUT)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Bảng xx. Sprint planning"):
            set_paragraph_text(paragraph, "Bảng 2.6. Sprint planning cho phiên bản V1.0")
            sprint_caption = paragraph
        elif text.startswith("Bảng 2.6. Sprint backlog cho phiên bản V1.0"):
            set_paragraph_text(paragraph, "")

    # Remove old spreadsheet-only filler lines directly after the sprint planning caption.
    removing = False
    for block in list(iter_body_blocks(document)):
        if isinstance(block, Paragraph) and block.text.strip() == "Bảng 2.6. Sprint planning cho phiên bản V1.0":
            removing = True
            continue
        if removing:
            if isinstance(block, Paragraph) and block.text.strip() == "2.5.1.2. Xây dựng kịch bản câu chuyện người dùng":
                break
            if isinstance(block, Paragraph):
                txt = block.text.strip()
                if (
                    not txt
                    or txt.startswith("Xem chi tiết")
                    or txt.startswith("https://docs.google.com")
                    or txt.startswith("Bảng 2.6. Sprint backlog")
                    or txt.startswith("Chi tiết xem tại")
                ):
                    remove_block(block)

    table = build_sprint_planning_table(document)
    insert_table_after(sprint_caption, table)

    blocks = list(iter_body_blocks(document))
    in_section = False
    story_headings = []
    story_tables = []
    last_table = None
    blanks_after_table = []
    for block in blocks:
        if isinstance(block, Paragraph) and block.text.strip() == "2.5.1.2. Xây dựng kịch bản câu chuyện người dùng":
            in_section = True
            continue
        if in_section and isinstance(block, Paragraph) and block.text.strip().startswith("2.5.1.3."):
            break
        if not in_section:
            continue
        if isinstance(block, Table):
            story_tables.append(block)
            last_table = block
        elif isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt.startswith("US-"):
                story_headings.append(block)
            elif not txt and last_table is not None:
                blanks_after_table.append(block)

    for (heading, table_block), (title, lines) in zip(zip(story_headings, story_tables), STORIES):
        set_paragraph_text(heading, title, bold=True)
        replace_story_table(table_block, lines)

    for extra_heading in story_headings[len(STORIES) :]:
        remove_block(extra_heading)
    for extra_table in story_tables[len(STORIES) :]:
        remove_block(extra_table)
    for extra_blank in blanks_after_table[len(STORIES) :]:
        remove_block(extra_blank)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
