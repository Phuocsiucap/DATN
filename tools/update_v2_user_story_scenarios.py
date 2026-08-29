from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


INPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2_3.docx")
OUTPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2_3_chi_tiet_2_5_2_2.docx")


SECTION_HEADING = "2.5.2.2. Xây dựng kịch bản câu chuyện người dùng"
NEXT_HEADING_PREFIX = "2.5.2.3."


def iter_body_blocks(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def remove_blocks_between(doc, start_text, end_prefix):
    blocks = list(iter_body_blocks(doc))
    start_idx = None
    end_idx = None
    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.text.strip() == start_text:
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError(f"Cannot find start heading: {start_text}")

    for i in range(start_idx + 1, len(blocks)):
        block = blocks[i]
        if isinstance(block, Paragraph) and block.text.strip().startswith(end_prefix):
            end_idx = i
            break
    if end_idx is None:
        raise RuntimeError(f"Cannot find next heading prefix: {end_prefix}")

    for block in blocks[start_idx + 1 : end_idx]:
        block._element.getparent().remove(block._element)

    # Return a fresh anchor because the body tree has changed.
    for block in iter_body_blocks(doc):
        if isinstance(block, Paragraph) and block.text.strip().startswith(end_prefix):
            return block
    raise RuntimeError("Cannot refind next heading after deletion")


def add_labeled_paragraph(anchor, label, text):
    p = anchor.insert_paragraph_before()
    if label:
        r = p.add_run(label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_plain(anchor, text=""):
    return anchor.insert_paragraph_before(text)


def add_us_heading(anchor, text):
    p = anchor.insert_paragraph_before()
    r = p.add_run(text)
    r.bold = True
    return p


def add_feature(anchor, text):
    add_labeled_paragraph(anchor, "Feature: ", text)


def add_scenario(anchor, title, actor, preconditions, main_flow, exceptions, result):
    add_labeled_paragraph(anchor, "Scenario: ", title)
    add_labeled_paragraph(anchor, "Tác nhân chính: ", actor)
    add_labeled_paragraph(anchor, "Tiền điều kiện: ", preconditions)
    for idx, step in enumerate(main_flow, start=1):
        add_labeled_paragraph(anchor, f"Luồng chính {idx}: ", step)
    for idx, item in enumerate(exceptions, start=1):
        add_labeled_paragraph(anchor, f"Ngoại lệ {idx}: ", item)
    add_labeled_paragraph(anchor, "Kết quả sau cùng: ", result)


USER_STORIES = [
    {
        "heading": "US-16: Tạo kịch bản, chia cảnh và timeline video",
        "feature": "Chuyển kế hoạch nội dung đã duyệt thành kịch bản sản xuất video có cấu trúc scene, lời thoại, visual prompt và timeline.",
        "scenarios": [
            {
                "title": "Sinh kịch bản video từ content project đã được duyệt",
                "actor": "Creator đã đăng nhập và sở hữu content project.",
                "preconditions": "Content project đã có kế hoạch AI ở trạng thái APPROVED, có chủ đề, góc triển khai, danh sách phần/tập và dữ liệu nội dung đầu vào hợp lệ.",
                "main_flow": [
                    "Người dùng mở content project, chọn tập hoặc phần cần sản xuất và bấm tạo kịch bản video.",
                    "API Service kiểm tra quyền sở hữu, trạng thái kế hoạch và tạo yêu cầu generate-video.script.requested gửi sang worker xử lý media.",
                    "AI Media Engine đọc nội dung đã chuẩn hóa, tóm tắt kế hoạch đã duyệt và tạo bản nháp gồm tiêu đề video, hook mở đầu, body, kết luận và lời kêu gọi hành động.",
                    "Hệ thống lưu kịch bản vào cơ sở dữ liệu, gắn với project, tập, người tạo và trạng thái SCRIPT_READY để các bước tiếp theo sử dụng.",
                ],
                "exceptions": [
                    "Nếu kế hoạch chưa được duyệt, hệ thống từ chối tạo kịch bản và yêu cầu người dùng quay lại bước duyệt kế hoạch.",
                    "Nếu dữ liệu đầu vào thiếu nội dung chính hoặc bị đánh dấu chất lượng thấp, hệ thống hiển thị thông báo dữ liệu không đủ để sản xuất video.",
                    "Nếu AI trả về nội dung rỗng hoặc lỗi định dạng, worker ghi log thất bại và cho phép chạy lại mà không làm mất dữ liệu cũ.",
                ],
                "result": "Người dùng có bản kịch bản video đầu tiên đủ nội dung để chuyển sang bước chia cảnh và timeline.",
            },
            {
                "title": "Tự động chia kịch bản thành các scene sản xuất",
                "actor": "Creator và AI Media Engine.",
                "preconditions": "Kịch bản tổng đã tạo thành công và chưa bị khóa bởi một workflow render đang chạy.",
                "main_flow": [
                    "Người dùng chọn chức năng chia cảnh hoặc hệ thống tự động thực hiện ngay sau khi kịch bản được tạo.",
                    "AI Media Engine tách kịch bản thành nhiều scene theo mạch nội dung: mở đầu, diễn biến, cao trào, kết thúc hoặc CTA.",
                    "Mỗi scene được gán thứ tự, lời thoại, mô tả hình ảnh, prompt tạo hình, thời lượng dự kiến và ghi chú cảm xúc.",
                    "Hệ thống kiểm tra tổng thời lượng scene có phù hợp độ dài video mục tiêu và lưu danh sách scene vào timeline draft.",
                ],
                "exceptions": [
                    "Nếu một scene quá dài, hệ thống tự chia nhỏ hoặc đánh dấu cần chỉnh sửa.",
                    "Nếu scene thiếu lời thoại hoặc thiếu visual prompt, hệ thống đánh dấu cảnh chưa đạt để người dùng chỉnh trước khi tạo voice.",
                    "Nếu thứ tự scene bị trùng hoặc thiếu, hệ thống chuẩn hóa lại index trước khi lưu.",
                ],
                "result": "Danh sách scene có đủ lời thoại và visual prompt, sẵn sàng cho chỉnh sửa AI và tạo audio.",
            },
            {
                "title": "Khởi tạo timeline theo thứ tự cảnh và thời lượng dự kiến",
                "actor": "Hệ thống sản xuất video.",
                "preconditions": "Danh sách scene đã hợp lệ, mỗi scene có nội dung thoại và mô tả hình ảnh.",
                "main_flow": [
                    "Hệ thống tạo timeline gồm các track: voice, subtitle, visual layer, background music nếu có và metadata xuất video.",
                    "Mỗi scene được ánh xạ thành một đoạn timeline với start_time, end_time dự kiến, thứ tự hiển thị và tài nguyên media cần tạo.",
                    "Timeline được gắn mã workflow để đồng bộ với các bước tạo giọng đọc, phụ đề và render MP4.",
                    "Người dùng xem được cấu trúc timeline ở màn hình sản xuất và biết scene nào đã sẵn sàng, scene nào còn thiếu dữ liệu.",
                ],
                "exceptions": [
                    "Nếu tổng thời lượng vượt quá giới hạn cấu hình video ngắn, hệ thống cảnh báo để người dùng rút gọn hoặc cho phép AI tối ưu lại.",
                    "Nếu project đã tồn tại timeline cũ, hệ thống tạo revision mới thay vì ghi đè mất lịch sử.",
                    "Nếu đang có render job chạy, hệ thống khóa chỉnh sửa timeline để tránh mất đồng bộ.",
                ],
                "result": "Timeline video được tạo có cấu trúc rõ ràng, là đầu vào thống nhất cho voice, subtitle và render.",
            },
        ],
    },
    {
        "heading": "US-17: Chỉnh sửa kịch bản hoặc cảnh bằng AI",
        "feature": "Cho phép người dùng chỉnh một scene hoặc toàn bộ kịch bản bằng prompt AI nhưng vẫn giữ ràng buộc nội dung, thời lượng và mạch truyện.",
        "scenarios": [
            {
                "title": "Chỉnh sửa một scene bằng prompt của người dùng",
                "actor": "Creator.",
                "preconditions": "Timeline đã có danh sách scene, scene được chọn thuộc project của người dùng và chưa ở trạng thái đang render.",
                "main_flow": [
                    "Người dùng mở chi tiết scene, nhập yêu cầu chỉnh sửa như tăng kịch tính, làm câu thoại tự nhiên hơn hoặc thêm yếu tố hấp dẫn.",
                    "Hệ thống gửi prompt, nội dung scene hiện tại, ngữ cảnh các scene liền trước và liền sau sang AI Media Engine.",
                    "AI trả về phiên bản scene mới gồm lời thoại, visual prompt, mô tả cảm xúc và thời lượng dự kiến.",
                    "Người dùng xem so sánh trước/sau và xác nhận áp dụng thay đổi.",
                ],
                "exceptions": [
                    "Nếu prompt vi phạm chính sách nội dung hoặc làm lệch chủ đề, hệ thống từ chối áp dụng và hiển thị lý do.",
                    "Nếu AI sinh nội dung làm thời lượng vượt giới hạn, hệ thống yêu cầu tối ưu lại hoặc cho phép người dùng chỉnh thủ công.",
                    "Nếu người dùng hủy thay đổi, hệ thống giữ nguyên scene hiện tại.",
                ],
                "result": "Scene được cập nhật có kiểm soát, đồng thời lưu revision để có thể truy vết hoặc khôi phục.",
            },
            {
                "title": "Chỉnh tone toàn bộ kịch bản nhưng giữ continuity",
                "actor": "Creator và AI Media Engine.",
                "preconditions": "Kịch bản gồm nhiều scene đã được tạo, series context nếu có đã được lưu trong MongoDB.",
                "main_flow": [
                    "Người dùng chọn chỉnh toàn bộ kịch bản và nhập hướng chỉnh như giọng kể hài hước, nghiêm túc, bí ẩn hoặc ngắn gọn hơn.",
                    "Hệ thống gom toàn bộ scene, thông tin nhân vật, bối cảnh, mục tiêu video và ràng buộc thời lượng gửi cho AI.",
                    "AI trả về phiên bản mới nhưng giữ nguyên thứ tự sự kiện, nhân vật, thông tin cốt lõi và CTA.",
                    "Hệ thống cập nhật script revision và đánh dấu các phần thay đổi để người dùng kiểm tra nhanh.",
                ],
                "exceptions": [
                    "Nếu AI làm thay đổi sự kiện chính, hệ thống phát hiện bằng bước validation và yêu cầu tạo lại.",
                    "Nếu series context không tồn tại, hệ thống chỉ chỉnh theo nội dung hiện tại và ghi chú không có dữ liệu kế thừa.",
                    "Nếu người dùng không xác nhận, revision mới được lưu nháp nhưng không thay thế bản đang dùng.",
                ],
                "result": "Kịch bản có tone phù hợp hơn mà không phá vỡ mạch nội dung của video hoặc series.",
            },
            {
                "title": "Khôi phục hoặc chọn lại revision sau khi chỉnh AI",
                "actor": "Creator.",
                "preconditions": "Một scene hoặc kịch bản đã có ít nhất hai revision.",
                "main_flow": [
                    "Người dùng mở lịch sử chỉnh sửa của scene hoặc kịch bản.",
                    "Hệ thống hiển thị thời gian chỉnh, prompt đã dùng, người thực hiện và nội dung thay đổi chính.",
                    "Người dùng chọn một revision cũ để xem trước và bấm khôi phục.",
                    "Hệ thống đặt revision được chọn làm bản hiện hành, đồng thời tạo log audit cho thao tác khôi phục.",
                ],
                "exceptions": [
                    "Nếu revision cũ không còn tương thích timeline mới, hệ thống cảnh báo các trường cần đồng bộ lại.",
                    "Nếu người dùng không có quyền với project, hệ thống không hiển thị lịch sử revision.",
                    "Nếu khôi phục làm thiếu voice/subtitle đã tạo, hệ thống chuyển workflow về trạng thái cần tạo lại audio và phụ đề.",
                ],
                "result": "Người dùng kiểm soát được lịch sử nội dung, giảm rủi ro mất bản kịch bản tốt.",
            },
        ],
    },
    {
        "heading": "US-18: Tạo giọng đọc cho video",
        "feature": "Sản xuất file audio giọng đọc từ lời thoại đã duyệt, lưu thời lượng chính xác để đồng bộ phụ đề và hình ảnh.",
        "scenarios": [
            {
                "title": "Tạo file audio MP3 từ lời thoại của timeline",
                "actor": "Creator và TTS Engine.",
                "preconditions": "Các scene đã có lời thoại hợp lệ, không còn cảnh rỗng và người dùng đã chọn cấu hình giọng đọc.",
                "main_flow": [
                    "Người dùng chọn giọng đọc, ngôn ngữ, tốc độ đọc và bấm tạo giọng đọc.",
                    "API Service kiểm tra trạng thái timeline rồi gửi generate-video.voice.requested cho worker.",
                    "Worker gom lời thoại theo đúng thứ tự scene, chuẩn hóa dấu câu, khoảng nghỉ và gửi sang TTS Engine.",
                    "TTS Engine trả về file audio MP3, hệ thống lưu đường dẫn file và thời lượng audio đo được.",
                ],
                "exceptions": [
                    "Nếu có scene thiếu lời thoại, hệ thống chặn tạo voice và chỉ rõ scene cần bổ sung.",
                    "Nếu dịch vụ TTS tạm lỗi, workflow chuyển sang VOICE_FAILED và cho phép retry.",
                    "Nếu audio trả về rỗng hoặc không đọc được duration, hệ thống không chuyển bước và ghi log lỗi.",
                ],
                "result": "Timeline có file audio chính thức cùng duration làm nguồn chuẩn cho bước đồng bộ phụ đề.",
            },
            {
                "title": "Tạo lại giọng đọc sau khi người dùng chỉnh kịch bản",
                "actor": "Creator.",
                "preconditions": "Audio cũ đã tồn tại nhưng script revision hiện tại mới hơn revision dùng để tạo audio.",
                "main_flow": [
                    "Hệ thống hiển thị cảnh báo audio đang không khớp với kịch bản mới.",
                    "Người dùng bấm tạo lại giọng đọc.",
                    "Worker tạo audio mới từ script revision hiện hành và gắn audio với revision mới.",
                    "Các phụ đề cũ bị đánh dấu cần đồng bộ lại vì duration đã thay đổi.",
                ],
                "exceptions": [
                    "Nếu người dùng hủy thao tác, audio cũ vẫn được giữ nhưng workflow không cho render bản cuối.",
                    "Nếu lần tạo lại thất bại, audio cũ không bị xóa để người dùng còn dữ liệu tham chiếu.",
                    "Nếu duration mới lệch quá nhiều so với mục tiêu, hệ thống cảnh báo cần rút gọn hoặc kéo dài script.",
                ],
                "result": "Audio luôn tương ứng với bản kịch bản đang dùng, tránh lỗi lời thoại và phụ đề không khớp.",
            },
            {
                "title": "Kiểm tra cấu hình giọng đọc trước khi gửi TTS",
                "actor": "Hệ thống.",
                "preconditions": "Người dùng đã chọn thông số giọng đọc trên giao diện.",
                "main_flow": [
                    "Hệ thống kiểm tra voice_id, language, speaking_rate và giới hạn ký tự của nhà cung cấp TTS.",
                    "Nếu cấu hình hợp lệ, hệ thống tạo request voice job và hiển thị trạng thái đang xử lý.",
                    "Worker ghi log từng bước để người dùng hoặc quản trị viên có thể theo dõi.",
                ],
                "exceptions": [
                    "Nếu voice_id không tồn tại, hệ thống yêu cầu chọn giọng đọc khác.",
                    "Nếu văn bản vượt giới hạn, hệ thống tự chia đoạn hoặc yêu cầu rút gọn tùy cấu hình.",
                    "Nếu người dùng gửi nhiều yêu cầu liên tiếp, hệ thống chống trùng job để tránh tạo nhiều audio không cần thiết.",
                ],
                "result": "Giảm lỗi TTS trước khi gọi dịch vụ bên ngoài và bảo đảm workflow có trạng thái minh bạch.",
            },
        ],
    },
    {
        "heading": "US-19: Tạo phụ đề và đồng bộ với giọng đọc",
        "feature": "Tạo subtitle và căn chỉnh timestamp của phụ đề, hình ảnh theo file audio để video không bị trôi tiếng.",
        "scenarios": [
            {
                "title": "Sinh phụ đề từ lời thoại và audio đã tạo",
                "actor": "Creator và Subtitle Sync Worker.",
                "preconditions": "Audio MP3 đã tạo thành công, duration đọc được và script revision không thay đổi sau khi tạo audio.",
                "main_flow": [
                    "Người dùng bấm tạo phụ đề hoặc hệ thống tự tạo sau khi voice job hoàn thành.",
                    "Worker tách lời thoại thành các cụm phụ đề ngắn, dễ đọc trên màn hình dọc.",
                    "Hệ thống gán start_time, end_time cho từng subtitle segment theo duration audio.",
                    "Phụ đề được lưu cùng timeline và hiển thị cho người dùng xem trước.",
                ],
                "exceptions": [
                    "Nếu audio thiếu duration, hệ thống không tạo phụ đề và yêu cầu tạo lại voice.",
                    "Nếu một dòng phụ đề quá dài, hệ thống tự chia thành nhiều segment nhỏ hơn.",
                    "Nếu script đã bị chỉnh sau khi tạo audio, hệ thống chặn đồng bộ để tránh sai lời thoại.",
                ],
                "result": "Video có danh sách subtitle rõ ràng, đúng thứ tự và có timestamp để render.",
            },
            {
                "title": "Căn chỉnh hình ảnh và phụ đề khớp 1:1 với audio",
                "actor": "Hệ thống timeline.",
                "preconditions": "Timeline đã có scene, audio và subtitle segment.",
                "main_flow": [
                    "Worker chạy thuật toán fit_video_clips_to_text để phân bổ thời lượng từng scene theo văn bản và audio.",
                    "Mỗi visual clip được kéo dài hoặc rút ngắn để khớp phần lời thoại tương ứng.",
                    "Subtitle segment được căn vào thời điểm lời thoại xuất hiện, không vượt quá thời lượng scene.",
                    "Hệ thống ghi lại timeline cuối cùng dùng cho render MP4.",
                ],
                "exceptions": [
                    "Nếu tổng duration subtitle vượt duration audio, hệ thống scale lại timestamp theo duration chuẩn.",
                    "Nếu một scene không có hình ảnh, hệ thống dùng placeholder hoặc đánh dấu thiếu media tùy cấu hình.",
                    "Nếu thuật toán phát hiện khoảng lặng dài, hệ thống phân bổ khoảng nghỉ vào transition thay vì để phụ đề trống bất thường.",
                ],
                "result": "Khung hình, giọng đọc và phụ đề được đồng bộ, giảm lỗi trôi chữ hoặc lệch tiếng.",
            },
            {
                "title": "Cho phép kiểm tra nhanh subtitle trước khi render",
                "actor": "Creator.",
                "preconditions": "Subtitle đã được sinh và timeline đã đồng bộ.",
                "main_flow": [
                    "Người dùng mở màn hình timeline preview và xem danh sách phụ đề theo mốc thời gian.",
                    "Hệ thống cho phép chỉnh nội dung phụ đề nhỏ mà không thay đổi lời thoại audio.",
                    "Khi lưu chỉnh sửa, hệ thống kiểm tra độ dài dòng và không cho thay đổi timestamp ngoài phạm vi audio.",
                    "Các thay đổi được lưu thành subtitle revision để render dùng bản mới nhất.",
                ],
                "exceptions": [
                    "Nếu người dùng chỉnh khác quá nhiều so với lời thoại, hệ thống cảnh báo phụ đề không còn khớp audio.",
                    "Nếu chỉnh sửa làm chữ tràn khung, hệ thống gợi ý chia dòng hoặc rút gọn.",
                    "Nếu timeline đang render, hệ thống khóa chỉnh sửa subtitle cho đến khi job kết thúc.",
                ],
                "result": "Người dùng có thể sửa lỗi chính tả và câu chữ phụ đề mà vẫn giữ đồng bộ kỹ thuật.",
            },
        ],
    },
    {
        "heading": "US-20: Ghép hình ảnh, giọng đọc, phụ đề thành video MP4",
        "feature": "Render video cuối cùng từ timeline đã đồng bộ bằng AI Media Engine, Remotion hoặc FFmpeg.",
        "scenarios": [
            {
                "title": "Render video MP4 thành công",
                "actor": "Creator, API Service và Render Worker.",
                "preconditions": "Timeline đã có audio, subtitle, visual clip hoặc hình ảnh cho từng scene và trạng thái QA kỹ thuật đạt điều kiện render.",
                "main_flow": [
                    "Người dùng bấm xuất video MP4 trên màn hình sản xuất.",
                    "API Service tạo render job, lưu trạng thái RENDERING và phát sự kiện generate-video.render.requested.",
                    "Render Worker lấy timeline, tải tài nguyên media, dựng layout video dọc, ghép voice, subtitle và hình ảnh theo timestamp.",
                    "Sau khi render xong, hệ thống lưu file MP4, thumbnail, duration, dung lượng file và chuyển workflow sang RENDERED.",
                ],
                "exceptions": [
                    "Nếu thiếu audio hoặc subtitle, hệ thống chặn render và hiển thị bước còn thiếu.",
                    "Nếu một media file không tải được, render job chuyển FAILED và ghi rõ tài nguyên lỗi.",
                    "Nếu quá trình render bị timeout, hệ thống cho phép retry từ cùng timeline để tránh phải tạo lại kịch bản.",
                ],
                "result": "Người dùng nhận được file video MP4 hoàn chỉnh để xem trước và duyệt đăng.",
            },
            {
                "title": "Render lại sau khi timeline bị chỉnh sửa",
                "actor": "Creator.",
                "preconditions": "Video MP4 cũ đã tồn tại nhưng timeline hoặc subtitle có revision mới hơn.",
                "main_flow": [
                    "Hệ thống hiển thị cảnh báo video hiện tại không còn là bản mới nhất.",
                    "Người dùng bấm render lại.",
                    "Render Worker tạo file MP4 mới từ timeline revision hiện hành và giữ file cũ dưới dạng lịch sử.",
                    "Màn hình preview tự chuyển sang bản video mới sau khi job hoàn tất.",
                ],
                "exceptions": [
                    "Nếu người dùng không render lại, hệ thống không cho duyệt bản cũ làm bản xuất bản cuối.",
                    "Nếu render mới thất bại, video cũ vẫn còn để xem lại nhưng không được đánh dấu latest.",
                    "Nếu có nhiều render job trùng nhau, hệ thống chỉ cho một job active trên cùng workflow.",
                ],
                "result": "Bản MP4 cuối luôn phản ánh đúng timeline mới nhất.",
            },
            {
                "title": "Theo dõi log và trạng thái render",
                "actor": "Creator và quản trị viên.",
                "preconditions": "Render job đã được tạo.",
                "main_flow": [
                    "Người dùng theo dõi trạng thái QUEUED, RENDERING, COMPLETED hoặc FAILED trên giao diện.",
                    "Worker cập nhật phần trăm tiến độ theo các bước tải media, dựng composition, encode và upload kết quả.",
                    "Nếu quản trị viên mở log, hệ thống hiển thị thông tin lỗi kỹ thuật, thời gian xử lý và worker phụ trách.",
                ],
                "exceptions": [
                    "Nếu mất kết nối thời gian thực, giao diện vẫn lấy lại trạng thái mới nhất bằng API polling.",
                    "Nếu job bị hủy, hệ thống dừng worker khi có thể và chuyển trạng thái CANCELED.",
                    "Nếu log quá dài, hệ thống chỉ hiển thị phần quan trọng và lưu đầy đủ ở backend.",
                ],
                "result": "Quá trình render minh bạch, dễ kiểm tra và dễ retry khi có sự cố.",
            },
        ],
    },
    {
        "heading": "US-21: Quản lý chuỗi video nhiều tập và duy trì ngữ cảnh xuyên suốt",
        "feature": "Lưu và tái sử dụng series context để các tập video trong cùng chuỗi có nhân vật, bối cảnh, giọng kể và mạch truyện nhất quán.",
        "scenarios": [
            {
                "title": "Lưu ngữ cảnh sau khi hoàn thành một tập",
                "actor": "AI Media Engine.",
                "preconditions": "Một video trong series đã có kịch bản hoặc đã render xong.",
                "main_flow": [
                    "Hệ thống trích xuất các yếu tố quan trọng: nhân vật, bối cảnh, sự kiện chính, tone giọng, cliffhanger và thông tin cần nhớ.",
                    "Series context được lưu vào MongoDB gắn với content project, series_id, episode_number và revision.",
                    "Nếu tập đã được duyệt, context được đánh dấu là nguồn tin cậy cho tập tiếp theo.",
                    "Người dùng có thể xem tóm tắt context trên màn hình quản lý series.",
                ],
                "exceptions": [
                    "Nếu nội dung tập quá ngắn, hệ thống chỉ lưu tóm tắt tối thiểu và cảnh báo context yếu.",
                    "Nếu có nhiều bản revision của cùng tập, hệ thống ưu tiên revision đã duyệt.",
                    "Nếu lưu MongoDB thất bại, workflow không mất video nhưng đánh dấu cần đồng bộ lại context.",
                ],
                "result": "Mỗi tập tạo ra dữ liệu ngữ cảnh có thể dùng cho các tập sau.",
            },
            {
                "title": "Tạo tập tiếp theo kế thừa context tập trước",
                "actor": "Creator và AI Media Engine.",
                "preconditions": "Series đã có ít nhất một tập và context đã được lưu.",
                "main_flow": [
                    "Người dùng chọn tạo tập tiếp theo trong cùng series.",
                    "Hệ thống tải series context gần nhất, kế hoạch nội dung và ràng buộc độ dài video.",
                    "AI tạo kịch bản mới có nhắc lại vừa đủ thông tin cũ, tiếp tục xung đột hoặc chủ đề còn dang dở.",
                    "Timeline mới được gắn episode_number kế tiếp và liên kết với series context nguồn.",
                ],
                "exceptions": [
                    "Nếu không tìm thấy context, hệ thống cho phép tạo tập mới độc lập nhưng cảnh báo continuity không đảm bảo.",
                    "Nếu context mâu thuẫn với dữ liệu mới, hệ thống yêu cầu người dùng chọn ưu tiên dữ liệu cũ hoặc kế hoạch mới.",
                    "Nếu người dùng sửa tập trước sau khi tạo tập sau, hệ thống đánh dấu các tập sau có thể cần rebuild context.",
                ],
                "result": "Các tập trong series giữ được mạch truyện và nhận diện nội dung nhất quán.",
            },
            {
                "title": "Cảnh báo lỗi continuity giữa các tập",
                "actor": "Hệ thống validation.",
                "preconditions": "Series có nhiều tập hoặc đang tạo tập mới dựa trên context cũ.",
                "main_flow": [
                    "Hệ thống so sánh nhân vật, tên riêng, mốc sự kiện, bối cảnh và lời hứa nội dung giữa tập mới và context cũ.",
                    "Nếu phát hiện thay đổi bất thường, hệ thống tạo cảnh báo continuity.",
                    "Người dùng chọn chấp nhận thay đổi có chủ ý hoặc yêu cầu AI chỉnh lại.",
                    "Kết quả xử lý được lưu vào audit log để theo dõi chất lượng series.",
                ],
                "exceptions": [
                    "Nếu cảnh báo là thay đổi có chủ đích, người dùng có thể đánh dấu bỏ qua.",
                    "Nếu AI không sửa được mâu thuẫn, hệ thống giữ trạng thái NEED_REVIEW.",
                    "Nếu thiếu quyền chỉnh series, người dùng chỉ được xem cảnh báo, không được áp dụng chỉnh sửa.",
                ],
                "result": "Series giảm lỗi sai tên, sai bối cảnh hoặc đứt mạch câu chuyện giữa các tập.",
            },
        ],
    },
    {
        "heading": "US-22: Xem trước và kiểm tra chất lượng video",
        "feature": "Cung cấp màn hình preview để người dùng kiểm tra hình ảnh, âm thanh, phụ đề và lỗi kỹ thuật trước khi duyệt.",
        "scenarios": [
            {
                "title": "Xem trước video MP4 đã render",
                "actor": "Creator.",
                "preconditions": "Video đã render thành công và file MP4 có thể truy cập qua media proxy hoặc storage.",
                "main_flow": [
                    "Người dùng mở màn hình QA Preview từ danh sách video hoặc workflow.",
                    "Trình phát tải video MP4, thumbnail, duration và trạng thái bản render mới nhất.",
                    "Người dùng phát, tua, tạm dừng và kiểm tra phụ đề, giọng đọc, hình ảnh theo từng đoạn.",
                    "Hệ thống hiển thị thông tin version để người dùng biết đang kiểm tra bản render nào.",
                ],
                "exceptions": [
                    "Nếu file MP4 không tồn tại, hệ thống báo lỗi storage và yêu cầu render lại.",
                    "Nếu video đang render, màn hình preview hiển thị trạng thái chờ thay vì phát file cũ sai version.",
                    "Nếu người dùng không sở hữu video, hệ thống từ chối truy cập.",
                ],
                "result": "Người dùng xem được video thật trước khi quyết định duyệt hoặc yêu cầu chỉnh sửa.",
            },
            {
                "title": "Ghi nhận lỗi QA trong quá trình xem trước",
                "actor": "Creator.",
                "preconditions": "Video preview đang mở.",
                "main_flow": [
                    "Người dùng phát hiện lỗi như phụ đề trễ, âm lượng nhỏ, hình ảnh không phù hợp hoặc lỗi chính tả.",
                    "Người dùng ghi nhận lỗi theo mốc thời gian hoặc chọn loại lỗi có sẵn.",
                    "Hệ thống lưu QA issue gắn với video, scene, timestamp và mức độ nghiêm trọng.",
                    "Workflow chuyển sang NEED_FIX nếu lỗi ảnh hưởng khả năng đăng bài.",
                ],
                "exceptions": [
                    "Nếu lỗi chỉ là ghi chú nhẹ, người dùng có thể lưu note mà không đổi trạng thái workflow.",
                    "Nếu timestamp nằm ngoài duration video, hệ thống yêu cầu chọn lại mốc thời gian.",
                    "Nếu video đã duyệt, hệ thống yêu cầu hủy duyệt trước khi ghi lỗi bắt buộc sửa.",
                ],
                "result": "Các lỗi chất lượng được ghi cụ thể, giúp quay lại đúng scene để chỉnh sửa.",
            },
            {
                "title": "Yêu cầu render lại từ màn hình QA",
                "actor": "Creator.",
                "preconditions": "Video có lỗi QA hoặc timeline đã được chỉnh sau khi preview.",
                "main_flow": [
                    "Người dùng chọn yêu cầu render lại sau khi sửa kịch bản, subtitle hoặc media.",
                    "Hệ thống xác nhận các bước phụ thuộc đã sẵn sàng và tạo render job mới.",
                    "Video cũ được giữ làm lịch sử, trạng thái bản hiện hành chuyển sang RENDERING.",
                    "Khi render xong, màn hình QA cập nhật sang bản mới để người dùng kiểm tra lại.",
                ],
                "exceptions": [
                    "Nếu lỗi cần tạo lại voice trước, hệ thống chuyển người dùng về bước TTS.",
                    "Nếu video đang được đăng hoặc đã đăng, hệ thống không cho render đè mà tạo bản mới riêng.",
                    "Nếu render lại thất bại, hệ thống giữ nguyên issue và ghi log nguyên nhân.",
                ],
                "result": "Quy trình QA khép kín từ phát hiện lỗi đến sửa và kiểm tra lại.",
            },
        ],
    },
    {
        "heading": "US-23: Duyệt nội dung trước khi đăng",
        "feature": "Thiết lập cổng phê duyệt để chỉ video đạt yêu cầu mới được lên lịch hoặc đăng ngay.",
        "scenarios": [
            {
                "title": "Duyệt video đạt yêu cầu",
                "actor": "Creator hoặc người có quyền duyệt.",
                "preconditions": "Video đã render xong, đã xem preview và không còn QA issue bắt buộc sửa.",
                "main_flow": [
                    "Người dùng bấm Duyệt video trên màn hình QA.",
                    "Hệ thống kiểm tra trạng thái render, issue QA và quyền sở hữu project.",
                    "Video chuyển sang APPROVED, ghi thời gian duyệt, người duyệt và version được duyệt.",
                    "Giao diện mở các chức năng lên lịch, đăng ngay hoặc đưa vào publishing queue.",
                ],
                "exceptions": [
                    "Nếu video chưa render xong, hệ thống không cho duyệt.",
                    "Nếu còn issue nghiêm trọng, hệ thống yêu cầu xử lý trước khi duyệt.",
                    "Nếu người dùng không có quyền, thao tác bị từ chối và ghi audit log.",
                ],
                "result": "Chỉ bản video đã kiểm tra mới được phép chuyển sang xuất bản.",
            },
            {
                "title": "Từ chối hoặc yêu cầu chỉnh sửa video",
                "actor": "Creator hoặc người kiểm duyệt.",
                "preconditions": "Video đang ở trạng thái RENDERED hoặc WAITING_REVIEW.",
                "main_flow": [
                    "Người dùng chọn Từ chối hoặc Yêu cầu chỉnh sửa.",
                    "Người dùng nhập lý do như sai nội dung, hình ảnh chưa phù hợp, phụ đề lệch hoặc chất lượng âm thanh thấp.",
                    "Hệ thống chuyển trạng thái sang REJECTED hoặc NEED_FIX và lưu lý do.",
                    "Các bước lên lịch và đăng bài bị khóa cho tới khi có bản sửa mới được duyệt.",
                ],
                "exceptions": [
                    "Nếu lý do trống, hệ thống yêu cầu nhập mô tả để đội phát triển hoặc người dùng biết cần sửa gì.",
                    "Nếu video đã được lên lịch, hệ thống tạm dừng queue item liên quan.",
                    "Nếu video đã đăng, thao tác từ chối chỉ được ghi nhận hậu kiểm, không thay đổi bài đã xuất bản.",
                ],
                "result": "Nội dung không đạt không bị xuất bản nhầm và có lý do rõ ràng để chỉnh sửa.",
            },
            {
                "title": "Ghi vết phê duyệt để kiểm soát trách nhiệm",
                "actor": "Hệ thống audit.",
                "preconditions": "Có thao tác duyệt, hủy duyệt hoặc từ chối video.",
                "main_flow": [
                    "Hệ thống ghi audit log gồm người thao tác, thời điểm, trạng thái cũ, trạng thái mới và video version.",
                    "Admin có thể tra cứu lịch sử phê duyệt trong trường hợp cần kiểm tra trách nhiệm.",
                    "Nếu trạng thái duyệt thay đổi, các queue item phụ thuộc được cập nhật theo.",
                ],
                "exceptions": [
                    "Nếu ghi audit log thất bại, hệ thống vẫn không được bỏ qua kiểm tra trạng thái duyệt.",
                    "Nếu có hai thao tác duyệt đồng thời, hệ thống lấy trạng thái cập nhật sau cùng và ghi đủ lịch sử.",
                    "Nếu người dùng mất phiên đăng nhập, thao tác duyệt không được thực hiện.",
                ],
                "result": "Quá trình phê duyệt minh bạch và có thể truy vết.",
            },
        ],
    },
    {
        "heading": "US-04: Kết nối tài khoản TikTok bằng QR",
        "feature": "Kết nối tài khoản TikTok cá nhân bằng QR/OAuth để hệ thống có quyền đăng video theo người dùng.",
        "scenarios": [
            {
                "title": "Tạo phiên QR để kết nối tài khoản TikTok mới",
                "actor": "Creator.",
                "preconditions": "Người dùng đã đăng nhập hệ thống và chưa có TikTok profile active hoặc muốn thêm profile mới.",
                "main_flow": [
                    "Người dùng vào trang Quản lý tài khoản mạng xã hội và chọn kết nối TikTok bằng QR.",
                    "API Service tạo pending session, yêu cầu TikTok cấp mã QR và trả về session_id, qr_url, expires_at.",
                    "Giao diện hiển thị mã QR và bắt đầu polling trạng thái kết nối.",
                    "Người dùng dùng ứng dụng TikTok quét mã và xác nhận quyền truy cập cần thiết.",
                ],
                "exceptions": [
                    "Nếu TikTok không trả được QR, hệ thống hiển thị lỗi và cho phép tạo phiên mới.",
                    "Nếu người dùng chưa đăng nhập, API không tạo session.",
                    "Nếu QR hết hạn, giao diện tự chuyển sang trạng thái expired và yêu cầu tạo mã mới.",
                ],
                "result": "Người dùng có một phiên QR hợp lệ để bắt đầu quá trình liên kết TikTok.",
            },
            {
                "title": "Hoàn tất kết nối và lưu social profile",
                "actor": "TikTok, API Service và Creator.",
                "preconditions": "Người dùng đã quét QR và TikTok trả về mã xác thực hợp lệ.",
                "main_flow": [
                    "Hệ thống nhận callback hoặc polling thấy pending session chuyển sang confirmed.",
                    "API Service đổi mã xác thực lấy access token, refresh token, scope và thông tin kênh.",
                    "Hệ thống lưu social profile theo user_id, platform TikTok, channel_id, avatar, username và trạng thái ACTIVE.",
                    "Giao diện hiển thị tài khoản đã kết nối và cho phép cấu hình chiến lược kênh.",
                ],
                "exceptions": [
                    "Nếu token không có scope đăng video, hệ thống lưu profile ở trạng thái NEED_PERMISSION và yêu cầu kết nối lại.",
                    "Nếu tài khoản TikTok đã tồn tại, hệ thống cập nhật token thay vì tạo bản ghi trùng.",
                    "Nếu lưu database lỗi, hệ thống không báo kết nối thành công để tránh profile giả.",
                ],
                "result": "Hệ thống có tài khoản TikTok hợp lệ gắn với đúng người dùng để đăng bài.",
            },
            {
                "title": "Ngắt kết nối hoặc kết nối lại tài khoản TikTok",
                "actor": "Creator.",
                "preconditions": "Người dùng đã có ít nhất một social profile TikTok.",
                "main_flow": [
                    "Người dùng chọn profile TikTok và bấm ngắt kết nối hoặc kết nối lại.",
                    "Nếu ngắt kết nối, hệ thống chuyển profile sang DISCONNECTED và dừng các queue item chưa xuất bản nếu cần.",
                    "Nếu kết nối lại, hệ thống tạo phiên QR mới và cập nhật token cho profile cũ sau khi xác thực thành công.",
                    "Hệ thống ghi audit log cho thay đổi trạng thái profile.",
                ],
                "exceptions": [
                    "Nếu profile đang có bài chuẩn bị đăng, hệ thống cảnh báo ảnh hưởng tới lịch đăng.",
                    "Nếu người dùng không sở hữu profile, thao tác bị từ chối.",
                    "Nếu kết nối lại thất bại, token cũ không được ghi đè bằng dữ liệu lỗi.",
                ],
                "result": "Người dùng chủ động quản lý trạng thái kết nối TikTok và đảm bảo token luôn hợp lệ.",
            },
        ],
    },
    {
        "heading": "US-05: Cấu hình chiến lược nội dung cho kênh",
        "feature": "Thiết lập định hướng kênh, tần suất, khung giờ và chính sách tự động đăng cho từng tài khoản TikTok.",
        "scenarios": [
            {
                "title": "Tạo chiến lược nội dung cho kênh TikTok",
                "actor": "Creator.",
                "preconditions": "Social profile TikTok ở trạng thái ACTIVE và thuộc sở hữu người dùng.",
                "main_flow": [
                    "Người dùng mở profile TikTok và chọn cấu hình chiến lược nội dung.",
                    "Người dùng nhập chủ đề ưu tiên, chủ đề cần tránh, phong cách giọng điệu, đối tượng khán giả và định dạng video mong muốn.",
                    "Người dùng chọn tần suất đăng, khung giờ đăng và bật/tắt auto_publish_enabled.",
                    "Hệ thống validate cấu hình rồi lưu vào SocialProfileStrategy gắn với profile.",
                ],
                "exceptions": [
                    "Nếu profile chưa active, hệ thống không cho tạo chiến lược.",
                    "Nếu tần suất đăng quá cao hoặc khung giờ không hợp lệ, hệ thống yêu cầu chỉnh lại.",
                    "Nếu chiến lược đã tồn tại, hệ thống cập nhật revision thay vì tạo trùng.",
                ],
                "result": "Mỗi kênh có bộ quy tắc vận hành rõ ràng để AI lập kế hoạch và scheduler đăng bài.",
            },
            {
                "title": "Cập nhật chiến lược sau khi hiệu suất thay đổi",
                "actor": "Creator.",
                "preconditions": "Kênh đã có dữ liệu chiến lược hoặc đã có bài đăng đo được metrics.",
                "main_flow": [
                    "Người dùng xem hiệu suất bài đăng và nhận thấy cần thay đổi chủ đề hoặc khung giờ.",
                    "Người dùng chỉnh strategy như tăng số bài/tuần, đổi giờ vàng hoặc thay đổi tone nội dung.",
                    "Hệ thống lưu bản cập nhật và áp dụng cho các lịch đăng mới.",
                    "Các queue item đã lên lịch trước đó được giữ nguyên hoặc được đề xuất cập nhật tùy lựa chọn người dùng.",
                ],
                "exceptions": [
                    "Nếu thay đổi làm xung đột queue hiện tại, hệ thống cảnh báo danh sách item bị ảnh hưởng.",
                    "Nếu auto_publish_enabled bị tắt, scheduler sẽ bỏ qua các item tự động thuộc profile đó.",
                    "Nếu người dùng thoát khi chưa lưu, cấu hình cũ vẫn được giữ.",
                ],
                "result": "Chiến lược kênh có thể thích nghi với dữ liệu thực tế mà không làm rối hàng đợi hiện có.",
            },
            {
                "title": "Phân tách chiến lược theo từng tài khoản sở hữu",
                "actor": "Hệ thống phân quyền.",
                "preconditions": "Hệ thống có nhiều người dùng hoặc nhiều social profile.",
                "main_flow": [
                    "Người dùng chỉ nhìn thấy strategy của các profile thuộc quyền sở hữu của mình.",
                    "API Service lọc dữ liệu theo current_user.id trước khi trả danh sách strategy.",
                    "Admin có thể xem tổng quan để hỗ trợ nhưng thao tác chỉnh sửa vẫn được kiểm soát bằng quyền.",
                ],
                "exceptions": [
                    "Nếu người dùng cố truy cập strategy_id không thuộc sở hữu, API trả lỗi không có quyền.",
                    "Nếu profile bị xóa hoặc disconnected, strategy liên quan không được dùng cho scheduler.",
                    "Nếu dữ liệu strategy thiếu profile_id, hệ thống đánh dấu lỗi dữ liệu để admin xử lý.",
                ],
                "result": "Dữ liệu chiến lược không bị lẫn giữa các người dùng và kênh TikTok.",
            },
        ],
    },
    {
        "heading": "US-24: Lên lịch đăng bài theo chiến lược kênh",
        "feature": "Đưa video đã duyệt vào publishing queue với thời điểm đăng phù hợp chiến lược kênh.",
        "scenarios": [
            {
                "title": "Lên lịch một video đã duyệt",
                "actor": "Creator.",
                "preconditions": "Video ở trạng thái APPROVED, có file MP4 hợp lệ và có ít nhất một social profile ACTIVE.",
                "main_flow": [
                    "Người dùng chọn video, kênh TikTok và thời điểm đăng mong muốn.",
                    "Hệ thống kiểm tra quyền với video và profile, đồng thời kiểm tra video version đã duyệt.",
                    "Publishing queue item được tạo với trạng thái QUEUED, scheduled_at, profile_id, video_id và caption.",
                    "Giao diện hiển thị item mới trong hàng đợi theo đúng thời gian đăng.",
                ],
                "exceptions": [
                    "Nếu video chưa duyệt, hệ thống chặn lên lịch.",
                    "Nếu profile mất kết nối hoặc thiếu token, hệ thống yêu cầu kết nối lại TikTok.",
                    "Nếu thời điểm đăng nằm trong quá khứ, hệ thống yêu cầu chọn thời gian mới hoặc dùng chức năng đăng ngay.",
                ],
                "result": "Video được xếp lịch rõ ràng và có trạng thái theo dõi trong publishing queue.",
            },
            {
                "title": "Gợi ý thời gian đăng theo chiến lược kênh",
                "actor": "Hệ thống scheduler hỗ trợ.",
                "preconditions": "Profile có SocialProfileStrategy chứa tần suất và khung giờ ưu tiên.",
                "main_flow": [
                    "Người dùng mở form lên lịch cho video đã duyệt.",
                    "Hệ thống đọc strategy của profile và đề xuất khung giờ phù hợp gần nhất.",
                    "Người dùng chọn một khung giờ đề xuất hoặc nhập thời điểm thủ công.",
                    "Hệ thống lưu lựa chọn cuối cùng vào queue item.",
                ],
                "exceptions": [
                    "Nếu strategy chưa cấu hình, hệ thống cho phép nhập thủ công và gợi ý tạo strategy.",
                    "Nếu khung giờ đề xuất đã có bài khác, hệ thống gợi ý slot tiếp theo.",
                    "Nếu người dùng chọn giờ ngoài strategy, hệ thống cảnh báo nhưng vẫn cho lưu nếu hợp lệ.",
                ],
                "result": "Lịch đăng bám sát chiến lược kênh nhưng vẫn linh hoạt cho người dùng.",
            },
            {
                "title": "Cập nhật hoặc hủy lịch đăng",
                "actor": "Creator.",
                "preconditions": "Queue item đang ở trạng thái QUEUED hoặc FAILED chưa đăng thành công.",
                "main_flow": [
                    "Người dùng mở chi tiết item trong hàng đợi.",
                    "Người dùng thay đổi scheduled_at, caption, profile hoặc chọn hủy lịch.",
                    "Hệ thống kiểm tra item chưa ở trạng thái PUBLISHING/PUBLISHED rồi cập nhật dữ liệu.",
                    "Queue hiển thị trạng thái mới và ghi audit log thay đổi lịch.",
                ],
                "exceptions": [
                    "Nếu item đang PUBLISHING, hệ thống không cho chỉnh để tránh xung đột với scheduler.",
                    "Nếu item đã PUBLISHED, người dùng chỉ được xem lịch sử, không được hủy.",
                    "Nếu thay đổi profile làm mất quyền đăng, hệ thống từ chối lưu.",
                ],
                "result": "Người dùng kiểm soát được lịch đăng trước thời điểm xuất bản.",
            },
        ],
    },
    {
        "heading": "US-25: Theo dõi trạng thái hàng đợi đăng bài",
        "feature": "Hiển thị và cập nhật trạng thái của các queue item: chờ đăng, đang đăng, đã đăng, thất bại hoặc bị hủy.",
        "scenarios": [
            {
                "title": "Xem danh sách hàng đợi theo người dùng và profile",
                "actor": "Creator.",
                "preconditions": "Người dùng đã có video được lên lịch hoặc đăng thủ công.",
                "main_flow": [
                    "Người dùng mở màn hình Hàng đợi đăng bài.",
                    "API Service trả danh sách item thuộc user hiện tại, có bộ lọc theo profile, trạng thái và khoảng thời gian.",
                    "Giao diện hiển thị tiêu đề video, profile, scheduled_at, trạng thái, caption và hành động phù hợp.",
                    "Người dùng có thể mở chi tiết item để xem log đăng bài.",
                ],
                "exceptions": [
                    "Nếu không có item nào, hệ thống hiển thị trạng thái rỗng và gợi ý lên lịch video đã duyệt.",
                    "Nếu người dùng yêu cầu profile không thuộc sở hữu, API trả lỗi không có quyền.",
                    "Nếu dữ liệu nhiều, hệ thống phân trang để không làm chậm giao diện.",
                ],
                "result": "Người dùng nắm được toàn bộ nội dung đang chờ hoặc đã xử lý trong hàng đợi của mình.",
            },
            {
                "title": "Cập nhật trạng thái thời gian thực khi scheduler xử lý",
                "actor": "Scheduler, API Service và Creator.",
                "preconditions": "Có queue item đến giờ đăng hoặc người dùng đang theo dõi màn hình queue.",
                "main_flow": [
                    "Scheduler chọn item đến hạn và chuyển trạng thái từ QUEUED sang PUBLISHING.",
                    "Giao diện cập nhật trạng thái bằng polling hoặc cơ chế realtime nếu có.",
                    "Sau khi TikTok trả kết quả, item chuyển sang PUBLISHED hoặc FAILED.",
                    "Người dùng thấy thời điểm xử lý, link bài đăng nếu thành công hoặc thông báo lỗi nếu thất bại.",
                ],
                "exceptions": [
                    "Nếu mất kết nối giao diện, lần tải lại sau vẫn hiển thị trạng thái mới nhất từ database.",
                    "Nếu scheduler bị dừng, item vẫn ở QUEUED và không bị mất lịch.",
                    "Nếu trạng thái bị cập nhật đồng thời, hệ thống sử dụng khóa xử lý để tránh đăng trùng.",
                ],
                "result": "Trạng thái queue phản ánh đúng vòng đời đăng bài và giúp người dùng yên tâm theo dõi.",
            },
            {
                "title": "Xử lý item thất bại và retry",
                "actor": "Creator hoặc Admin.",
                "preconditions": "Queue item ở trạng thái FAILED và có failure_reason.",
                "main_flow": [
                    "Người dùng mở item thất bại để xem nguyên nhân như token hết hạn, TikTok lỗi, file MP4 không hợp lệ hoặc caption sai định dạng.",
                    "Người dùng sửa lỗi tương ứng, ví dụ kết nối lại TikTok hoặc render lại video.",
                    "Người dùng bấm retry hoặc scheduler retry theo chính sách nếu còn lượt.",
                    "Hệ thống tạo lần thử mới, tăng retry_count và cập nhật trạng thái.",
                ],
                "exceptions": [
                    "Nếu lỗi không thể retry như video bị xóa, hệ thống yêu cầu tạo queue item mới.",
                    "Nếu vượt quá số lần retry, item giữ FAILED và cần xử lý thủ công.",
                    "Nếu token thiếu quyền, hệ thống chuyển người dùng sang màn hình kết nối TikTok.",
                ],
                "result": "Các lỗi đăng bài có đường xử lý rõ ràng thay vì chỉ hiển thị thất bại chung chung.",
            },
        ],
    },
    {
        "heading": "US-26: Theo dõi hiệu suất bài đăng trên mạng xã hội",
        "feature": "Thu thập và hiển thị chỉ số hiệu suất bài đăng TikTok để đánh giá hiệu quả nội dung và cải thiện chiến lược.",
        "scenarios": [
            {
                "title": "Thu thập metrics sau khi bài đăng thành công",
                "actor": "Metrics Worker hoặc API Service.",
                "preconditions": "Queue item đã PUBLISHED, có platform_post_id hoặc link bài đăng TikTok.",
                "main_flow": [
                    "Hệ thống lên lịch lấy metrics sau khi bài đăng đã xuất bản một khoảng thời gian.",
                    "Worker gọi API nền tảng hoặc đọc dữ liệu có sẵn để lấy views, likes, comments, shares và thời điểm cập nhật.",
                    "Metrics được lưu vào SocialPostMetrics gắn với bài đăng, profile và video.",
                    "Giao diện hiển thị lần cập nhật gần nhất để người dùng biết dữ liệu có mới hay không.",
                ],
                "exceptions": [
                    "Nếu TikTok chưa trả metrics, hệ thống lưu trạng thái pending và thử lại sau.",
                    "Nếu token hết hạn, hệ thống đánh dấu NEED_RECONNECT cho profile.",
                    "Nếu bài đăng bị xóa trên TikTok, hệ thống ghi trạng thái unavailable thay vì xóa dữ liệu lịch sử.",
                ],
                "result": "Mỗi bài đã đăng có dữ liệu đo lường để phục vụ phân tích hiệu quả.",
            },
            {
                "title": "Xem dashboard hiệu suất bài đăng",
                "actor": "Creator.",
                "preconditions": "Người dùng có ít nhất một bài đăng đã có metrics.",
                "main_flow": [
                    "Người dùng mở trang thống kê bài đăng.",
                    "Hệ thống hiển thị danh sách bài, biểu đồ xu hướng và các chỉ số tổng hợp theo profile hoặc khoảng thời gian.",
                    "Người dùng lọc theo kênh, chủ đề, ngày đăng hoặc trạng thái để so sánh.",
                    "Người dùng mở chi tiết bài để xem liên kết với video, caption và lịch sử metrics.",
                ],
                "exceptions": [
                    "Nếu chưa có dữ liệu, hệ thống hiển thị hướng dẫn ngắn rằng cần đăng bài trước khi có metrics.",
                    "Nếu metrics thiếu một chỉ số, giao diện hiển thị N/A thay vì tính sai.",
                    "Nếu dữ liệu thuộc profile khác, API không trả về cho người dùng hiện tại.",
                ],
                "result": "Người dùng đánh giá được video nào hoạt động tốt để điều chỉnh chiến lược kênh.",
            },
            {
                "title": "Sử dụng metrics để cải thiện chiến lược nội dung",
                "actor": "Creator.",
                "preconditions": "Dashboard có dữ liệu từ nhiều bài đăng hoặc nhiều khung giờ.",
                "main_flow": [
                    "Người dùng so sánh hiệu suất theo chủ đề, giờ đăng và định dạng video.",
                    "Hệ thống chỉ ra bài có performance cao và các nhóm nội dung nên ưu tiên.",
                    "Người dùng cập nhật SocialProfileStrategy dựa trên kết quả phân tích.",
                    "Các lịch đăng mới sử dụng strategy đã cập nhật.",
                ],
                "exceptions": [
                    "Nếu dữ liệu quá ít, hệ thống cảnh báo chưa đủ cơ sở kết luận.",
                    "Nếu hiệu suất bất thường do lỗi dữ liệu, admin có thể kiểm tra log metrics.",
                    "Nếu người dùng không muốn thay đổi strategy, hệ thống chỉ giữ metrics làm tham khảo.",
                ],
                "result": "Metrics không chỉ để xem mà còn quay lại cải thiện quy trình lập kế hoạch và đăng bài.",
            },
        ],
    },
    {
        "heading": "US-27: Đăng bài thủ công ngay lập tức lên mạng xã hội TikTok",
        "feature": "Cho phép người dùng chủ động xuất bản ngay một video đã duyệt lên TikTok mà không chờ scheduler.",
        "scenarios": [
            {
                "title": "Đăng ngay video đã duyệt",
                "actor": "Creator.",
                "preconditions": "Video đã APPROVED, có file MP4 hợp lệ, profile TikTok ACTIVE và token có scope đăng video.",
                "main_flow": [
                    "Người dùng mở video hoặc queue item và bấm Đăng ngay.",
                    "Hệ thống xác nhận thao tác vì bài sẽ được gửi trực tiếp lên TikTok.",
                    "API Service kiểm tra quyền, trạng thái duyệt, trạng thái token và file MP4.",
                    "Hệ thống gọi TikTok API để upload/publish video, sau đó tạo SocialPost và cập nhật trạng thái PUBLISHED.",
                ],
                "exceptions": [
                    "Nếu video chưa APPROVED, nút đăng ngay bị vô hiệu hoặc API trả lỗi.",
                    "Nếu token hết hạn, hệ thống yêu cầu kết nối lại TikTok trước khi đăng.",
                    "Nếu TikTok trả lỗi caption hoặc file, item chuyển FAILED và hiển thị nguyên nhân.",
                ],
                "result": "Người dùng có thể xuất bản nội dung gấp mà vẫn tuân thủ kiểm duyệt và phân quyền.",
            },
            {
                "title": "Đăng ngay từ item đang ở hàng đợi",
                "actor": "Creator.",
                "preconditions": "Queue item đang ở trạng thái QUEUED, chưa đến giờ hoặc người dùng muốn đăng sớm.",
                "main_flow": [
                    "Người dùng chọn item trong hàng đợi và bấm Publish Now.",
                    "Hệ thống khóa item để scheduler không xử lý song song.",
                    "Item chuyển sang PUBLISHING và được gửi ngay tới TikTok API.",
                    "Nếu thành công, scheduled_at cũ được giữ làm dữ liệu lịch sử nhưng trạng thái chuyển PUBLISHED.",
                ],
                "exceptions": [
                    "Nếu scheduler vừa lấy item xử lý, hệ thống báo item đang PUBLISHING và không tạo lần đăng mới.",
                    "Nếu item đã PUBLISHED, hệ thống không cho đăng lại để tránh duplicate.",
                    "Nếu người dùng hủy ở bước xác nhận, trạng thái vẫn là QUEUED.",
                ],
                "result": "Một item đã lên lịch có thể được xuất bản sớm an toàn, không bị đăng trùng.",
            },
            {
                "title": "Ghi nhận lỗi đăng thủ công",
                "actor": "Hệ thống publishing.",
                "preconditions": "Người dùng đã kích hoạt đăng ngay nhưng quá trình gọi TikTok không thành công.",
                "main_flow": [
                    "Hệ thống bắt lỗi từ TikTok API, storage hoặc kiểm tra quyền.",
                    "Queue item hoặc bản ghi publish chuyển sang FAILED, lưu failure_reason và thời điểm thất bại.",
                    "Giao diện hiển thị lý do dễ hiểu và hành động tiếp theo như retry, kết nối lại TikTok hoặc render lại.",
                    "Audit log ghi nhận người kích hoạt đăng thủ công để phục vụ truy vết.",
                ],
                "exceptions": [
                    "Nếu lỗi chỉ là tạm thời, hệ thống cho phép retry ngay.",
                    "Nếu lỗi do thiếu quyền, retry bị chặn cho tới khi profile được cấp quyền lại.",
                    "Nếu không xác định được lỗi, hệ thống lưu log kỹ thuật cho admin.",
                ],
                "result": "Đăng thủ công thất bại vẫn có thông tin xử lý rõ ràng và không làm mất trạng thái video.",
            },
        ],
    },
    {
        "heading": "US-28: Tự động đăng bài theo lịch trình lên mạng xã hội TikTok",
        "feature": "Background Scheduler tự động quét publishing queue và xuất bản video đúng lịch khi đủ điều kiện.",
        "scenarios": [
            {
                "title": "Scheduler tự động đăng item đến hạn",
                "actor": "Background Scheduler.",
                "preconditions": "ENABLE_SCHEDULER được bật, queue item ở trạng thái QUEUED, scheduled_at nhỏ hơn hoặc bằng thời điểm hiện tại, video APPROVED và profile TikTok ACTIVE.",
                "main_flow": [
                    "Scheduler chạy định kỳ theo cấu hình SCHEDULER_POLL_SECONDS.",
                    "Scheduler truy vấn các queue item đến hạn và khóa từng item để xử lý độc quyền.",
                    "Hệ thống kiểm tra lại video, profile, token, strategy và trạng thái auto_publish_enabled.",
                    "Nếu đủ điều kiện, hệ thống upload/publish video lên TikTok và cập nhật item thành PUBLISHED.",
                ],
                "exceptions": [
                    "Nếu item chưa đến giờ, scheduler bỏ qua.",
                    "Nếu video bị hủy duyệt sau khi lên lịch, scheduler không đăng và chuyển item sang BLOCKED hoặc FAILED tùy cấu hình.",
                    "Nếu TikTok API lỗi tạm thời, item được ghi failure_reason và có thể retry theo chính sách.",
                ],
                "result": "Video được đăng tự động đúng lịch mà không cần người dùng thao tác thủ công.",
            },
            {
                "title": "Scheduler tôn trọng chiến lược kênh và công tắc auto publish",
                "actor": "Background Scheduler và SocialProfileStrategy.",
                "preconditions": "Profile có chiến lược kênh và queue item thuộc profile đó.",
                "main_flow": [
                    "Scheduler đọc SocialProfileStrategy của profile trước khi đăng.",
                    "Nếu auto_publish_enabled bật và lịch đăng hợp lệ, scheduler tiếp tục xử lý.",
                    "Nếu auto_publish_enabled tắt, scheduler bỏ qua item hoặc giữ QUEUED để người dùng đăng thủ công.",
                    "Hệ thống ghi log lý do bỏ qua để người dùng hiểu vì sao bài chưa đăng.",
                ],
                "exceptions": [
                    "Nếu strategy bị xóa, scheduler dùng chính sách an toàn là không tự đăng.",
                    "Nếu schedule_enabled tạm tắt ở cấp hệ thống, scheduler không xử lý bất kỳ item nào.",
                    "Nếu profile disconnected, item được đánh dấu cần kết nối lại.",
                ],
                "result": "Tự động đăng chỉ diễn ra khi người dùng và hệ thống đều cho phép.",
            },
            {
                "title": "Xử lý retry và chống đăng trùng trong auto publish",
                "actor": "Scheduler và hệ thống queue.",
                "preconditions": "Có item đến hạn đang hoặc từng được scheduler xử lý.",
                "main_flow": [
                    "Trước khi gọi TikTok, scheduler đặt trạng thái PUBLISHING và ghi locked_at.",
                    "Nếu publish thành công, hệ thống lưu platform_post_id và không cho xử lý lại item đó.",
                    "Nếu publish thất bại, hệ thống tăng retry_count và quyết định giữ FAILED hoặc đưa lại QUEUED theo chính sách.",
                    "Scheduler ở lần chạy sau chỉ chọn item đủ điều kiện retry và chưa vượt giới hạn.",
                ],
                "exceptions": [
                    "Nếu worker bị dừng giữa chừng, item quá thời gian khóa có thể được giải phóng theo cơ chế recovery.",
                    "Nếu TikTok trả về kết quả không chắc chắn, hệ thống kiểm tra platform_post_id trước khi retry để tránh đăng hai lần.",
                    "Nếu retry vượt giới hạn, item giữ FAILED và yêu cầu người dùng xử lý thủ công.",
                ],
                "result": "Auto scheduler vận hành bền vững, hạn chế mất bài và đặc biệt tránh đăng trùng.",
            },
            {
                "title": "Theo dõi lịch sử chạy scheduler",
                "actor": "Admin và hệ thống vận hành.",
                "preconditions": "Scheduler đã chạy ít nhất một chu kỳ.",
                "main_flow": [
                    "Hệ thống ghi lại số item quét được, số item đăng thành công, số item thất bại và thời gian chạy.",
                    "Admin xem log để kiểm tra scheduler có hoạt động ổn định hay không.",
                    "Nếu số lỗi tăng bất thường, admin có thể kiểm tra token, TikTok API, storage hoặc worker.",
                    "Các thông tin vận hành được dùng để điều chỉnh cấu hình poll interval và retry.",
                ],
                "exceptions": [
                    "Nếu không có item đến hạn, scheduler vẫn ghi heartbeat tối thiểu để biết tiến trình còn sống.",
                    "Nếu log database lỗi, scheduler vẫn ưu tiên không đăng trùng và báo lỗi vận hành.",
                    "Nếu scheduler bị tắt chủ động, hệ thống hiển thị trạng thái paused.",
                ],
                "result": "Người vận hành có dữ liệu kiểm soát auto publish thay vì phụ thuộc vào suy đoán.",
            },
        ],
    },
]


def main():
    doc = Document(INPUT_DOCX)
    anchor = remove_blocks_between(doc, SECTION_HEADING, NEXT_HEADING_PREFIX)

    add_plain(anchor, "Các kịch bản câu chuyện người dùng của Release V2.0 được mô tả chi tiết theo từng nghiệp vụ chính. Mỗi kịch bản làm rõ tác nhân, tiền điều kiện, luồng xử lý, tình huống ngoại lệ và kết quả kỳ vọng để bảo đảm nhóm phát triển có căn cứ triển khai, kiểm thử và nghiệm thu.")
    add_plain(anchor, "")

    for story in USER_STORIES:
        add_us_heading(anchor, story["heading"])
        add_feature(anchor, story["feature"])
        for scenario in story["scenarios"]:
            add_scenario(anchor, **scenario)
            add_plain(anchor, "")

    doc.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")
    print(f"Stories: {len(USER_STORIES)}")
    print(f"Scenarios: {sum(len(s['scenarios']) for s in USER_STORIES)}")


if __name__ == "__main__":
    main()
