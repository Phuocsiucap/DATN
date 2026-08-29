import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc.docx")
SOURCE_MD = Path(r"D:\DATN\tài liệu\CHUONG_2_HOAN_THIEN.md")
OUTPUT_DOCX = Path(r"D:\DATN\tài liệu\DATN_BC_NguyenVanPhuoc_cap_nhat_chuong_2.docx")
IMAGE_DIR = Path(r"D:\DATN\tài liệu")


SPRINT_PLANNING_MD = """| **Sprint** | **Mục tiêu** | **Ngày bắt đầu** | **Ngày kết thúc** | **Tiến độ** |
|---|---|---|---|---|
| Sprint 1 | Hoàn thành nền tảng xác thực, phân quyền, đăng xuất, thu thập dữ liệu, chuẩn hóa dữ liệu và nhóm nội dung | 20/07/2026 | 09/08/2026 | 100% |
| Sprint 2 | Hoàn thành các chức năng AI chọn chủ đề, đề xuất kiểu video, chia nội dung nhiều tập và duyệt kế hoạch | 10/08/2026 | 23/08/2026 | 100% |

| **Tiêu đề** | **Mục tiêu** | **Người thực hiện** | **Estimate Effort (h)** | **Ngày bắt đầu** | **Ngày kết thúc** | **Priority** |
|---|---|---|---:|---|---|---|
| US-01: Đăng nhập vào hệ thống | Người dùng đăng nhập để sử dụng chức năng theo quyền hạn | Nguyễn Văn Phước | 8 | 20/07/2026 | 26/07/2026 | High |
| US-02: Kiểm soát phạm vi truy cập dữ liệu | Người dùng được gán vai trò phù hợp và chỉ sử dụng chức năng theo quyền hạn | Nguyễn Văn Phước | 8 | 20/07/2026 | 26/07/2026 | Medium |
| US-03: Đăng xuất khỏi hệ thống | Người dùng hủy phiên đăng nhập an toàn khỏi hệ thống | Nguyễn Văn Phước | 4 | 20/07/2026 | 26/07/2026 | High |
| US-06: Tạo job lấy dữ liệu | Người dùng tạo được private crawl job theo nguồn và từ khóa | Nguyễn Văn Phước | 12 | 27/07/2026 | 02/08/2026 | High |
| US-07: Theo dõi tiến độ crawl | Người dùng xem được trạng thái, tiến độ và log crawl của job thuộc sở hữu | Nguyễn Văn Phước | 12 | 27/07/2026 | 02/08/2026 | High |
| US-08: Lấy dữ liệu VNExpress | Hệ thống lấy được bài viết, metadata và hình ảnh từ VNExpress | Nguyễn Văn Phước | 12 | 27/07/2026 | 02/08/2026 | High |
| US-09: Chuẩn hóa dữ liệu thô | Dữ liệu crawl được làm sạch và chuyển thành dữ liệu chuẩn | Nguyễn Văn Phước | 20 | 03/08/2026 | 09/08/2026 | High |
| US-10: Nhóm nội dung theo truyện/tập | Nội dung được sắp xếp đúng thứ tự để chuẩn bị cho AI planning | Nguyễn Văn Phước | 20 | 03/08/2026 | 09/08/2026 | Medium |
| US-11: Quality Score & Deduplication | Hệ thống phát hiện trùng lặp hình ảnh và nội dung bằng fingerprint | Nguyễn Văn Phước | 20 | 03/08/2026 | 09/08/2026 | High |
| US-12: Tạo content project | Người dùng tạo được dự án nội dung từ dữ liệu đã crawl | Nguyễn Văn Phước | 12 | 10/08/2026 | 16/08/2026 | High |
| US-13: AI đề xuất chủ đề | AI đề xuất chủ đề, góc triển khai và kiểu video phù hợp | Nguyễn Văn Phước | 20 | 10/08/2026 | 16/08/2026 | High |
| US-14: AI chia nội dung thành nhiều phần | Kế hoạch chuỗi video có mạch truyện rõ ràng giữa các phần | Nguyễn Văn Phước | 32 | 17/08/2026 | 23/08/2026 | High |
| US-15: Duyệt hoặc chỉnh sửa kế hoạch | Người dùng duyệt, từ chối hoặc yêu cầu tạo lại kế hoạch | Nguyễn Văn Phước | 12 | 17/08/2026 | 23/08/2026 | Medium |"""


ENHANCED_2512 = """#### 2.5.1.2. Xây dựng kịch bản câu chuyện người dùng

**US-01: Đăng nhập vào hệ thống**

Feature: Xác thực người dùng

Scenario 1: Đăng nhập thành công bằng tài khoản hợp lệ
Given người dùng đã có tài khoản hợp lệ và đang ở màn hình đăng nhập
When người dùng nhập đúng email, mật khẩu và chọn nút Đăng nhập
Then hệ thống xác thực thông tin đăng nhập, cấp JWT access token và điều hướng người dùng vào dashboard theo đúng vai trò

Scenario 2: Đăng nhập thất bại do sai thông tin
Given người dùng đang ở màn hình đăng nhập
When người dùng nhập sai email hoặc mật khẩu
Then hệ thống không cấp phiên đăng nhập và hiển thị thông báo tài khoản hoặc mật khẩu không chính xác

Scenario 3: Truy cập chức năng khi chưa đăng nhập
Given người dùng chưa có access token hợp lệ
When người dùng truy cập màn hình yêu cầu xác thực
Then hệ thống chặn truy cập và chuyển người dùng về màn hình đăng nhập

**US-02: Kiểm soát phạm vi truy cập dữ liệu**

Feature: Phân quyền và kiểm soát phạm vi dữ liệu

Scenario 1: Người dùng CREATOR chỉ xem dữ liệu cá nhân
Given người dùng đăng nhập với vai trò CREATOR
When người dùng mở danh sách crawl job, content project hoặc kế hoạch AI
Then hệ thống chỉ hiển thị dữ liệu do chính người dùng đó tạo và không trả về dữ liệu thuộc tài khoản khác

Scenario 2: Quản trị viên SYSTEM_ADMIN xem dữ liệu toàn hệ thống
Given quản trị viên đăng nhập với vai trò SYSTEM_ADMIN
When quản trị viên truy cập trang quản trị hoặc báo cáo hệ thống
Then hệ thống cho phép xem toàn bộ người dùng, job, dữ liệu, nhật ký xử lý và ghi nhận thao tác quản trị vào audit log

Scenario 3: Từ chối truy cập vượt quyền
Given người dùng không có quyền với tài nguyên được yêu cầu
When người dùng gọi API hoặc mở màn hình ngoài phạm vi quyền hạn
Then hệ thống phản hồi 403 Forbidden và hiển thị thông báo không có quyền truy cập

**US-03: Đăng xuất khỏi hệ thống**

Feature: Đăng xuất và kết thúc phiên làm việc

Scenario 1: Đăng xuất thành công
Given người dùng đã đăng nhập vào hệ thống
When người dùng chọn chức năng Đăng xuất
Then hệ thống xóa access token khỏi phía client, kết thúc phiên làm việc hiện tại và chuyển người dùng về màn hình đăng nhập

Scenario 2: Gửi request sau khi đã đăng xuất
Given người dùng đã thực hiện đăng xuất
When người dùng gửi request tới API yêu cầu xác thực
Then hệ thống phản hồi 401 Unauthorized và yêu cầu người dùng đăng nhập lại

**US-06: Tạo job lấy dữ liệu**

Feature: Tạo private crawl job

Scenario 1: Hiển thị form tạo job crawl
Given người dùng đã đăng nhập
When người dùng truy cập màn hình thu thập dữ liệu
Then hệ thống hiển thị form gồm nguồn dữ liệu, từ khóa hoặc URL, giới hạn crawl và tùy chọn xử lý

Scenario 2: Tạo job crawl VNExpress thành công
Given người dùng nhập nguồn VNExpress và tham số hợp lệ
When người dùng chọn Tạo job
Then hệ thống tạo crawl job ở trạng thái PENDING, gắn owner_id của người dùng và phát sự kiện job_created qua Kafka

Scenario 3: Không tạo job khi thiếu dữ liệu bắt buộc
Given người dùng đang ở form tạo job
When người dùng không nhập nguồn, URL hoặc từ khóa bắt buộc
Then hệ thống không tạo job và hiển thị thông báo yêu cầu bổ sung dữ liệu đầu vào

**US-07: Theo dõi tiến độ crawl**

Feature: Giám sát trạng thái và log crawl job

Scenario 1: Xem danh sách job thuộc sở hữu
Given người dùng đã tạo một hoặc nhiều crawl job
When người dùng mở màn hình Crawl Jobs
Then hệ thống hiển thị danh sách job thuộc sở hữu của người dùng cùng trạng thái PENDING, RUNNING, COMPLETED hoặc FAILED

Scenario 2: Theo dõi tiến độ xử lý real-time
Given crawler đang xử lý một job
When hệ thống cập nhật số lượng bài viết đã phát hiện, đã tải và đã chuẩn hóa
Then giao diện hiển thị phần trăm tiến độ mới nhất và cập nhật trạng thái job mà không cần tải lại trang

Scenario 3: Xem nhật ký lỗi khi crawl thất bại
Given crawl job chuyển sang trạng thái FAILED
When người dùng mở chi tiết job
Then hệ thống hiển thị log lỗi, bước lỗi, thời điểm xảy ra lỗi và giữ dữ liệu đã xử lý được để phục vụ chạy lại

**US-08: Lấy dữ liệu VNExpress**

Feature: Thu thập bài viết, metadata và hình ảnh từ VNExpress

Scenario 1: Crawl bài viết VNExpress bằng URL hợp lệ
Given crawler nhận được job VNExpress từ Kafka
When crawler tải nội dung bài viết từ URL hợp lệ
Then hệ thống trích xuất tiêu đề, mô tả, nội dung chính, chuyên mục, thời gian đăng và danh sách hình ảnh

Scenario 2: Crawl theo từ khóa hoặc chuyên mục
Given job crawl có cấu hình từ khóa hoặc chuyên mục VNExpress
When crawler tìm thấy danh sách bài viết phù hợp
Then hệ thống tạo danh sách bản ghi thô, gắn metadata nguồn, URL gốc và thời điểm thu thập

Scenario 3: Xử lý bài viết thiếu ảnh hoặc thiếu metadata
Given một bài viết VNExpress thiếu ảnh đại diện hoặc thiếu mô tả
When crawler hoàn tất trích xuất dữ liệu
Then hệ thống vẫn lưu nội dung hợp lệ và đánh dấu các trường thiếu để bước chuẩn hóa chấm điểm chất lượng

**US-09: Chuẩn hóa dữ liệu thô**

Feature: Làm sạch và chuẩn hóa dữ liệu crawl

Scenario 1: Làm sạch nội dung HTML và ký tự nhiễu
Given dữ liệu thô đã được lưu trong MongoDB
When Normalization Service xử lý bản ghi
Then hệ thống loại bỏ thẻ HTML, ký tự rác, khoảng trắng dư, nội dung quảng cáo và tạo văn bản sạch lưu vào PostgreSQL

Scenario 2: Chuẩn hóa metadata bài viết
Given bản ghi thô có tiêu đề, URL, thời gian đăng và danh sách ảnh
When hệ thống chuyển đổi sang mô hình dữ liệu chuẩn
Then hệ thống chuẩn hóa tiêu đề, slug, source_url, published_at, image_urls và liên kết với crawl job gốc

Scenario 3: Đánh dấu dữ liệu cần kiểm tra
Given dữ liệu thô thiếu tiêu đề, thiếu nội dung hoặc nội dung quá ngắn
When hệ thống không đạt ngưỡng dữ liệu tối thiểu
Then hệ thống đặt trạng thái NEEDS_REVIEW và ghi lý do để người dùng hoặc admin kiểm tra lại

**US-10: Nhóm nội dung theo truyện/tập**

Feature: Gom nhóm nội dung theo series và thứ tự tập

Scenario 1: Nhận diện các bài viết thuộc cùng một truyện
Given hệ thống có nhiều bài viết đã chuẩn hóa
When Story Processing Service phân tích tiêu đề, URL và nội dung
Then hệ thống gom các bài viết liên quan vào cùng một content series và lưu quan hệ giữa series với từng tập

Scenario 2: Sắp xếp tập theo đúng thứ tự
Given một series có nhiều tập được phát hiện
When hệ thống trích xuất số tập hoặc thứ tự xuất hiện từ metadata
Then hệ thống sắp xếp danh sách tập theo trình tự nội dung và đánh dấu tập đầu, tập tiếp theo, tập cuối nếu xác định được

Scenario 3: Xử lý trường hợp thiếu số tập
Given một số bài viết không có số tập rõ ràng
When hệ thống không chắc chắn về thứ tự
Then hệ thống giữ dữ liệu trong series và gắn cờ NEEDS_REVIEW_ORDER để tránh tạo kế hoạch sai mạch truyện

**US-11: Quality Score & Deduplication**

Feature: Chấm điểm chất lượng và phát hiện trùng lặp dữ liệu

Scenario 1: Tính điểm chất lượng nội dung
Given bài viết đã được chuẩn hóa
When hệ thống kiểm tra độ dài nội dung, số lượng ảnh, metadata bắt buộc và trạng thái xử lý
Then hệ thống tính quality_score cho từng bản ghi và phân loại dữ liệu thành READY hoặc NEEDS_REVIEW

Scenario 2: Phát hiện trùng lặp nội dung văn bản
Given hai bài viết có nội dung tương đồng nhưng khác URL
When hệ thống tạo text fingerprint từ tiêu đề chuẩn hóa và phần nội dung chính
Then hệ thống đánh dấu các bản ghi có fingerprint trùng hoặc gần trùng và tăng duplicate_count

Scenario 3: Phát hiện trùng lặp hình ảnh
Given hai bài viết sử dụng cùng ảnh nhưng khác subdomain CDN
When hệ thống áp dụng image fingerprinting cho danh sách ảnh
Then hệ thống nhận diện ảnh trùng lặp và tránh chọn nhiều dữ liệu giống nhau cho AI planning

**US-12: Tạo content project**

Feature: Khởi tạo dự án nội dung từ dữ liệu đã crawl

Scenario 1: Tạo content project từ bài viết READY
Given người dùng có dữ liệu đã chuẩn hóa ở trạng thái READY
When người dùng chọn bài viết và tạo dự án nội dung
Then hệ thống tạo content project mới ở trạng thái DRAFT và liên kết project với dữ liệu nguồn, người sở hữu

Scenario 2: Tạo project từ series nhiều tập
Given người dùng chọn một content series đã được gom nhóm
When người dùng yêu cầu tạo project
Then hệ thống khởi tạo project chứa danh sách tập theo đúng thứ tự và lưu context series để chuyển sang AI planning

Scenario 3: Chặn tạo project từ dữ liệu chưa đạt chất lượng
Given dữ liệu nguồn đang ở trạng thái NEEDS_REVIEW
When người dùng tạo content project
Then hệ thống cảnh báo dữ liệu chưa đạt điều kiện và không cho chuyển sang AI planning cho đến khi dữ liệu được xác nhận

**US-13: AI đề xuất chủ đề**

Feature: AI phân tích nguồn nội dung và đề xuất hướng triển khai video

Scenario 1: Đề xuất chủ đề và góc khai thác
Given content project đã được tạo từ dữ liệu hợp lệ
When người dùng chọn tạo kế hoạch bằng AI
Then hệ thống gửi nội dung, metadata, mục tiêu sản xuất tới AI Planning Service và AI trả về chủ đề, góc triển khai, đối tượng mục tiêu, giọng điệu đề xuất

Scenario 2: Đề xuất kiểu video phù hợp
Given AI đã phân tích nội dung nguồn
When hệ thống tạo kết quả planning
Then AI đề xuất kiểu video như kể chuyện, tóm tắt tin tức hoặc series nhiều phần và giải thích lý do theo đặc điểm nội dung

Scenario 3: Xử lý lỗi AI provider
Given người dùng yêu cầu AI đề xuất chủ đề
When AI provider quá tải hoặc trả về lỗi
Then hệ thống lưu trạng thái planning là FAILED và hiển thị thông báo lỗi cùng tùy chọn thử lại

**US-14: AI chia nội dung thành nhiều phần**

Feature: AI chia nội dung dài thành kế hoạch chuỗi video

Scenario 1: Chia nội dung thành nhiều phần có cấu trúc
Given content project chứa bài viết dài hoặc series nhiều tập
When người dùng yêu cầu AI chia nội dung
Then hệ thống tạo kế hoạch gồm nhiều phần video, mỗi phần có tiêu đề, mục tiêu nội dung, tóm tắt, hook mở đầu và thời lượng dự kiến

Scenario 2: Giữ mạch truyện giữa các phần
Given kế hoạch gồm nhiều phần liên tiếp
When AI tạo nội dung cho từng phần
Then hệ thống lưu context của phần trước để phần sau kế thừa nhân vật, sự kiện chính và điểm dừng của phần trước

Scenario 3: Điều chỉnh số phần theo yêu cầu
Given người dùng muốn thay đổi số tập hoặc độ dài từng phần
When người dùng cập nhật yêu cầu và tạo lại kế hoạch
Then AI tái phân bổ nội dung và bảo đảm trình tự câu chuyện rõ ràng, không lặp ý, không bỏ sót ý chính

**US-15: Duyệt hoặc chỉnh sửa kế hoạch**

Feature: Duyệt, chỉnh sửa hoặc tạo lại kế hoạch nội dung AI

Scenario 1: Duyệt kế hoạch phù hợp
Given AI đã tạo kế hoạch nội dung cho project
When người dùng kiểm tra và chọn Chấp nhận kế hoạch
Then hệ thống cập nhật trạng thái kế hoạch thành APPROVED và cho phép chuyển sang bước sản xuất video ở phiên bản tiếp theo

Scenario 2: Chỉnh sửa kế hoạch trước khi duyệt
Given kế hoạch AI có một số tiêu đề, hook hoặc thứ tự phần chưa phù hợp
When người dùng chỉnh sửa nội dung kế hoạch và lưu lại
Then hệ thống cập nhật phiên bản kế hoạch mới và giữ lịch sử thay đổi để đối chiếu khi cần

Scenario 3: Từ chối và yêu cầu tạo lại kế hoạch
Given người dùng chưa hài lòng với kế hoạch hiện tại
When người dùng nhập lý do và chọn tạo lại kế hoạch
Then hệ thống lưu phản hồi của người dùng và gửi lại context cùng yêu cầu điều chỉnh cho AI Planning Service
"""


def iter_body_blocks(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def remove_from_chapter2(document):
    deleting = False
    for block in list(iter_body_blocks(document)):
        if isinstance(block, Paragraph) and block.text.strip().startswith("CHƯƠNG 2"):
            deleting = True
        if deleting:
            element = block._element
            element.getparent().remove(element)


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\{width=[^}]+\}", "", text)
    return text.strip()


def add_run_with_bold(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def set_table_layout(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8)


def add_table(document, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = clean_inline(row[c_idx]) if c_idx < len(row) else ""
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9EAF7")
                tc_pr.append(shd)
    set_table_layout(table)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)
        if not is_separator:
            rows.append(cells)
        i += 1
    return rows, i


def prepare_markdown(text):
    text = re.sub(
        r"\*\*Bảng 2\.6\. Sprint planning cho phiên bản V1\.0\*\*\n\n\| Sprint \|.*?\n\n\*\*Bảng phân chia Sprint 1 và Sprint 2 theo tuần cho phiên bản V1\.0\*\*",
        "**Bảng 2.6. Sprint planning cho phiên bản V1.0**\n\n"
        + SPRINT_PLANNING_MD
        + "\n\n**Bảng phân chia Sprint 1 và Sprint 2 theo tuần cho phiên bản V1.0**",
        text,
        flags=re.S,
    )
    text = re.sub(
        r"#### 2\.5\.1\.2\. Xây dựng kịch bản câu chuyện người dùng.*?(?=#### 2\.5\.1\.3\.)",
        ENHANCED_2512 + "\n\n",
        text,
        flags=re.S,
    )
    return text


def add_markdown(document, text):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            image_path = IMAGE_DIR / rel_path
            if image_path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(6.0))
            else:
                p = document.add_paragraph()
                p.add_run(f"[Thiếu hình: {rel_path}]").italic = True
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text_part = clean_inline(heading.group(2))
            p = document.add_paragraph()
            run = p.add_run(text_part)
            run.bold = True
            if level == 1:
                run.font.size = Pt(16)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith("- "):
            p = document.add_paragraph()
            p.style = "Normal"
            p.add_run("- ")
            add_run_with_bold(p, stripped[2:].strip())
            i += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            p = document.add_paragraph()
            p.style = "Normal"
            p.add_run(ordered.group(0).split(" ", 1)[0] + " ")
            add_run_with_bold(p, ordered.group(1).strip())
            i += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            p = document.add_paragraph()
            run = p.add_run(clean_inline(stripped))
            run.bold = True
            p.paragraph_format.space_before = Pt(4)
            i += 1
            continue
        p = document.add_paragraph()
        add_run_with_bold(p, stripped)
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        i += 1


def main():
    doc = Document(SOURCE_DOCX)
    remove_from_chapter2(doc)
    chapter2_md = prepare_markdown(SOURCE_MD.read_text(encoding="utf-8"))
    add_markdown(doc, chapter2_md)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
