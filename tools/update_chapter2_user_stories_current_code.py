from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.table import Table


SOURCE = Path(r"D:\DATN\tài liệu\NguyenVanPhuoc_Chuong2_V2_Final.docx")
OUTPUT = Path(r"D:\DATN\tài liệu\NguyenVanPhuoc_Chuong2_V2_CodeUpdate_20260829_NoTabs.docx")


STORIES = [
    ("1", "Xác thực & phân quyền", "Nền tảng người dùng", "US-01", "Đăng nhập vào hệ thống", "Là người dùng, tôi muốn đăng nhập bằng email/mật khẩu để sử dụng đúng workspace và vai trò được cấp.", 2, 8, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("2", "Xác thực & phân quyền", "Nền tảng người dùng", "US-02", "Kiểm soát quyền và phạm vi dữ liệu theo vai trò", "Là người dùng hoặc quản trị viên, tôi muốn hệ thống giới hạn dữ liệu theo vai trò, chủ sở hữu và trạng thái chia sẻ để đảm bảo an toàn dữ liệu.", 3, 13, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("3", "Xác thực & phân quyền", "Nền tảng người dùng", "US-03", "Đăng xuất và kết thúc phiên client", "Là người dùng, tôi muốn đăng xuất khỏi hệ thống để ngăn người khác tiếp tục sử dụng phiên làm việc trên thiết bị của tôi.", 1, 3, "Release V1.0 (23/08)", "Sprint 1", "Medium", "Done"),
    ("4", "Tài khoản mạng xã hội", "OAuth TikTok & chiến lược kênh", "US-04", "Kết nối tài khoản TikTok bằng OAuth QR", "Là nhà sáng tạo nội dung, tôi muốn kết nối TikTok bằng luồng OAuth QR để hệ thống lưu token, scope và metadata hồ sơ phục vụ xuất bản.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("5", "Tài khoản mạng xã hội", "OAuth TikTok & chiến lược kênh", "US-05", "Cấu hình chiến lược Social Profile và tự động hóa kênh", "Là nhà sáng tạo nội dung, tôi muốn cấu hình chủ đề, giọng điệu, lịch đăng, ngưỡng chất lượng và các cờ tự động hóa cho từng kênh.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("6", "Thu thập dữ liệu", "Crawl orchestration", "US-06", "Tạo crawl job đa nguồn", "Là người dùng, tôi muốn tạo crawl job từ nhiều nguồn để hệ thống tự tạo tác vụ thu thập dữ liệu và theo dõi tiến độ xử lý.", 3, 13, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("7", "Thu thập dữ liệu", "Crawl orchestration", "US-07", "Theo dõi, hủy và chạy lại crawl job", "Là người dùng, tôi muốn xem tiến độ theo thời gian gần thực, hủy job đang chạy hoặc retry job lỗi để kiểm soát dữ liệu đầu vào.", 2, 8, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("8", "Thu thập dữ liệu", "Crawler workers", "US-08", "Thu thập dữ liệu từ VNExpress và Bilibili", "Là hệ thống, tôi muốn chọn crawler phù hợp theo source type, chạy task bất đồng bộ và retry khi lỗi để tăng độ ổn định.", 5, 21, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("9", "Xử lý dữ liệu", "Normalization pipeline", "US-09", "Chuẩn hóa dữ liệu và lưu ContentItem", "Là hệ thống, tôi muốn chuẩn hóa nội dung thô thành ContentItem có metadata, trạng thái chất lượng và dữ liệu full-text để các module sau có thể dùng lại.", 3, 13, "Release V1.0 (23/08)", "Sprint 1", "High", "Done"),
    ("10", "Xử lý dữ liệu", "Story grouping", "US-10", "Gom nhóm Story/Episode từ nội dung đã chuẩn hóa", "Là biên tập viên, tôi muốn các nội dung liên quan được gom thành story hoặc episode để dễ chọn nguồn khi lập kế hoạch video.", 3, 13, "Release V1.0 (23/08)", "Sprint 1", "Medium", "Done"),
    ("11", "Xử lý dữ liệu", "Data quality & deduplication", "US-11", "Đánh giá chất lượng, reprocess và dedup dữ liệu", "Là quản trị viên dữ liệu, tôi muốn lọc, đánh dấu trùng lặp và yêu cầu xử lý lại nội dung để giữ kho dữ liệu sạch.", 3, 13, "Release V1.0 (23/08)", "Sprint 1", "Medium", "Done"),
    ("12", "Lập kế hoạch nội dung", "MediaWorkflow", "US-12", "Tạo MediaWorkflow từ content, story hoặc crawl job", "Là nhà sáng tạo nội dung, tôi muốn tạo workflow video từ nội dung đã chọn, story hoặc kết quả crawl để bắt đầu quy trình sản xuất.", 3, 13, "Release V1.0 (23/08)", "Sprint 2", "High", "Done"),
    ("13", "Lập kế hoạch nội dung", "AI candidate planning", "US-13", "Auto Planning chọn ứng viên theo SocialProfileStrategy", "Là hệ thống, tôi muốn tự chọn ứng viên nội dung theo chiến lược từng Social Profile sau khi crawl hoàn tất để gợi ý workflow phù hợp.", 5, 21, "Release V1.0 (23/08)", "Sprint 2", "High", "Done"),
    ("14", "Lập kế hoạch nội dung", "Planning review", "US-14", "Duyệt kế hoạch theo profile và series", "Là nhà sáng tạo nội dung, tôi muốn xem workflow được nhóm theo profile/series cùng nguồn gốc và story_data để quyết định duyệt hoặc tạo lại.", 3, 13, "Release V1.0 (23/08)", "Sprint 2", "High", "Done"),
    ("15", "Lập kế hoạch nội dung", "Planning feedback", "US-15", "Phê duyệt, từ chối hoặc cập nhật metadata kế hoạch", "Là nhà sáng tạo nội dung, tôi muốn duyệt, từ chối hoặc chỉnh thông tin workflow trước khi chuyển sang sản xuất video.", 2, 8, "Release V1.0 (23/08)", "Sprint 2", "Medium", "Done"),
    ("16", "Sản xuất video", "Story timeline", "US-16", "Sinh draft story/timeline từ MediaWorkflow", "Là nhà sáng tạo nội dung, tôi muốn AI sinh timeline video từ nguồn đã chọn để có bản nháp gồm cảnh, chữ, giọng đọc và media clip.", 5, 21, "Release V2.0 (13/09)", "Sprint 3", "High", "Executing"),
    ("17", "Sản xuất video", "Story editing", "US-17", "Chỉnh sửa và review timeline bằng AI hoặc thủ công", "Là nhà sáng tạo nội dung, tôi muốn chỉnh story_data/timeline thủ công hoặc bằng prompt AI để tinh chỉnh nội dung trước khi render.", 3, 13, "Release V2.0 (13/09)", "Sprint 3", "High", "Executing"),
    ("18", "Sản xuất video", "Voice generation", "US-18", "Tạo giọng đọc cho timeline video", "Là nhà sáng tạo nội dung, tôi muốn tạo voice-over từ timeline bằng Edge TTS hoặc ElevenLabs để video có âm thanh đồng bộ.", 3, 13, "Release V2.0 (13/09)", "Sprint 3", "High", "Executing"),
    ("19", "Sản xuất video", "Timeline alignment", "US-19", "Căn chỉnh voice, phụ đề và visual timeline", "Là hệ thống, tôi muốn đồng bộ audio clip, text clip và visual clip trong timeline để nội dung hiển thị đúng nhịp.", 3, 13, "Release V2.0 (13/09)", "Sprint 3", "High", "Executing"),
    ("20", "Sản xuất video", "Video rendering", "US-20", "Render video MP4 bằng Remotion Worker", "Là nhà sáng tạo nội dung, tôi muốn xuất bản nháp thành file MP4 cuối cùng để có thể duyệt và đưa vào hàng đợi đăng.", 5, 21, "Release V2.0 (13/09)", "Sprint 3", "High", "Executing"),
    ("21", "Sản xuất video", "ContentSeries", "US-21", "Quản lý ContentSeries và ngữ cảnh liên tục", "Là nhà sáng tạo nội dung, tôi muốn gom workflow vào series, kiểm tra tính nhất quán và tái sử dụng ngữ cảnh cho các tập tiếp theo.", 3, 13, "Release V2.0 (13/09)", "Sprint 3", "Medium", "Executing"),
    ("22", "Sản xuất video", "Video workspace", "US-22", "Theo dõi xưởng sản xuất video và preview workflow", "Là nhà sáng tạo nội dung, tôi muốn xem kanban, bộ lọc, khả năng thao tác và tiến độ từng workflow trong một workspace tập trung.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "Medium", "Executing"),
    ("23", "Sản xuất video", "Video approval", "US-23", "Duyệt video đã render trước khi đưa vào queue", "Là nhà sáng tạo nội dung, tôi muốn xem video MP4 đã render và chỉ đánh dấu approved khi nội dung đạt yêu cầu.", 2, 8, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("24", "Xuất bản", "Publishing queue", "US-24", "Đưa video đã duyệt vào hàng đợi và lên lịch", "Là nhà sáng tạo nội dung, tôi muốn đưa video approved vào PublishingQueueItem với lịch đăng phù hợp hoặc trạng thái cần duyệt.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("25", "Xuất bản", "Queue operations", "US-25", "Theo dõi, lọc và xử lý Publishing Queue", "Là người vận hành, tôi muốn lọc queue, xem chi tiết nguồn/video và yêu cầu thay đổi, bỏ qua, duyệt hoặc đăng ngay từng item.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("26", "Xuất bản", "Metrics", "US-26", "Ghi nhận metrics và xem hiệu suất bài đăng", "Là nhà sáng tạo nội dung, tôi muốn ghi nhận lượt xem, thích, bình luận, chia sẻ và xem overview tăng trưởng của các bài đã đăng.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "Medium", "Executing"),
    ("27", "Xuất bản", "TikTok publishing", "US-27", "Đăng ngay TikTok bằng Direct Post hoặc Inbox Upload", "Là nhà sáng tạo nội dung, tôi muốn đăng video lên TikTok theo scope tài khoản, tạo SocialPost và cập nhật trạng thái queue.", 5, 21, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("28", "Xuất bản", "Publish scheduler", "US-28", "Tự động đăng queue item đến hạn theo scheduler", "Là hệ thống, tôi muốn kiểm tra queue định kỳ và chỉ tự đăng khi chiến lược kênh cho phép auto publish.", 5, 21, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("31", "Xuất bản", "Rendered workflow queue sync", "US-31", "Tự đồng bộ video đã render sang hàng đợi chờ duyệt", "Là người vận hành, tôi muốn các MediaWorkflow đã render được tự tạo PublishingQueueItem khi mở queue để không bỏ sót video cần duyệt.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("32", "Xuất bản", "Review feedback loop", "US-32", "Trả workflow về biên tập khi reviewer yêu cầu chỉnh sửa", "Là reviewer, tôi muốn khi yêu cầu chỉnh sửa ở Publishing Queue thì workflow liên quan quay lại trạng thái EDITING và lưu rõ ghi chú review.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "High", "Executing"),
    ("33", "Sản xuất video", "AI series decision", "US-33", "Lưu metadata series do AI đề xuất theo chủ đề và mạch truyện", "Là nhà sáng tạo nội dung, tôi muốn AI khi lập kế hoạch hoặc sinh kịch bản có thể tạo/gán series với mô tả, loại series và tổng số tập để quản lý nội dung dài kỳ chính xác hơn.", 3, 13, "Release V2.0 (13/09)", "Sprint 4", "Medium", "Executing"),
    ("34", "Sản xuất video", "Real media preview", "US-34", "Hiển thị thumbnail, avatar và video preview từ dữ liệu thật", "Là người vận hành sản xuất, tôi muốn workspace và approval queue hiển thị media thật từ nguồn crawl/workflow thay vì ảnh demo để kiểm tra đúng nội dung trước khi duyệt.", 2, 8, "Release V2.0 (13/09)", "Sprint 4", "Medium", "Executing"),
    ("29", "Triển khai & vận hành", "Deployment", "US-29", "Triển khai hệ thống microservices lên môi trường chính thức", "Là quản trị viên hệ thống, tôi muốn triển khai các service API, ingestion, AI media và frontend theo cấu hình môi trường để vận hành ổn định.", 5, 21, "Release Production (20/09)", "Sprint 5", "High", "Planned"),
    ("30", "Triển khai & vận hành", "Administration", "US-30", "Quản trị hệ thống và cấu hình lịch tự động hóa", "Là quản trị viên, tôi muốn quản lý người dùng, cấu hình hệ thống, audit log và trạng thái scheduler để kiểm soát vận hành.", 3, 13, "Release Production (20/09)", "Sprint 5", "Medium", "Planned"),
]


SCENARIOS = {
    "US-01": ("Xác thực JWT cho người dùng", [
        ("Đăng nhập thành công", "Người dùng nhập email và mật khẩu hợp lệ, API trả về access token và thông tin vai trò; frontend lưu phiên và chuyển người dùng vào workspace phù hợp.", "Sai mật khẩu hoặc tài khoản không hợp lệ sẽ bị từ chối với thông báo lỗi rõ ràng."),
        ("Truy cập tài nguyên cần xác thực", "Frontend gửi token trong header, backend giải mã người dùng hiện tại và cho phép gọi API theo quyền.", "Token hết hạn hoặc thiếu token sẽ nhận lỗi xác thực và yêu cầu đăng nhập lại."),
    ]),
    "US-02": ("Phân quyền và giới hạn phạm vi dữ liệu", [
        ("Người dùng thường xem dữ liệu của mình", "API lọc ContentItem, CrawlJob, MediaWorkflow và queue theo owner hoặc dữ liệu GLOBAL có thể chia sẻ.", "Dữ liệu PRIVATE của người khác không xuất hiện trong danh sách và không truy cập được qua detail API."),
        ("Quản trị viên truy cập phạm vi hệ thống", "SYSTEM_ADMIN có thể xem dữ liệu toàn hệ thống, cập nhật nội dung, reprocess và đánh dấu duplicate.", "Các thao tác admin bị chặn nếu người gọi không có quyền phù hợp."),
    ]),
    "US-03": ("Kết thúc phiên làm việc", [
        ("Đăng xuất khỏi ứng dụng", "Người dùng chọn đăng xuất, frontend xóa token và trạng thái phiên, sau đó chuyển về màn hình đăng nhập.", "Các request tiếp theo không còn thông tin xác thực nên bị backend từ chối."),
        ("Bảo vệ route sau khi logout", "Người dùng mở lại route nội bộ, guard xác thực phát hiện chưa đăng nhập và điều hướng về trang login.", "Không có dữ liệu cá nhân nào được hiển thị từ cache phiên cũ."),
    ]),
    "US-06": ("Tạo crawl job bất đồng bộ", [
        ("Tạo job với nhiều nguồn", "Người dùng khai báo danh sách nguồn, API lưu CrawlJob/CrawlJobSource và phát event crawl.job.created; orchestrator tạo KafkaTask CRAWL_URL cho từng nguồn.", "Nếu không có nguồn hợp lệ, job được đánh dấu failed thay vì tạo task rỗng."),
        ("Khởi động tiến trình khám phá", "Orchestrator cập nhật stage DISCOVERING, thống kê nguồn và phát job_progress để UI theo dõi.", "Event trùng lặp được xử lý idempotent để tránh tạo task hai lần."),
    ]),
    "US-07": ("Theo dõi và điều khiển crawl job", [
        ("Xem tiến độ qua màn hình crawl", "Frontend gọi detail hoặc SSE progress để hiển thị status, stage, discovered, processed, failed, duplicates và progress.", "Endpoint logs hiện trả danh sách rỗng vì log chi tiết đã chuyển khỏi Postgres."),
        ("Hủy hoặc retry job", "Người dùng hủy job đang chạy hoặc retry job lỗi; API cập nhật trạng thái và tạo lại luồng xử lý khi hợp lệ.", "Crawler kiểm tra cờ cancel trong khi chạy để dừng task đúng cách."),
    ]),
    "US-08": ("Crawler worker theo loại nguồn", [
        ("Thu thập từ VNExpress", "Crawler runner chọn VNExpressCrawler, fetch danh sách bài, chuyển đổi dữ liệu về schema chung và lưu batch nội dung thô.", "Lỗi tạm thời được retry với backoff và chuyển dead-letter nếu vượt quá giới hạn."),
        ("Thu thập từ Bilibili", "Runner chọn BilibiliCrawler theo source_type, xử lý output đã chuẩn hóa hoặc tự chuẩn hóa trước khi phát normalized events.", "Task cập nhật failed_count của job nếu nguồn không trả dữ liệu hợp lệ."),
    ]),
    "US-09": ("Chuẩn hóa và lưu nội dung", [
        ("Tạo ContentItem từ dữ liệu crawl", "Pipeline lưu metadata, full text, media, trạng thái chất lượng và liên kết source để các module sau truy vấn.", "Nội dung thiếu dữ liệu quan trọng được gắn trạng thái cần review hoặc usable with warning."),
        ("Truy vấn kho nội dung", "Người dùng xem danh sách hoặc final view, backend nạp full text từ repository và tính profile_matches theo SocialProfileStrategy.", "Người dùng thường chỉ thấy dữ liệu mình sở hữu hoặc nội dung GLOBAL."),
    ]),
    "US-10": ("Gom nhóm story/episode", [
        ("Tạo story từ các ContentItem liên quan", "Hệ thống gom các nội dung cùng chủ đề thành Story, lưu episode_order và metadata giúp chọn nguồn video nhanh hơn.", "Nếu dữ liệu không đủ tương quan, nội dung vẫn tồn tại như item độc lập."),
        ("Xem detail story", "Người dùng mở Story để xem danh sách episode, nguồn và dữ liệu phục vụ planning.", "Các story không thuộc phạm vi quyền sẽ không hiển thị."),
    ]),
    "US-11": ("Kiểm soát chất lượng dữ liệu", [
        ("Đánh dấu nội dung trùng lặp", "Quản trị viên gọi mark-duplicate để chuyển trạng thái nội dung và tránh đưa vào planning tự động.", "Người dùng không có quyền admin không thể thao tác duplicate thủ công."),
        ("Yêu cầu xử lý lại nội dung", "Admin gọi reprocess để phát lại luồng xử lý cho ContentItem có dữ liệu lỗi hoặc cũ.", "API trả lỗi nếu ContentItem không tồn tại hoặc người gọi không có quyền."),
    ]),
    "US-12": ("Khởi tạo MediaWorkflow", [
        ("Tạo workflow từ content hoặc story đã chọn", "Người dùng chọn content_ids/story_ids, API kiểm tra quyền truy cập, tạo MediaWorkflow READY và lưu input metadata.", "Nếu nguồn không hợp lệ hoặc không thuộc profile, workflow không được tạo."),
        ("Tạo workflow từ crawl job", "API chọn candidate content tốt nhất theo quality_score và updated_at, tạo workflow READY hoặc NEEDS_REVIEW tùy trạng thái dữ liệu.", "Job không có nội dung phù hợp sẽ trả lỗi để người dùng bổ sung dữ liệu."),
    ]),
    "US-13": ("Auto Planning theo chiến lược profile", [
        ("Tạo candidate sau khi crawl hoàn tất", "Consumer crawl.job.completed đọc các SocialProfile active có receive_system_content và auto_project_queue_enabled, sau đó tạo PlanningRun và PlanningCandidate.", "Job failed hoặc profile thiếu chiến lược sẽ không tạo candidate tự động."),
        ("Chọn nội dung ưu tiên", "Hệ thống đánh giá candidate theo chủ đề, độ phù hợp và chất lượng nguồn để gắn selected_content_id cho workflow.", "Candidate không đạt ngưỡng vẫn được ghi nhận lý do để người dùng xem lại."),
    ]),
    "US-14": ("Duyệt kế hoạch theo profile/series", [
        ("Xem danh sách kế hoạch", "Planning page nhóm workflow theo SocialProfile và ContentSeries, hiển thị source article, story_data và lý do chọn.", "Nếu strategy thiếu topic config, UI hiển thị cảnh báo cấu hình."),
        ("Tạo lại hoặc xem nguồn", "Người dùng mở kế hoạch để xem nguồn gốc, yêu cầu regenerate hoặc chuyển sang workspace sản xuất.", "Workflow bị reject không được tự chuyển tiếp sang bước render."),
    ]),
    "US-15": ("Phản hồi kế hoạch", [
        ("Approve hoặc reject workflow", "Người dùng duyệt hoặc từ chối, API cập nhật MediaWorkflow và lưu PlanningFeedback.", "Lý do reject được giữ trong metadata để truy vết quyết định."),
        ("Cập nhật metadata kế hoạch", "Người dùng sửa title, series_id, caption, tags hoặc draft_json trước khi sản xuất.", "Series bị xóa sẽ tự bỏ liên kết khỏi các workflow liên quan."),
    ]),
    "US-16": ("Sinh story/timeline", [
        ("Tạo story từ workflow", "Người dùng bấm create story, API tạo KafkaTask GENERATE_VIDEO_SCRIPT và phát generate-video.script.requested; worker chạy các stage loading, generating, normalizing và saving draft.", "Nếu đang có task script active, API chặn tạo task trùng."),
        ("Tạo direct script từ một ContentItem", "Người dùng có thể tạo workflow và script trực tiếp từ content cụ thể để bỏ qua bước chọn candidate.", "ContentItem không truy cập được sẽ bị từ chối."),
    ]),
    "US-17": ("Chỉnh sửa và review timeline", [
        ("Chỉnh sửa bằng prompt AI", "Người dùng nhập prompt edit, API lưu story hiện tại rồi tạo KafkaTask GENERATE_VIDEO_EDIT; worker chỉnh timeline dựa trên nguồn và yêu cầu.", "Prompt trống hoặc workflow không thuộc người dùng sẽ bị từ chối."),
        ("Review story bằng AI", "Người dùng yêu cầu review, worker kiểm tra logic, độ bám nguồn và overlap clip rồi trả lại bản đề xuất.", "Nếu review phát hiện thay đổi draft, UI cảnh báo cần tạo lại voice/video."),
    ]),
    "US-18": ("Tạo giọng đọc", [
        ("Sinh voice-over tự động", "Người dùng chọn provider, voice và speed; API tạo task GENERATE_VIDEO_VOICE, worker sinh MP3, căn duration và lưu audio artifact.", "Workflow chưa có draft hợp lệ sẽ không thể tạo voice."),
        ("Dùng provider phù hợp", "Hệ thống dùng Edge TTS mặc định hoặc ElevenLabs khi có cấu hình, đồng thời fallback khi dịch vụ ngoài không sẵn sàng.", "Voice tag và ký tự không phù hợp được làm sạch trước khi render audio."),
    ]),
    "US-19": ("Đồng bộ timeline", [
        ("Căn voice với subtitle", "Worker upsert audio clip, fit frame bằng whisper khi có thể và đồng bộ text clip theo timeline.", "Nếu user sửa text sau khi tạo voice, UI đánh dấu cần regenerate voice/video."),
        ("Chuẩn hóa visual clip", "Timeline service chuẩn hóa video, text, audio clip, chống overlap và link text với visual clip tương ứng.", "Dữ liệu timeline rỗng được bổ sung cấu trúc fallback để không làm hỏng render."),
    ]),
    "US-20": ("Render video cuối", [
        ("Xuất MP4 bằng Remotion", "Người dùng bấm export video, API tạo task GENERATE_VIDEO_RENDER; worker dựng video qua Remotion và lưu FINAL_VIDEO artifact.", "Render bị chặn nếu còn task script/edit/review/voice active."),
        ("Cập nhật trạng thái sau render", "Workflow chuyển sang RENDERED hoặc WAITING_APPROVAL tùy chính sách duyệt của profile.", "Lỗi render được ghi vào task/workflow để UI hiển thị và retry."),
    ]),
    "US-21": ("Quản lý series nội dung", [
        ("Tạo và gán ContentSeries", "Người dùng tạo series cho profile, gán workflow và dùng context series để giữ mạch nội dung nhiều tập.", "Xóa series sẽ unset series_id khỏi workflow thay vì xóa workflow."),
        ("Kiểm tra tính nhất quán", "API rebuild context và consistency-check để cảnh báo số tập thiếu, scene rỗng hoặc dữ liệu không liền mạch.", "Cảnh báo không chặn workflow nhưng giúp người dùng sửa trước khi xuất bản."),
    ]),
    "US-22": ("Xưởng sản xuất video", [
        ("Theo dõi workflow theo kanban", "Frontend hiển thị các cột draft sẵn sàng, đang chỉnh sửa, đang render, chờ duyệt và sẵn sàng xuất bản, kèm bộ lọc profile/series/status.", "Workflow không có quyền truy cập sẽ không xuất hiện."),
        ("Mở workspace chi tiết", "Người dùng xem progress, capabilities, source_content, story, voice, video preview và các action có thể thực hiện ở từng stage.", "Action bị disable khi thiếu draft, thiếu voice, thiếu final video hoặc đang có task active."),
    ]),
    "US-23": ("Duyệt video đã render", [
        ("Approve video", "Người dùng xem MP4 cuối, gọi approve-video để đánh dấu video_approved và lưu thông tin review Module 4.", "API yêu cầu phải có file video đã render trước khi approve."),
        ("Tự động đưa vào queue sau approve", "Nếu strategy bật auto_queue_enabled, hệ thống tạo hoặc cập nhật PublishingQueueItem sau khi video được duyệt.", "Nếu kênh yêu cầu manual approval, item được giữ ở trạng thái phù hợp để người dùng duyệt lịch."),
    ]),
    "US-04": ("Kết nối TikTok OAuth QR", [
        ("Bắt đầu phiên QR", "Người dùng tạo QR OAuth, frontend theo dõi status pending và backend lưu state phiên kết nối.", "Người dùng có thể stop phiên QR chưa hoàn tất để tránh callback cũ."),
        ("Hoàn tất kết nối profile", "Sau callback, hệ thống upsert SocialProfile với external_id, access/refresh token, scopes, avatar và thống kê cơ bản.", "Thiếu scope video.publish hoặc video.upload sẽ giới hạn kiểu xuất bản khả dụng."),
    ]),
    "US-05": ("Chiến lược Social Profile", [
        ("Cấu hình nội dung và lịch", "Người dùng cập nhật topic, avoid topic, tone, target audience, post frequency, active hours, schedule_days, schedule_times và timezone.", "Giá trị lịch không hợp lệ sẽ không được scheduler dùng để đăng tự động."),
        ("Cấu hình tự động hóa", "Người dùng bật/tắt receive_system_content, auto_project_queue_enabled, video_render_mode, auto_queue_enabled, auto_publish_enabled, approval_mode và min_score.", "Scheduler chỉ đăng khi các cờ chiến lược cho phép."),
    ]),
    "US-24": ("Đưa video vào hàng đợi", [
        ("Queue video đã duyệt", "Người dùng gọi queue-post, API kiểm tra final_video và video_approved rồi tạo PublishingQueueItem với caption, tags, scheduled_at và profile.", "Video chưa approve hoặc chưa render sẽ không được queue."),
        ("Chọn lịch thủ công hoặc gợi ý", "Người dùng duyệt lịch đăng bằng slot AI hoặc nhập lịch thủ công trong trang approval/schedule.", "Item có thể ở trạng thái queued, needs_approval hoặc approved tùy chiến lược profile."),
    ]),
    "US-25": ("Vận hành Publishing Queue", [
        ("Lọc và xem chi tiết item", "Người vận hành lọc theo profile, nền tảng, trạng thái, ngày và search, sau đó mở dialog xem source article, video, caption và lịch.", "Item không thuộc profile được phép xem sẽ bị ẩn."),
        ("Yêu cầu sửa hoặc bỏ qua", "Người dùng request changes, reject/skip hoặc approve schedule để điều chỉnh trước khi đăng.", "Lý do thay đổi được lưu cùng queue item để truy vết."),
    ]),
    "US-26": ("Theo dõi hiệu suất bài đăng", [
        ("Ghi nhận metrics", "Hệ thống hoặc người vận hành gửi SocialPostMetric gồm views, likes, comments, shares và snapshot thời điểm đo.", "Không giả định worker tự kéo metrics nếu chưa có cấu hình tích hợp tương ứng."),
        ("Xem overview tăng trưởng", "API nhóm SocialPost theo nội dung, lấy metric mới nhất và tính growth để hiển thị hiệu suất.", "Bài chưa có metrics vẫn xuất hiện với dữ liệu hiệu suất rỗng."),
    ]),
    "US-27": ("Đăng ngay TikTok", [
        ("Direct Post", "Người dùng chọn publish now khi profile có scope video.publish; service lấy MP4 đã render, gọi TikTok Direct Post và tạo SocialPost.", "Nếu TikTok trả lỗi, queue item chuyển failed và lưu error message."),
        ("Inbox Upload", "Nếu profile chỉ có scope video.upload, hệ thống dùng inbox upload thay vì direct post và cập nhật trạng thái theo phản hồi.", "Profile thiếu cả hai scope không thể đăng từ hệ thống."),
    ]),
    "US-28": ("Scheduler xuất bản tự động", [
        ("Quét queue đến hạn", "Publish scheduler chạy định kỳ, tìm item queued/approved đến hạn, profile active và có token để xử lý.", "Scheduler bỏ qua item nếu schedule_enabled hoặc auto_publish_enabled không bật."),
        ("Tôn trọng approval mode", "Với approval_mode manual, item queued chưa approve không được đăng tự động; với auto, scheduler có thể publish nếu đủ điều kiện.", "Admin có thể xem snapshot settings, jobs và last_run để giám sát."),
    ]),
    "US-29": ("Triển khai microservices", [
        ("Cấu hình môi trường triển khai", "Quản trị viên thiết lập biến môi trường cho API, ingestion engine, AI media engine, frontend, database, object storage và Kafka.", "Thiếu cấu hình bắt buộc sẽ khiến service fail fast thay vì chạy sai ngầm."),
        ("Khởi chạy các service", "Các service được deploy theo vai trò riêng: api-service, data-ingestion-engine, ai-media-engine, frontend và scheduler trong API lifespan.", "Log và health check được dùng để xác nhận từng service đã sẵn sàng."),
    ]),
    "US-30": ("Quản trị vận hành", [
        ("Quản lý người dùng và cấu hình", "SYSTEM_ADMIN quản lý người dùng, vai trò, system setting và audit log để kiểm soát quyền vận hành.", "Thao tác quản trị bị chặn nếu tài khoản không có role phù hợp."),
        ("Giám sát scheduler và queue", "Admin theo dõi snapshot scheduler, queue tồn đọng, item failed và trạng thái publish để can thiệp kịp thời.", "Các lỗi publish được giữ lại trên queue item để retry hoặc xử lý thủ công."),
    ]),
}


V1_ORDER = ["US-01", "US-02", "US-03", "US-06", "US-07", "US-08", "US-09", "US-10", "US-11", "US-12", "US-13", "US-14", "US-15"]
V2_ORDER = ["US-16", "US-17", "US-18", "US-19", "US-20", "US-21", "US-22", "US-23", "US-04", "US-05", "US-24", "US-25", "US-26", "US-27", "US-28", "US-31", "US-32", "US-33", "US-34"]
PROD_ORDER = ["US-29", "US-30"]

GHERKIN_SCENARIOS = {
    "US-01": ("Xác thực người dùng", [
        ("Đăng nhập thành công", "người dùng đã có tài khoản hợp lệ trong hệ thống", "nhập email, mật khẩu và bấm Đăng nhập", "hệ thống trả JWT token, thông tin vai trò và điều hướng người dùng vào workspace phù hợp"),
        ("Từ chối đăng nhập sai", "người dùng nhập sai mật khẩu hoặc tài khoản không tồn tại", "gửi yêu cầu đăng nhập", "hệ thống không cấp token và hiển thị thông báo lỗi xác thực"),
    ]),
    "US-02": ("Phân quyền và phạm vi dữ liệu", [
        ("Người dùng chỉ xem dữ liệu được phép", "người dùng đăng nhập với vai trò thường", "truy cập danh sách ContentItem, CrawlJob, MediaWorkflow hoặc Publishing Queue", "hệ thống chỉ trả dữ liệu thuộc owner hiện tại hoặc dữ liệu GLOBAL được chia sẻ"),
        ("Quản trị viên truy cập dữ liệu hệ thống", "quản trị viên đăng nhập với vai trò SYSTEM_ADMIN", "mở màn hình quản trị hoặc gọi API quản trị nội dung", "hệ thống cho phép xem dữ liệu toàn hệ thống và thực hiện thao tác reprocess hoặc mark-duplicate"),
    ]),
    "US-03": ("Đăng xuất phiên làm việc", [
        ("Đăng xuất khỏi ứng dụng", "người dùng đang đăng nhập và có JWT token hợp lệ ở client", "bấm Đăng xuất", "frontend xóa token, xóa trạng thái phiên và chuyển người dùng về màn hình đăng nhập"),
        ("Chặn truy cập sau đăng xuất", "người dùng đã đăng xuất khỏi ứng dụng", "mở lại một route nội bộ cần xác thực", "hệ thống yêu cầu đăng nhập lại và không hiển thị dữ liệu từ phiên cũ"),
    ]),
    "US-04": ("Kết nối TikTok OAuth QR", [
        ("Bắt đầu phiên kết nối TikTok", "người dùng đã đăng nhập và chưa hoàn tất kết nối TikTok", "bấm Kết nối TikTok bằng QR", "hệ thống tạo phiên OAuth QR, trả trạng thái pending và cho phép frontend polling trạng thái kết nối"),
        ("Hoàn tất kết nối Social Profile", "TikTok OAuth callback trả token và scope hợp lệ", "backend xử lý callback", "hệ thống tạo hoặc cập nhật SocialProfile với external_id, access token, refresh token, scope, avatar và thống kê hồ sơ"),
    ]),
    "US-05": ("Cấu hình chiến lược Social Profile", [
        ("Cập nhật chủ đề và lịch đăng", "người dùng đã có Social Profile active", "cấu hình content_topics, avoid_topics, tone, target_audience, active_hours, schedule_days, schedule_times và timezone", "hệ thống lưu SocialProfileStrategy để dùng cho planning, queue và scheduler"),
        ("Bật tự động hóa kênh", "người dùng đang ở màn hình cấu hình chiến lược kênh", "bật receive_system_content, auto_project_queue_enabled, video_render_mode, auto_queue_enabled hoặc auto_publish_enabled", "hệ thống áp dụng các cờ tự động hóa vào luồng chọn nội dung, render, queue và đăng bài"),
    ]),
    "US-06": ("Tạo crawl job đa nguồn", [
        ("Tạo crawl job hợp lệ", "người dùng đã đăng nhập và nhập ít nhất một nguồn crawl", "bấm Tạo Crawl Job", "API lưu CrawlJob, CrawlJobSource và phát sự kiện crawl.job.created cho orchestrator"),
        ("Orchestrator tạo task crawl", "orchestrator nhận sự kiện crawl.job.created", "kiểm tra danh sách nguồn và claim event thành công", "hệ thống tạo KafkaTask CRAWL_URL cho từng nguồn và cập nhật stage DISCOVERING"),
    ]),
    "US-07": ("Theo dõi, hủy và retry crawl job", [
        ("Theo dõi tiến độ crawl", "crawl job đã được tạo và đang xử lý", "người dùng mở màn hình chi tiết hoặc progress stream", "hệ thống hiển thị status, stage, progress, discovered, processed, failed và duplicates"),
        ("Retry job lỗi", "crawl job kết thúc với trạng thái lỗi hoặc partial failure", "người dùng bấm Retry", "hệ thống tạo lại luồng xử lý phù hợp và cập nhật trạng thái job để worker tiếp tục xử lý"),
    ]),
    "US-08": ("Crawler worker theo source type", [
        ("Thu thập dữ liệu VNExpress", "KafkaTask CRAWL_URL có source_type VNEXPRESS", "crawler runner nhận task", "hệ thống dùng VNExpressCrawler, lưu dữ liệu thô và cập nhật counters của crawl job"),
        ("Làm sạch nội dung bài VNExpress", "bài VNExpress có sapo, đoạn liên quan, CTA Xem thêm hoặc thông tin tác giả nằm trong thẻ p", "crawler trích xuất phần thân bài", "hệ thống bỏ đoạn mô tả đã lưu riêng, bỏ link điều hướng/bài liên quan và chỉ giữ nội dung chính để tránh trùng lặp trong pipeline"),
    ]),
    "US-09": ("Chuẩn hóa và lưu ContentItem", [
        ("Tạo ContentItem từ dữ liệu thô", "crawler đã lưu raw document từ nguồn crawl", "normalization pipeline xử lý dữ liệu", "hệ thống tạo ContentItem có metadata, full text, media metadata và trạng thái chất lượng"),
        ("Gắn trạng thái cần review", "nội dung thiếu dữ liệu quan trọng hoặc chất lượng thấp", "pipeline hoàn tất bước chuẩn hóa", "hệ thống đánh dấu NEEDS_REVIEW hoặc USABLE_WITH_WARNING để tránh planning sai"),
    ]),
    "US-10": ("Gom nhóm Story/Episode", [
        ("Gom các nội dung liên quan", "nhiều ContentItem đã được chuẩn hóa và có quan hệ chủ đề", "story grouping pipeline chạy", "hệ thống tạo Story, liên kết episode và lưu episode_order"),
        ("Xem chi tiết story", "người dùng có quyền truy cập story", "mở trang chi tiết story", "hệ thống hiển thị danh sách episode, nguồn nội dung và metadata phục vụ planning"),
    ]),
    "US-11": ("Đánh giá chất lượng và dedup dữ liệu", [
        ("Đánh dấu nội dung trùng lặp", "quản trị viên đăng nhập với vai trò SYSTEM_ADMIN", "chọn ContentItem và bấm Mark Duplicate", "hệ thống cập nhật trạng thái duplicate để nội dung không được chọn vào planning tự động"),
        ("Yêu cầu reprocess nội dung", "ContentItem đã tồn tại nhưng cần xử lý lại", "quản trị viên bấm Reprocess", "hệ thống tạo yêu cầu xử lý lại và cập nhật trạng thái cho luồng normalization/dedup"),
    ]),
    "US-12": ("Khởi tạo MediaWorkflow", [
        ("Tạo MediaWorkflow từ nội dung đã chọn", "người dùng đã chọn content_ids hoặc story_ids có quyền truy cập", "bấm Tạo workflow video", "API tạo MediaWorkflow READY, lưu input metadata và primary_content_id nếu có"),
        ("Tạo MediaWorkflow từ crawl job", "crawl job có candidate content đạt chất lượng", "người dùng tạo workflow từ kết quả crawl", "hệ thống chọn candidate theo quality_score, tạo MediaWorkflow READY hoặc NEEDS_REVIEW"),
    ]),
    "US-13": ("Auto Planning theo SocialProfileStrategy", [
        ("Tạo planning run sau crawl completed", "crawl job hoàn tất với trạng thái SUCCEEDED hoặc PARTIAL_SUCCESS", "consumer nhận sự kiện crawl.job.completed", "hệ thống tạo PlanningRun và PlanningCandidate cho các SocialProfile bật receive_system_content và auto_project_queue_enabled"),
        ("Chọn candidate phù hợp", "PlanningRun có nhiều candidate nội dung", "AI planning đánh giá theo topic, min_score và profile strategy", "hệ thống lưu selected_content_id, selection_reasons và workflow đề xuất"),
    ]),
    "US-14": ("Duyệt kế hoạch theo profile và series", [
        ("Xem danh sách kế hoạch", "người dùng đã có SocialProfile và MediaWorkflow được tạo", "mở trang Planning ở chế độ duyệt thành phẩm", "hệ thống nhóm workflow theo profile/series và hiển thị source_content, story_data cùng lý do chọn"),
        ("Regenerate kế hoạch", "workflow chưa đạt kỳ vọng nội dung", "người dùng yêu cầu tạo lại hoặc chỉnh kế hoạch", "hệ thống cập nhật workflow hoặc tạo task tương ứng để chuẩn bị lại nội dung"),
    ]),
    "US-15": ("Phản hồi kế hoạch nội dung", [
        ("Phê duyệt workflow", "MediaWorkflow đang ở trạng thái chờ duyệt kế hoạch", "người dùng bấm Approve", "hệ thống cập nhật trạng thái workflow và lưu PlanningFeedback"),
        ("Từ chối workflow", "MediaWorkflow không phù hợp với chiến lược nội dung", "người dùng bấm Reject và nhập lý do", "hệ thống lưu feedback từ chối và không chuyển workflow sang bước sản xuất video"),
    ]),
    "US-16": ("Sinh story/timeline từ MediaWorkflow", [
        ("Tạo draft story", "MediaWorkflow có nguồn nội dung hợp lệ và chưa có script task active", "người dùng bấm Create Story", "API tạo KafkaTask GENERATE_VIDEO_SCRIPT và worker lưu draft timeline vào MediaWorkflow"),
        ("Tạo direct script từ ContentItem", "người dùng chọn một ContentItem có quyền truy cập", "gọi chức năng Direct Script", "hệ thống tạo MediaWorkflow trực tiếp từ ContentItem và enqueue task sinh script"),
    ]),
    "US-17": ("Chỉnh sửa và review timeline", [
        ("Chỉnh sửa timeline bằng AI", "MediaWorkflow đã có draft story/timeline", "người dùng nhập prompt chỉnh sửa và bấm Edit Story", "hệ thống tạo KafkaTask GENERATE_VIDEO_EDIT, lưu revision hiện tại và cập nhật timeline theo prompt"),
        ("Review story bằng AI", "MediaWorkflow đã có draft cần kiểm tra", "người dùng bấm Review Story", "worker kiểm tra logic, độ bám nguồn và overlap clip rồi trả lại bản đề xuất"),
    ]),
    "US-18": ("Tạo giọng đọc cho video", [
        ("Sinh voice-over từ timeline", "MediaWorkflow có draft timeline hợp lệ", "người dùng chọn provider, voice, speed và bấm Generate Voice", "hệ thống tạo KafkaTask GENERATE_VIDEO_VOICE, sinh MP3 và lưu audio artifact"),
        ("Fallback provider giọng đọc", "provider chính không sẵn sàng hoặc thiếu cấu hình", "worker xử lý tác vụ voice", "hệ thống dùng provider fallback phù hợp và vẫn lưu audio clip vào timeline nếu sinh thành công"),
    ]),
    "US-19": ("Đồng bộ voice, phụ đề và visual timeline", [
        ("Căn voice với text clip", "file MP3 voice-over đã được tạo", "worker chạy bước alignment", "hệ thống cập nhật duration, audio clip, text clip và timestamp để timeline đồng bộ"),
        ("Cảnh báo khi draft thay đổi", "người dùng chỉnh story sau khi đã tạo voice hoặc render", "frontend phát hiện revision timeline thay đổi", "hệ thống cảnh báo cần regenerate voice/video để tránh lệch nội dung"),
    ]),
    "US-20": ("Render video bằng Remotion Worker", [
        ("Render video MP4", "MediaWorkflow đã có timeline đủ dữ liệu và không có pre-render task active", "người dùng bấm Export Video", "API tạo KafkaTask GENERATE_VIDEO_RENDER và worker xuất file MP4 bằng Remotion"),
        ("Lưu final video artifact", "Remotion render hoàn tất thành công", "worker cập nhật kết quả render", "hệ thống lưu FINAL_VIDEO artifact, cập nhật final_video và chuyển workflow sang trạng thái chờ duyệt hoặc rendered"),
    ]),
    "US-21": ("Quản lý ContentSeries", [
        ("Tạo và gán series thủ công", "người dùng có SocialProfile và nhiều workflow liên quan", "tạo ContentSeries hoặc gán workflow vào series hiện có", "hệ thống lưu series_id, description, current_part và total_parts để dùng context liên tục cho các tập tiếp theo"),
        ("Kiểm tra tính nhất quán series", "series đã có nhiều workflow hoặc episode", "người dùng chạy consistency-check", "hệ thống cảnh báo scene rỗng, thiếu số tập hoặc dữ liệu không liền mạch"),
    ]),
    "US-22": ("Xưởng sản xuất video", [
        ("Xem kanban workflow bốn cột", "người dùng có danh sách MediaWorkflow", "mở trang Xưởng sản xuất video", "hệ thống hiển thị các cột Draft kịch bản, Đang biên tập & Voice, Đang render MP4 và Chờ duyệt video theo trạng thái thực tế của workflow"),
        ("Mở studio chi tiết", "người dùng chọn workflow có story/timeline", "mở Video Production Workspace ở bước video", "hệ thống hiển thị studio toàn màn hình gồm rail cảnh, preview theo aspect ratio, panel chỉnh scene, inspector thông tin/AI/script và action Export MP4"),
    ]),
    "US-23": ("Duyệt video đã render", [
        ("Approve video", "MediaWorkflow đã có file MP4 final_video", "người dùng xem preview và bấm Approve Video", "hệ thống đánh dấu video_approved và lưu metadata review Module 4"),
        ("Phân biệt trạng thái chờ duyệt và đã duyệt", "workflow đã render xong nhưng chưa được reviewer duyệt", "người dùng xem danh sách dự án video", "hệ thống hiển thị trạng thái RENDERED là Chờ duyệt video, còn VIDEO_APPROVED/QUEUED_FOR_PUBLISHING/PUBLISHED được xếp vào nhóm đã sẵn sàng xuất bản"),
    ]),
    "US-24": ("Đưa video vào Publishing Queue", [
        ("Queue video đã duyệt", "MediaWorkflow có final_video và video_approved", "người dùng bấm Queue Post", "hệ thống tạo PublishingQueueItem với profile, caption, hashtags, scheduled_at và trạng thái phù hợp"),
        ("Chọn lịch theo chiến lược profile", "SocialProfileStrategy có schedule_days, schedule_times, timezone và approval_mode", "hệ thống tính lịch cho queue item", "queue item nhận scheduled_at phù hợp và trạng thái approved hoặc needs_approval tùy cấu hình approval_mode"),
    ]),
    "US-25": ("Theo dõi và xử lý Publishing Queue", [
        ("Lọc danh sách queue", "người dùng mở trang Approvals hoặc Schedule", "chọn profile, platform, status, date hoặc từ khóa tìm kiếm", "hệ thống trả danh sách queue item đúng bộ lọc và quyền truy cập"),
        ("Xem chi tiết nguồn và video", "queue item liên kết với ContentItem hoặc MediaWorkflow đã render", "người dùng mở dialog chi tiết", "hệ thống hiển thị bài nguồn, caption, lịch đăng và video preview được resolve từ workflow nếu article_link không phải URL video"),
    ]),
    "US-26": ("Ghi nhận và xem metrics bài đăng", [
        ("Ghi nhận SocialPostMetric", "SocialPost đã được tạo sau khi publish", "hệ thống hoặc người vận hành gửi số views, likes, comments và shares", "API lưu SocialPostMetric với mốc thời gian snapshot"),
        ("Xem overview hiệu suất", "bài đăng đã có hoặc chưa có metric", "người dùng mở trang metrics/overview", "hệ thống nhóm bài theo nội dung, lấy metric mới nhất và tính tăng trưởng nếu có dữ liệu"),
    ]),
    "US-27": ("Đăng TikTok thủ công", [
        ("Đăng bằng Direct Post", "profile TikTok có scope video.publish và queue item có video đã duyệt", "người dùng bấm Publish Now với chế độ direct", "hệ thống gọi TikTok Direct Post, tạo SocialPost và chuyển queue item sang published nếu thành công"),
        ("Đăng bằng Inbox Upload", "profile TikTok chỉ có scope video.upload", "người dùng bấm Publish Now với chế độ inbox", "hệ thống upload video vào TikTok inbox và cập nhật trạng thái queue theo phản hồi"),
    ]),
    "US-28": ("Tự động đăng theo Publish Scheduler", [
        ("Quét queue đến hạn", "Background Scheduler đang bật và có queue item queued/approved đến hạn", "scheduler chạy chu kỳ quét", "hệ thống kiểm tra profile active, token hợp lệ và điều kiện chiến lược trước khi publish"),
        ("Tôn trọng approval mode", "SocialProfileStrategy có approval_mode manual", "scheduler gặp item queued chưa được approve", "hệ thống bỏ qua item và không tự động đăng cho đến khi được approve"),
    ]),
    "US-29": ("Triển khai microservices", [
        ("Cấu hình môi trường Production", "quản trị viên chuẩn bị môi trường triển khai", "thiết lập biến môi trường cho API, ingestion engine, AI media engine, frontend, database, Kafka và storage", "các service đọc đúng cấu hình và sẵn sàng khởi chạy"),
        ("Kiểm tra health check sau deploy", "các service đã được khởi chạy trên môi trường Production", "quản trị viên kiểm tra log và health endpoint", "hệ thống xác nhận API, worker, scheduler và frontend hoạt động ổn định"),
    ]),
    "US-30": ("Quản trị & Scheduler Config", [
        ("Cấu hình khoảng thời gian Background Scheduler", "quản trị viên đăng nhập với vai trò SYSTEM_ADMIN", "điều chỉnh publish_queue_interval_minutes thành 5 phút và bấm Lưu", "hệ thống cập nhật cấu hình vào SystemSetting và Background Scheduler tự động áp dụng chu kỳ quét mới"),
        ("Xem nhật ký audit log hệ thống", "các thao tác tạo job, duyệt bài, đăng bài thủ công/tự động đã diễn ra", "quản trị viên vào trang Audit Logs", "hệ thống hiển thị danh sách nhật ký gồm người thực hiện, hành động, mục tiêu và mốc thời gian"),
    ]),
    "US-31": ("Rendered Workflow Queue Sync", [
        ("Tự tạo queue item khi mở danh sách queue", "MediaWorkflow có trạng thái RENDERED, VIDEO_APPROVED hoặc QUEUED_FOR_PUBLISHING, có final video và chưa có queued_post_id hợp lệ", "người vận hành mở danh sách Publishing Queue", "hệ thống tự tạo PublishingQueueItem trạng thái needs_approval, đặt scheduled_at mặc định sau 2 giờ và ghi queued_post_id vào metadata workflow"),
        ("Không tạo trùng queue item", "MediaWorkflow đã có metadata queued_post_id trỏ tới PublishingQueueItem còn tồn tại", "hệ thống chạy bước đồng bộ rendered workflow sang queue", "workflow bị bỏ qua và không phát sinh queue item thứ hai cho cùng một video"),
    ]),
    "US-32": ("Review Feedback Loop", [
        ("Reviewer yêu cầu chỉnh sửa từ queue", "queue item chưa đạt yêu cầu nội dung hoặc kỹ thuật", "reviewer bấm Request Changes và nhập ghi chú", "hệ thống chuyển queue item sang changes_requested, bỏ cờ video_approved và đưa MediaWorkflow liên quan về status/current_stage EDITING"),
        ("Lưu vết quyết định review vào workflow", "queue item có workflow liên quan theo queued_post_id hoặc theo cặp profile/content", "hệ thống xử lý yêu cầu changes_requested", "metadata workflow được cập nhật module4_review với decision changes_requested, reviewer, thời gian, ghi chú và previous_decision"),
    ]),
    "US-33": ("AI Series Decision Metadata", [
        ("AI tạo series mới có metadata đầy đủ", "Auto Planning hoặc Generate Video Script không tìm thấy series phù hợp cho nội dung mới", "LLM trả series_decision với series_title, series_description, series_type và total_parts", "hệ thống tạo ContentSeries active với description, type, tổng số tập và context_json ghi nguồn tạo từ auto_planning hoặc generate_video_script"),
        ("AI gán vào series theo mạch nội dung", "đã có ContentSeries active cùng chủ đề hoặc có mạch truyện liên quan", "AI đánh giá topic relevance, story continuity và reusable theme", "hệ thống gán workflow vào series phù hợp mà không bắt buộc category phải trùng tuyệt đối"),
    ]),
    "US-34": ("Real Media Preview", [
        ("Hiển thị thumbnail và avatar thật trong video workspace", "MediaWorkflow có primary_content chứa thumbnail_url, thumbnailUrl, image_url hoặc media_jsonb và profile có avatar_url", "người dùng mở Xưởng sản xuất video", "frontend hiển thị thumbnail từ dữ liệu nguồn và avatar profile, chỉ dùng fallback icon khi thật sự không có media"),
        ("Resolve video preview từ workflow cho approval queue", "queue item article_link không phải URL video nhưng có MediaWorkflow liên quan chứa FINAL_VIDEO hoặc rendered_video", "người dùng mở chi tiết queue item", "backend trả video_url từ workflow artifact/metadata/draft_json để reviewer xem đúng MP4 cần duyệt"),
    ]),
}

SPRINT_TASKS = {
    "US-01": ["Xây dựng API đăng nhập JWT và guard frontend", "Kiểm thử đăng nhập sai và route yêu cầu token"],
    "US-02": ["Áp dụng owner/privacy filters cho API dữ liệu", "Kiểm tra role SYSTEM_ADMIN trên thao tác quản trị"],
    "US-03": ["Xóa token và trạng thái phiên ở client", "Chặn truy cập route nội bộ sau khi logout"],
    "US-04": ["Xây dựng phiên TikTok OAuth QR start/status/stop", "Upsert SocialProfile với token, scope và metadata"],
    "US-05": ["Thiết kế form cấu hình SocialProfileStrategy", "Lưu cờ schedule, approval và auto publishing"],
    "US-06": ["Lưu CrawlJob và danh sách CrawlJobSource", "Phát crawl.job.created và tạo KafkaTask CRAWL_URL"],
    "US-07": ["Xây dựng detail/SSE progress cho crawl job", "Bổ sung cancel và retry theo trạng thái job"],
    "US-08": ["Tích hợp crawler theo source_type cho VNExpress/Bilibili", "Làm sạch nội dung VNExpress để bỏ sapo trùng, CTA và đoạn liên quan"],
    "US-09": ["Chuẩn hóa nội dung thô thành ContentItem", "Lưu full text/media metadata và trạng thái chất lượng"],
    "US-10": ["Gom ContentItem thành Story/Episode", "Lưu episode_order và metadata nguồn"],
    "US-11": ["Bổ sung final view, profile_matches và lọc chất lượng", "Xây dựng reprocess và mark-duplicate cho admin"],
    "US-12": ["Tạo MediaWorkflow từ content_ids/story_ids", "Tạo workflow từ crawl job candidate tốt nhất"],
    "US-13": ["Tạo PlanningRun và PlanningCandidate sau crawl completed", "Chọn candidate theo SocialProfileStrategy"],
    "US-14": ["Hiển thị kế hoạch theo profile/series", "Mở nguồn, story_data và regenerate workflow"],
    "US-15": ["Approve/reject MediaWorkflow và lưu PlanningFeedback", "Cập nhật title, series_id, caption, tags và draft_json"],
    "US-16": ["Tạo KafkaTask GENERATE_VIDEO_SCRIPT", "Sinh và lưu draft timeline từ nguồn MediaWorkflow"],
    "US-17": ["Tạo task edit/review story bằng prompt AI", "Cho phép chỉnh story_data/timeline thủ công trong workspace"],
    "US-18": ["Tích hợp Edge TTS/ElevenLabs cho voice-over", "Lưu MP3 artifact và audio clip vào timeline"],
    "US-19": ["Đồng bộ audio, text clip và visual clip", "Cảnh báo regenerate khi draft đổi sau voice/render"],
    "US-20": ["Render MP4 bằng Remotion worker", "Lưu FINAL_VIDEO artifact và cập nhật trạng thái workflow"],
    "US-21": ["CRUD ContentSeries và gán workflow", "Lưu description/current_part/total_parts và kiểm tra consistency series"],
    "US-22": ["Xây dựng kanban bốn cột đúng trạng thái workflow", "Thiết kế studio chi tiết với scene rail, preview, editor panel và inspector"],
    "US-23": ["Approve video đã render và lưu review metadata", "Tách trạng thái Chờ duyệt video khỏi nhóm đã duyệt/sẵn sàng xuất bản"],
    "US-24": ["Tạo PublishingQueueItem từ video approved", "Tính scheduled_at theo schedule_days, schedule_times, timezone và approval_mode"],
    "US-25": ["Lọc queue theo profile, trạng thái và ngày", "Hiển thị detail preview gồm source, caption, lịch và video_url resolve từ workflow"],
    "US-26": ["Ghi nhận SocialPostMetric", "Tính overview hiệu suất và growth theo SocialPost"],
    "US-27": ["Publish now qua TikTok Direct Post", "Fallback Inbox Upload theo scope video.upload"],
    "US-28": ["Quét queue đến hạn trong publish scheduler", "Tôn trọng schedule_enabled, approval_mode và auto_publish_enabled"],
    "US-31": ["Đồng bộ rendered workflow sang PublishingQueueItem khi list queue", "Lưu queued_post_id để chống tạo trùng queue item"],
    "US-32": ["Xử lý request changes từ approval queue", "Đưa workflow liên quan về EDITING và ghi module4_review"],
    "US-33": ["Mở rộng AI series_decision với description, type và total_parts", "Gán series theo topic relevance, story continuity và reusable theme"],
    "US-34": ["Trả thumbnail/avatar thật cho video workspace", "Resolve final video từ workflow artifacts/metadata/draft_json cho approval preview"],
    "US-29": ["Đóng gói API, ingestion, AI media và frontend", "Cấu hình database, Kafka, storage và biến môi trường"],
    "US-30": ["Xây dựng trang quản trị user/role/settings", "Giám sát audit log, scheduler snapshot và queue lỗi"],
}


def iter_blocks(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def delete_block(block):
    el = block._element
    el.getparent().remove(el)


def find_paragraph(document, prefix):
    for block in iter_blocks(document):
        if isinstance(block, Paragraph) and block.text.strip().startswith(prefix):
            return block
    raise ValueError(f"Cannot find paragraph starting with {prefix!r}")


def remove_between(document, start_prefix, end_prefix):
    blocks = list(iter_blocks(document))
    start_idx = end_idx = None
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if start_idx is None and text.startswith(start_prefix):
                start_idx = idx
            elif start_idx is not None and text.startswith(end_prefix):
                end_idx = idx
                break
    if start_idx is None or end_idx is None:
        raise ValueError(f"Cannot locate section {start_prefix!r} -> {end_prefix!r}")
    for block in blocks[start_idx + 1:end_idx]:
        delete_block(block)
    return find_paragraph(document, end_prefix)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)
    cell._tc.append(OxmlElement("w:p"))


def set_cell(cell, text, bold=False, font_size=None):
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def normalize_legacy_terms(document):
    replacements = [
        ("content project", "MediaWorkflow"),
        ("Content project", "MediaWorkflow"),
        ("project plan", "MediaWorkflow"),
        ("Project plan", "MediaWorkflow"),
    ]
    for paragraph in document.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in replacements:
            updated = updated.replace(old, new)
        if updated != text:
            set_paragraph_text(paragraph, updated)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                updated = text
                for old, new in replacements:
                    updated = updated.replace(old, new)
                if updated != text:
                    set_cell(cell, updated)


def update_table(table, rows):
    for ridx, values in enumerate(rows, start=1):
        if ridx >= len(table.rows):
            break
        for cidx, value in enumerate(values):
            if cidx < len(table.rows[ridx].cells):
                set_cell(table.rows[ridx].cells[cidx], value)


def replace_table(document, old_table, headers, rows, font_size=10):
    parent = old_table._element.getparent()
    new_table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    if old_table.style:
        new_table.style = old_table.style

    for cidx, header in enumerate(headers):
        set_cell(new_table.rows[0].cells[cidx], header, bold=True, font_size=font_size)
    for ridx, values in enumerate(rows, start=1):
        for cidx, value in enumerate(values):
            set_cell(new_table.rows[ridx].cells[cidx], value, font_size=font_size)

    old_table._element.addprevious(new_table._element)
    parent.remove(old_table._element)
    return new_table


def story_by_id(story_id):
    for row in STORIES:
        if row[3] == story_id:
            return row
    raise KeyError(story_id)


def make_sprint_backlog_rows(story_ids):
    rows = []
    for story_id in story_ids:
        _, _, _, uid, title, _, _, effort, _, _, _, status = story_by_id(story_id)
        tasks = SPRINT_TASKS[uid]
        first = effort // 2
        split = [first, effort - first]
        for task, est in zip(tasks, split):
            actual = est if status == "Done" else 0
            rows.append((f"{uid}: {title} ({effort}h)", task, "Nguyễn Văn Phước", status, est, actual))
    return rows


V1_SPRINT_PLAN = [
    ("Sprint 1", "Hoàn thiện nền tảng đăng nhập, phân quyền, crawl job và pipeline xử lý dữ liệu đầu vào.", "US-01, US-02, US-03, US-06, US-07, US-08, US-09, US-10, US-11", "10/08/2026", "16/08/2026", "Người dùng đăng nhập được, tạo/theo dõi crawl job, dữ liệu VNExpress/Bilibili được làm sạch và lưu thành ContentItem/Story.", "Done"),
    ("Sprint 2", "Khởi tạo MediaWorkflow và duyệt kế hoạch nội dung theo SocialProfileStrategy.", "US-12, US-13, US-14, US-15", "17/08/2026", "23/08/2026", "MediaWorkflow được tạo từ content/story/crawl job, PlanningRun chọn candidate và người dùng duyệt/từ chối kế hoạch.", "Done"),
]


V2_SPRINT_PLAN = [
    ("Sprint 3", "Sinh timeline, chỉnh sửa story, tạo voice, căn timeline, render MP4 và quản lý ContentSeries.", "US-16, US-17, US-18, US-19, US-20, US-21", "24/08/2026", "30/08/2026", "Workflow có draft story, voice artifact, timeline đồng bộ, FINAL_VIDEO và series context đủ metadata để dùng tiếp.", "Executing"),
    ("Sprint 4", "Hoàn thiện workspace/studio video, TikTok profile strategy, queue duyệt, lịch đăng, publish scheduler và feedback loop.", "US-04, US-05, US-22, US-23, US-24, US-25, US-26, US-27, US-28, US-31, US-32, US-33, US-34", "31/08/2026", "13/09/2026", "Workspace hiển thị media thật, queue tự đồng bộ video rendered, reviewer yêu cầu sửa được trả workflow về EDITING, TikTok publish hoạt động theo strategy.", "Executing"),
]


PROD_SPRINT_PLAN = [
    ("Sprint 5", "Đóng gói triển khai và hoàn thiện công cụ quản trị vận hành.", "US-29, US-30", "14/09/2026", "20/09/2026", "Các service chạy ổn định trên môi trường chính thức, admin quản lý user/settings/audit log/scheduler.", "Planned"),
]


V2_RELEASE_ROWS = [
    ("Phiên bản", "V2.0 - Sản xuất video, duyệt queue và xuất bản TikTok theo SocialProfileStrategy", "Cập nhật"),
    ("Ngày phát hành dự kiến", "13/09/2026", "Đang thực hiện"),
    ("Sprint bao phủ", "Sprint 3 và Sprint 4", "Đang thực hiện"),
    ("Phạm vi user story", "US-16, US-17, US-18, US-19, US-20, US-21, US-22, US-23, US-04, US-05, US-24, US-25, US-26, US-27, US-28, US-31, US-32, US-33, US-34", "Cập nhật theo code"),
    ("Mục tiêu phát hành", "Hoàn thiện luồng từ MediaWorkflow đến story/timeline, voice, render FINAL_VIDEO, duyệt video, tự đồng bộ PublishingQueueItem, xử lý request changes, lên lịch và publish TikTok.", "Đang thực hiện"),
    ("Tiêu chí chấp nhận", "Workspace hiển thị đúng kanban bốn cột và media thật; studio chỉnh được scene/script/config; queue không tạo trùng; request changes trả workflow về EDITING; scheduler/publish tôn trọng strategy.", "Đang kiểm thử"),
    ("Tính năng hoàn thành", "Script/scene/AI edit/TTS, subtitle sync/render MP4, ContentSeries context, AI series metadata, QA approval, queue sync, real media preview, TikTok QR/profile strategy, manual/auto publish và metrics.", "Đang hoàn thiện"),
]


V2_REVIEW_ROWS = [
    ("Frontend", "Video workspace và studio", "Kanban bốn cột, sidebar series, card thumbnail/avatar thật, studio gồm scene rail, preview aspect ratio, editor panel và inspector.", "Rà soát responsive và trạng thái disable action khi workflow thiếu artifact.", "Đạt"),
    ("Backend API", "Workspace summary, queue, approval", "API trả profile.avatar, thumbnail_url/thumbnailUrl, series profile_id/profileId; queue detail resolve video_url từ workflow.", "Bổ sung test cho fallback content_media và final video từ artifacts/metadata/draft_json.", "Đạt"),
    ("AI Media", "Script, series decision, render", "LLM schema có series_description, series_type, total_parts; tạo ContentSeries từ auto planning/script; Remotion lưu FINAL_VIDEO.", "Theo dõi chất lượng matching series theo mạch truyện và reusable theme.", "Đạt"),
    ("Publishing", "Queue sync và feedback loop", "list_user_queue tự sync workflow rendered sang queue; queued_post_id chống trùng; request changes cập nhật module4_review và trả workflow về EDITING.", "Kiểm tra idempotency khi queue item bị xóa hoặc workflow có nhiều artifact final.", "Đạt"),
    ("Scheduler/TikTok", "Lên lịch và đăng bài", "Publish scheduler kiểm tra schedule_enabled, approval_mode, auto_publish_enabled; hỗ trợ Direct Post và Inbox Upload theo scope.", "Theo dõi lỗi token/scope và retry publish failed.", "Đạt"),
]


V2_RETRO_ROWS = [
    ("Điểm làm tốt", "Luồng V2 đã rõ trạng thái hơn: RENDERED là Chờ duyệt video, còn VIDEO_APPROVED/QUEUED_FOR_PUBLISHING/PUBLISHED tách khỏi giai đoạn biên tập."),
    ("Điểm làm tốt", "Workspace bỏ dữ liệu demo và ưu tiên media thật từ ContentItem/MediaWorkflow, giúp reviewer kiểm tra đúng video cần duyệt."),
    ("Vấn đề còn tồn tại", "Một số trạng thái và metadata workflow phụ thuộc nhiều nguồn fallback nên cần test regression kỹ để tránh queue trùng hoặc mất video_url."),
    ("Bài học", "Các user story cần bám entity code thực tế như MediaWorkflow, ContentSeries, PublishingQueueItem, KafkaTask và SocialProfileStrategy thay vì mô tả nghiệp vụ quá chung."),
    ("Hành động tiếp theo", "Bổ sung test cho sync_rendered_workflows_to_queue, mark_queue_workflow_changes_requested, series_decision metadata và workspace summary media fallback."),
]


TABLE_CAPTIONS = [
    "Bảng 2.1. Product Roadmap của dự án",
    "Bảng 2.2. Danh sách các user story trong Product Backlog",
    "Bảng 2.3. Ước tính độ phức tạp câu chuyện người dùng",
    "Bảng 2.4. Phân bổ câu chuyện người dùng theo Release và Sprint",
    "Bảng 2.5. Thứ tự ưu tiên phát triển User Story",
    "Bảng 2.6. Sprint Planning cho phiên bản V1.0",
    "Bảng 2.7. Sprint Backlog cho phiên bản V1.0",
    "Bảng 2.8. Kiểm thử chức năng xác thực và phân quyền",
    "Bảng 2.9. Kiểm thử chức năng crawl job và thu thập dữ liệu",
    "Bảng 2.10. Kiểm thử chức năng chuẩn hóa dữ liệu và story grouping",
    "Bảng 2.11. Kiểm thử chức năng lập kế hoạch nội dung và MediaWorkflow",
    "Bảng 2.12. Thông tin phát hành phiên bản V1.0",
    "Bảng 2.13. Sprint Review cho phiên bản V1.0",
    "Bảng 2.14. Sprint Retrospective cho phiên bản V1.0",
    "Bảng 2.15. Sprint Planning cho phiên bản V2.0",
    "Bảng 2.16. Sprint Backlog cho phiên bản V2.0",
    "Bảng 2.17. Thông tin phát hành phiên bản V2.0",
    "Bảng 2.18. Sprint Review cho phiên bản V2.0",
    "Bảng 2.19. Sprint Retrospective cho phiên bản V2.0",
    "Bảng 2.20. Sprint Planning cho giai đoạn Production",
    "Bảng 2.21. Sprint Backlog cho giai đoạn Production",
    "Bảng 2.22. Kiểm thử các chức năng giai đoạn Production",
]


def insert_after_heading(anchor, lines):
    for line in lines:
        paragraph = anchor.insert_paragraph_before()
        if isinstance(line, tuple):
            label, body = line
            run = paragraph.add_run(label)
            run.bold = True
            paragraph.add_run(body)
        else:
            paragraph.add_run(line)


def make_story_section(story_ids):
    lines = []
    for story_id in story_ids:
        _, _, _, uid, title, *_ = story_by_id(story_id)
        feature, scenarios = GHERKIN_SCENARIOS[uid]
        lines.extend([
            (f"{uid}: {title}", ""),
            ("Feature: ", feature),
        ])
        for index, (name, given, when, then) in enumerate(scenarios, start=1):
            lines.extend([
                (f"Scenario {index}: ", name),
                ("Given ", given),
                ("When ", when),
                ("Then ", then),
            ])
        lines.append("")
    return lines


def set_caption_text(paragraph, text):
    set_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        run.bold = True
    try:
        paragraph.style = "Caption"
    except KeyError:
        pass


def insert_caption_before_table(document, table, text):
    paragraph_el = OxmlElement("w:p")
    table._element.addprevious(paragraph_el)
    paragraph = Paragraph(paragraph_el, document)
    paragraph.add_run(text)
    set_caption_text(paragraph, text)


def previous_nonempty_paragraph(blocks, table_index):
    for index in range(table_index - 1, -1, -1):
        block = blocks[index]
        if isinstance(block, Paragraph) and block.text.strip():
            return block
        if isinstance(block, Table):
            break
    return None


def ensure_table_captions(document):
    blocks = list(iter_blocks(document))
    table_index = 0
    for block_index, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue
        if table_index >= len(TABLE_CAPTIONS):
            raise ValueError("More tables than expected; add captions to TABLE_CAPTIONS.")
        caption = TABLE_CAPTIONS[table_index]
        previous = previous_nonempty_paragraph(blocks, block_index)
        if previous is not None and previous.text.strip().startswith("Bảng "):
            set_caption_text(previous, caption)
        else:
            insert_caption_before_table(document, block, caption)
        table_index += 1
    if table_index != len(TABLE_CAPTIONS):
        raise ValueError(f"Expected {len(TABLE_CAPTIONS)} tables but found {table_index}.")


def remove_tabs_from_table_cells(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    ppr = paragraph._p.pPr
                    if ppr is not None:
                        for tabs in list(ppr.xpath("./w:tabs")):
                            tabs.getparent().remove(tabs)
                    for run in paragraph.runs:
                        if "\t" in run.text:
                            run.text = run.text.replace("\t", " ")
                        for tab in list(run._element.xpath(".//w:tab")):
                            tab.getparent().remove(tab)


def main():
    doc = Document(SOURCE)
    tables = doc.tables

    backlog_rows = [(tt, theme, epic, uid, title) for tt, theme, epic, uid, title, *_ in STORIES]
    estimate_rows = [(uid, desc, sp, effort) for _, _, _, uid, _, desc, sp, effort, *_ in STORIES]
    allocation_rows = [(release, sprint, uid, title, "Implementation", status) for _, _, _, uid, title, _, _, _, release, sprint, _, status in STORIES]
    priority_rows = [(uid, sp, effort, sprint, priority) for _, _, _, uid, _, _, sp, effort, _, sprint, priority, _ in STORIES]

    replace_table(doc, tables[1], ["TT", "Theme", "Epic", "Mã PBI", "Tên PBI"], backlog_rows, font_size=10)
    replace_table(doc, tables[2], ["Mã PBI", "User story", "Story Point", "Effort (giờ)"], estimate_rows, font_size=10)
    replace_table(doc, tables[3], ["Release", "Sprint", "Mã PBI", "Tên PBI", "Giai đoạn", "Trạng thái"], allocation_rows, font_size=10)
    replace_table(doc, tables[4], ["Mã PBI", "Story Point", "Effort (giờ)", "Sprint", "Priority"], priority_rows, font_size=10)
    replace_table(doc, tables[5], ["Sprint", "Mục tiêu sprint", "PBI chính", "Ngày bắt đầu", "Ngày kết thúc", "Definition of Done", "Trạng thái"], V1_SPRINT_PLAN, font_size=9)
    replace_table(doc, tables[6], ["Product Backlog Item", "Sprint Task", "Owner", "Status", "Est. Effort (h)", "Actual (h)"], make_sprint_backlog_rows(V1_ORDER), font_size=9)
    replace_table(doc, tables[14], ["Sprint", "Mục tiêu sprint", "PBI chính", "Ngày bắt đầu", "Ngày kết thúc", "Definition of Done", "Trạng thái"], V2_SPRINT_PLAN, font_size=9)
    replace_table(doc, tables[15], ["Product Backlog Item", "Sprint Task", "Owner", "Status", "Est. Effort (h)", "Actual (h)"], make_sprint_backlog_rows(V2_ORDER), font_size=9)
    replace_table(doc, tables[21], ["Hạng mục", "Nội dung", "Trạng thái"], V2_RELEASE_ROWS, font_size=9)
    replace_table(doc, tables[22], ["Nhóm", "Phạm vi review", "Kết quả đạt được", "Điều chỉnh sau review", "Trạng thái"], V2_REVIEW_ROWS, font_size=9)
    replace_table(doc, tables[23], ["Loại ghi nhận", "Nội dung retrospective"], V2_RETRO_ROWS, font_size=10)
    replace_table(doc, tables[24], ["Sprint", "Mục tiêu sprint", "PBI chính", "Ngày bắt đầu", "Ngày kết thúc", "Definition of Done", "Trạng thái"], PROD_SPRINT_PLAN, font_size=9)
    replace_table(doc, tables[25], ["Product Backlog Item", "Sprint Task", "Owner", "Status", "Est. Effort (h)", "Actual (h)"], make_sprint_backlog_rows(PROD_ORDER), font_size=9)

    v1_plan_anchor = remove_between(doc, "2.5.1.1. Mục tiêu", "Bảng 2.6.")
    insert_after_heading(v1_plan_anchor, [
        "Phiên bản V1.0 tập trung xây dựng nền tảng bắt buộc của hệ thống SocialContentHub: xác thực, phân quyền, crawl đa nguồn, chuẩn hóa dữ liệu, gom nhóm story và lập kế hoạch nội dung bằng MediaWorkflow.",
        "Sau khi đối chiếu với code hiện tại, phạm vi V1.0 vẫn giữ vai trò là luồng nền: người dùng tạo CrawlJob, worker xử lý nguồn VNExpress/Bilibili qua KafkaTask, dữ liệu được làm sạch và lưu thành ContentItem/Story, sau đó PlanningRun chọn candidate theo SocialProfileStrategy để người dùng duyệt.",
        "Các điểm đã chỉnh trong báo cáo gồm việc bỏ cách gọi cũ content project/project plan, thay bằng MediaWorkflow, và bổ sung yêu cầu làm sạch nội dung VNExpress để loại sapo trùng, link bài liên quan, CTA Xem thêm và artifact tác giả.",
    ])

    v2_plan_anchor = remove_between(doc, "2.5.2.1. Mục tiêu", "Bảng 2.11.")
    insert_after_heading(v2_plan_anchor, [
        "Phiên bản V2.0 phản ánh các thay đổi lớn nhất của code hiện tại: sản xuất video theo MediaWorkflow, studio chỉnh timeline, voice, render MP4, quản lý ContentSeries, duyệt video, đồng bộ Publishing Queue và xuất bản TikTok theo chiến lược từng Social Profile.",
        "So với bản mô tả cũ, V2.0 không còn chỉ dừng ở preview/render. Workspace hiện có kanban bốn cột, màn hình studio chi tiết với scene rail, preview theo aspect ratio, panel chỉnh scene, inspector thông tin/AI/script; trạng thái RENDERED được hiểu là Chờ duyệt video, còn VIDEO_APPROVED/QUEUED_FOR_PUBLISHING/PUBLISHED thuộc nhóm đã sẵn sàng xuất bản.",
        "Báo cáo cũng bổ sung các user story mới cho các luồng code vừa xuất hiện: tự đồng bộ workflow đã render sang PublishingQueueItem, trả workflow về EDITING khi reviewer yêu cầu chỉnh sửa, lưu metadata series do AI đề xuất và hiển thị thumbnail/avatar/video preview từ dữ liệu thật.",
    ])

    prod_plan_anchor = remove_between(doc, "2.5.3.1. Mục tiêu", "Bảng 2.16.")
    insert_after_heading(prod_plan_anchor, [
        "Giai đoạn Production tập trung đóng gói triển khai và vận hành hệ thống sau khi các luồng nghiệp vụ V1.0/V2.0 đã ổn định.",
        "Phạm vi chính gồm cấu hình môi trường cho API service, data-ingestion-engine, ai-media-engine, frontend, Kafka, database, storage, đồng thời hoàn thiện trang quản trị người dùng, SystemSetting, audit log và snapshot scheduler.",
        "Các user story Production được giữ riêng để tránh nhầm với luồng sản xuất nội dung đang phát triển ở V2.0.",
    ])

    v1_anchor = remove_between(doc, "2.5.1.2. Xây dựng kịch bản", "2.5.1.3. Phát triển tính năng")
    insert_after_heading(v1_anchor, make_story_section(V1_ORDER))

    v2_anchor = remove_between(doc, "2.5.2.2. Xây dựng kịch bản", "2.5.2.3. Phát triển tính năng")
    insert_after_heading(v2_anchor, make_story_section(V2_ORDER))

    v2_dev_anchor = remove_between(doc, "2.5.2.3. Phát triển tính năng", "2.5.2.4. Phát hành")
    insert_after_heading(v2_dev_anchor, [
        "Ở phiên bản V2.0, phần phát triển tính năng được điều chỉnh theo kiến trúc code hiện tại: MediaWorkflow là đối tượng trung tâm, liên kết với SocialProfile, ContentItem, ContentSeries, KafkaTask, artifact voice/video và PublishingQueueItem. Luồng không còn mô tả chung là tạo content project, mà đi theo chuỗi rõ ràng: sinh timeline, chỉnh story, tạo voice, render MP4, duyệt video, đồng bộ queue, lên lịch và đăng TikTok.",
        ("Nhóm US-16 đến US-20: ", "AI Media Engine nhận MediaWorkflow có nguồn nội dung hợp lệ, sinh story/timeline gồm scene, subtitle, voice text, visual prompt và timing; người dùng có thể chỉnh bằng AI hoặc thủ công; voice-over được tạo qua provider phù hợp; timeline được căn lại theo audio/subtitle/visual; cuối cùng Remotion Worker xuất FINAL_VIDEO và cập nhật trạng thái workflow sang nhóm rendered/chờ duyệt."),
        ("Nhóm US-21 và US-33: ", "ContentSeries được quản lý như ngữ cảnh dài kỳ của kênh. Code hiện tại cho phép AI trả series_decision có series_description, series_type và total_parts; khi tạo series mới, hệ thống lưu description/type/tổng số tập và context_json ghi nguồn tạo. Khi đã có series phù hợp, AI gán workflow theo topic relevance, story continuity và reusable theme, không còn phụ thuộc cứng vào việc category phải trùng tuyệt đối."),
        ("Nhóm US-22 và US-34: ", "Video workspace đã chuyển sang mô hình xưởng sản xuất/studio. Danh sách dự án dùng kanban bốn cột gồm Draft kịch bản, Đang biên tập & Voice, Đang render MP4 và Chờ duyệt video; card ưu tiên thumbnail từ primary_content hoặc media_jsonb và avatar SocialProfile, chỉ dùng fallback icon khi không có media thật. Workspace chi tiết hiển thị rail cảnh, preview theo aspect ratio, panel chỉnh scene, inspector thông tin/AI/script, nút reload và lệnh xuất MP4."),
        ("Nhóm US-23 đến US-25, US-31 và US-32: ", "Luồng duyệt/xuất bản đã được tách rõ. RENDERED nghĩa là video đang chờ duyệt; VIDEO_APPROVED, QUEUED_FOR_PUBLISHING và PUBLISHED thuộc nhóm đã sẵn sàng xuất bản hoặc đã xử lý. Khi người dùng mở queue, service tự đồng bộ MediaWorkflow đã render thành PublishingQueueItem nếu chưa có queued_post_id hợp lệ, đặt trạng thái needs_approval và lịch mặc định sau 2 giờ. Nếu reviewer yêu cầu chỉnh sửa, queue item chuyển changes_requested, workflow liên quan quay lại EDITING và metadata module4_review lưu decision, reviewer, thời gian, ghi chú và previous_decision."),
        ("Nhóm US-26 đến US-28: ", "Sau khi queue item được duyệt hoặc đến hạn, hệ thống ghi nhận SocialPost/SocialPostMetric, hỗ trợ Direct Post hoặc Inbox Upload tùy scope TikTok, đồng thời publish scheduler chỉ tự đăng khi profile active, token hợp lệ và SocialProfileStrategy cho phép schedule_enabled, approval_mode/auto_publish_enabled phù hợp."),
        "Các kiểm thử của V2.0 cần tập trung vào tính đúng trạng thái workflow, idempotency khi đồng bộ queue, khả năng resolve final video từ artifact/metadata/draft_json, đường quay lại EDITING khi request changes, và tính nhất quán dữ liệu hiển thị giữa backend workspace summary với frontend studio.",
    ])

    prod_anchor = remove_between(doc, "2.5.3.2. Xây dựng kịch bản", "2.5.3.3. Phát triển tính năng")
    insert_after_heading(prod_anchor, make_story_section(PROD_ORDER))

    note_anchor = find_paragraph(doc, "Bảng 2.2. Danh sách các user story")
    insert_after_heading(note_anchor, [
        ("Ghi chú cập nhật 29/08/2026: ", "Danh sách user story dưới đây đã được điều chỉnh theo code hiện tại: dữ liệu được thu thập bằng crawl job/KafkaTask, lập kế hoạch qua MediaWorkflow và PlanningRun, sản xuất video qua studio timeline/script/voice/render, video render xong được đồng bộ sang Publishing Queue để duyệt, reviewer có thể trả workflow về EDITING khi yêu cầu chỉnh sửa, sau đó hệ thống lên lịch và đăng TikTok theo SocialProfileStrategy."),
    ])

    normalize_legacy_terms(doc)
    ensure_table_captions(doc)
    remove_tabs_from_table_cells(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
